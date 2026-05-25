from __future__ import annotations

from pathlib import Path

from docx import Document

from build_sign_language_tech_doc import (
    FILL_OK,
    FILL_WARN,
    add_bullet,
    add_callout,
    add_code,
    add_flow,
    add_h1,
    add_h2,
    add_h3,
    add_numbered_list,
    add_para,
    add_table,
    add_title,
    configure_document,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "sign-language-concepts-complete-ko.docx"


def bullets(doc, items):
    for item in items:
        add_bullet(doc, item)


def numbered(doc, items):
    add_numbered_list(doc, items)


def concept_block(doc, title, body, project_link=None, common_mistake=None):
    add_h3(doc, title)
    add_para(doc, body)
    if project_link:
        add_para(doc, project_link, bold_lead="이 프로젝트에서는: ")
    if common_mistake:
        add_para(doc, common_mistake, bold_lead="초보자가 헷갈리기 쉬운 점: ")


def build_doc():
    doc = Document()
    configure_document(doc)

    add_title(
        doc,
        "수어지교 개념 완전판",
        "프로젝트 안에 등장하는 거의 모든 기술 개념을 기초부터 자세히 설명한 초보자용 참고서",
    )
    add_callout(
        doc,
        "문서의 목적",
        "이 문서는 기능 문서가 아니라 개념 문서다. 기존 기술 흐름 문서와 개념 심화 문서에 나온 용어를 더 넓게 확장해, 개발을 처음 배우는 사람이 프로젝트를 읽기 전에 필요한 배경지식을 한곳에서 볼 수 있도록 구성했다.",
        FILL_OK,
    )
    add_para(
        doc,
        "구성은 기초 소프트웨어 개념에서 시작해 네트워크, 프런트엔드, 백엔드, 인증과 보안, 데이터 저장소, 영상 저장소, 자연어 처리, 프로젝트 연결 흐름, 디버깅 개념으로 이어진다. 각 개념은 정의, 왜 필요한지, 이 프로젝트에서 어디에 나타나는지까지 함께 설명한다."
    )

    add_h1(doc, "1. 읽는 순서와 큰 그림")
    add_para(
        doc,
        "처음부터 모든 개념을 완벽히 외울 필요는 없다. 큰 그림을 먼저 잡고, 코드에서 모르는 단어가 보일 때 이 문서로 돌아오는 방식이 가장 효율적이다."
    )
    numbered(
        doc,
        [
            "개발 자체가 낯설다면 2장과 3장을 먼저 읽는다.",
            "API나 서버 통신이 헷갈리면 4장을 읽는다.",
            "React Native 화면 코드가 어렵다면 5장을 읽는다.",
            "Spring Boot 백엔드 코드가 어렵다면 6장을 읽는다.",
            "로그인, JWT, 보안이 어렵다면 7장을 읽는다.",
            "Firestore와 Storage가 헷갈리면 8장과 9장을 읽는다.",
            "번역 로직의 토큰, 형태소, 정규화가 어렵다면 10장을 읽는다.",
            "실제 프로젝트에서 개념이 연결되는 모습을 보고 싶다면 11장과 12장을 읽는다.",
        ],
    )
    add_flow(
        doc,
        "수어지교를 보는 가장 큰 지도",
        [
            "사용자",
            "  ↓ 화면 입력",
            "프런트엔드: React Native / Expo",
            "  ↓ HTTP API 요청",
            "백엔드: Spring Boot",
            "  ├─ 인증: Google idToken, JWT",
            "  ├─ 번역: 토큰화, 정규화, 사전 매칭",
            "  ├─ 퀴즈: 세션 생성, 정답 판정",
            "  └─ 사용자 기록: 오답, 북마크, 학습량",
            "  ↓",
            "Firestore: 구조화된 데이터",
            "Firebase Storage: 수어 영상 파일",
            "외부 API: ETRI, 우리말샘, KRDict, OpenAI",
        ],
    )

    add_h1(doc, "2. 소프트웨어 개발 기초 개념")
    concept_block(
        doc,
        "프로그램",
        "프로그램은 컴퓨터가 어떤 일을 하도록 만든 명령의 묶음이다. 사용자가 앱을 실행하면 프로그램이 입력을 받고 계산하고 결과를 보여 준다. 프로그램은 사람이 읽는 소스 코드로 작성되고, 실행 환경에서 실제로 동작한다.",
        "프런트엔드는 Expo 앱이라는 프로그램이고, 백엔드는 Spring Boot 서버라는 별도 프로그램이다.",
    )
    concept_block(
        doc,
        "소스 코드",
        "소스 코드는 개발자가 작성하는 명령문이다. 컴퓨터가 바로 이해하는 기계어와 달리 사람이 읽을 수 있는 형태로 작성된다. Java, TypeScript, JavaScript 같은 언어로 작성된 파일들이 소스 코드다.",
        "frontendcodes 폴더에는 주로 TypeScript/JavaScript 코드가 있고, backend 폴더에는 Java 코드가 있다.",
    )
    concept_block(
        doc,
        "저장소(Repository)",
        "저장소는 프로젝트 파일을 모아 두고 변경 이력을 관리하는 공간이다. 보통 Git 저장소라고 부른다. 저장소에는 소스 코드, 설정 파일, 문서, 리소스가 함께 들어간다.",
        "현재 작업 중인 Sign-Language 폴더가 이 프로젝트의 저장소다.",
    )
    concept_block(
        doc,
        "폴더 구조",
        "큰 프로젝트는 역할별로 폴더를 나눈다. 화면 코드는 화면 폴더에, 서버 코드는 서버 폴더에, 설정은 설정 위치에 둔다. 구조를 모르면 파일을 찾는 데 시간이 오래 걸린다.",
        "frontendcodes는 프런트엔드, backend는 백엔드, docs는 지금 만든 문서들이 있는 폴더다.",
    )
    concept_block(
        doc,
        "런타임(Runtime)",
        "런타임은 프로그램을 실제로 실행하는 환경이다. JavaScript는 Node.js나 브라우저 같은 런타임에서 실행되고, Java는 JVM 위에서 실행된다.",
        "프런트엔드는 Node.js와 Expo 도구를 사용해 실행하고, 백엔드는 Java 17과 JVM 위에서 실행된다.",
    )
    concept_block(
        doc,
        "의존성(Dependency)",
        "의존성은 프로젝트가 직접 만들지 않고 가져다 쓰는 외부 라이브러리다. 개발자는 모든 기능을 직접 만들지 않고 검증된 패키지를 가져와 시간을 줄인다.",
        "프런트엔드는 package.json으로 Expo, React Native, expo-router 등을 관리하고, 백엔드는 build.gradle로 Spring Boot와 Firebase 관련 라이브러리를 관리한다.",
        "의존성이 많아질수록 버전 충돌이나 업데이트 관리가 중요해진다.",
    )
    concept_block(
        doc,
        "패키지 관리자",
        "패키지 관리자는 의존성을 설치하고 버전을 맞추는 도구다. JavaScript 생태계에서는 npm이 흔하고, Java 프로젝트에서는 Gradle이나 Maven이 흔하다.",
        "frontendcodes는 npm install로 패키지를 설치하고, backend는 Gradle Wrapper인 gradlew.bat로 빌드와 실행을 한다.",
    )
    concept_block(
        doc,
        "빌드(Build)",
        "빌드는 사람이 작성한 소스 코드와 리소스를 실행 가능한 형태로 준비하는 과정이다. 컴파일, 번들링, 리소스 복사, 설정 반영 같은 작업이 포함될 수 있다.",
        "백엔드는 Gradle이 Java 코드를 컴파일하고 Spring Boot 앱을 실행한다. 프런트엔드는 Expo가 앱 번들링과 실행 환경 구성을 담당한다.",
    )
    concept_block(
        doc,
        "개발 환경과 운영 환경",
        "개발 환경은 개발자가 테스트하고 수정하는 환경이고, 운영 환경은 실제 사용자가 접속하는 환경이다. 개발 환경에서는 편의를 위해 보안을 완화할 때가 있지만, 운영 환경에서는 보안과 안정성이 훨씬 중요하다.",
        "현재 SecurityConfig에서 여러 API가 열려 있는 것은 개발 편의에 가까우며, 실제 운영에서는 JWT 검증과 권한 확인을 강화해야 한다.",
    )
    concept_block(
        doc,
        "환경 변수",
        "환경 변수는 코드 밖에서 프로그램에 주입하는 설정 값이다. 비밀 키나 배포 환경별 주소를 코드에 직접 적지 않고 환경 변수로 넣으면 보안과 운영이 쉬워진다.",
        "JWT_SECRET, GOOGLE_CLIENT_ID, FIREBASE_STORAGE_BUCKET, OPENAI_API_KEY 같은 값이 환경 변수로 들어간다.",
        "환경 변수가 누락되면 코드가 맞아도 서버가 시작되지 않거나 외부 서비스 연결이 실패할 수 있다.",
    )
    add_table(
        doc,
        ["개념", "프런트엔드 예시", "백엔드 예시"],
        [
            ("언어", "TypeScript / JavaScript", "Java"),
            ("실행 도구", "Expo, Node.js", "Java 17, Gradle"),
            ("설정 파일", "package.json, app config", "build.gradle, application.yml"),
            ("주요 책임", "화면과 상태", "로직과 데이터 처리"),
            ("외부 의존성", "expo-router, expo-video", "Spring Boot, Firebase Admin, JJWT"),
        ],
        [1.3, 2.6, 2.6],
    )

    add_h1(doc, "3. 웹과 앱 서비스의 구조 개념")
    concept_block(
        doc,
        "사용자",
        "사용자는 서비스를 이용하는 사람이다. 개발 관점에서는 사용자가 어떤 화면을 보고 어떤 행동을 하는지부터 흐름이 시작된다.",
        "사용자는 번역기에 문장을 입력하거나 퀴즈 선택지를 누른다.",
    )
    concept_block(
        doc,
        "클라이언트",
        "클라이언트는 사용자의 기기에서 실행되며 서버에게 요청을 보내는 프로그램이다. 브라우저, 모바일 앱, 데스크톱 앱이 모두 클라이언트가 될 수 있다.",
        "이 프로젝트에서는 Expo/React Native 앱이 클라이언트다.",
    )
    concept_block(
        doc,
        "서버",
        "서버는 클라이언트의 요청을 받아 처리하고 응답을 돌려주는 프로그램이다. 서버는 보통 항상 켜져 있으면서 요청을 기다린다.",
        "backend 폴더의 Spring Boot 애플리케이션이 서버다.",
    )
    concept_block(
        doc,
        "프런트엔드",
        "프런트엔드는 사용자가 보는 화면과 상호작용을 담당한다. 그러나 단순 디자인만이 아니라 입력, 상태, API 호출, 로딩, 에러, 결과 표시까지 포함한다.",
        "translator.tsx, quiz.tsx, login.tsx 같은 화면 파일들이 프런트엔드의 중심이다.",
    )
    concept_block(
        doc,
        "백엔드",
        "백엔드는 화면 뒤에서 실제 규칙을 처리한다. 데이터 검증, 저장, 인증, 외부 API 호출, 비즈니스 로직이 들어간다.",
        "TranslationService, QuizService, UserService, AuthController 같은 코드가 백엔드의 중심이다.",
    )
    concept_block(
        doc,
        "데이터 저장소",
        "데이터 저장소는 앱이 계속 기억해야 하는 값을 저장하는 공간이다. 사용자 정보, 퀴즈 문제, 북마크, 오답 기록처럼 앱을 껐다 켜도 남아야 하는 정보가 여기에 들어간다.",
        "Firestore가 구조화된 데이터 저장소 역할을 한다.",
    )
    concept_block(
        doc,
        "파일 저장소",
        "파일 저장소는 이미지, 동영상, 음성처럼 큰 파일을 저장하는 공간이다. 데이터베이스에 큰 파일 자체를 넣는 것보다 별도 파일 저장소를 쓰는 것이 일반적이다.",
        "Firebase Storage가 수어 영상 파일 저장소 역할을 한다.",
    )
    add_flow(
        doc,
        "화면과 서버가 분리되는 이유",
        [
            "화면은 사용자 경험에 집중한다",
            "서버는 공통 규칙과 보안에 집중한다",
            "데이터 저장소는 장기 보관에 집중한다",
            "파일 저장소는 큰 파일 제공에 집중한다",
        ],
    )

    add_h1(doc, "4. 네트워크, API, HTTP 개념")
    concept_block(
        doc,
        "네트워크",
        "네트워크는 서로 다른 컴퓨터나 프로그램이 데이터를 주고받는 연결이다. 프런트엔드와 백엔드가 같은 컴퓨터에서 실행되어도 서로 다른 프로그램이면 네트워크 요청 형태로 대화할 수 있다.",
        "프런트엔드는 http://localhost:8080 같은 주소로 백엔드에 요청한다.",
    )
    concept_block(
        doc,
        "localhost",
        "localhost는 지금 실행 중인 자기 컴퓨터를 가리키는 이름이다. 웹에서 localhost:8080이라고 하면 내 컴퓨터의 8080 포트에서 실행 중인 서버를 뜻한다.",
        "웹 브라우저에서 테스트할 때는 localhost가 PC를 가리키지만, 실제 모바일 기기에서는 기기 자신을 가리킬 수 있어 LAN IP가 필요할 수 있다.",
    )
    concept_block(
        doc,
        "포트(Port)",
        "포트는 한 컴퓨터 안에서 어떤 프로그램과 통신할지 구분하는 번호다. 같은 컴퓨터에서 여러 서버가 실행될 수 있으므로 포트 번호가 필요하다.",
        "백엔드는 기본적으로 8080 포트에서 실행된다.",
    )
    concept_block(
        doc,
        "URL",
        "URL은 인터넷 자원의 주소다. 프로토콜, 호스트, 포트, 경로, 쿼리 문자열 같은 요소로 이루어진다.",
        "http://localhost:8080/translate에서 /translate는 번역 API 경로다.",
    )
    concept_block(
        doc,
        "API",
        "API는 서로 다른 프로그램이 대화하기 위한 약속이다. 어떤 경로로 어떤 데이터를 보내면 어떤 응답이 돌아오는지 정해 둔 계약이다.",
        "POST /translate는 문장을 보내면 번역 결과를 받는 API다.",
    )
    concept_block(
        doc,
        "REST",
        "REST는 HTTP 메서드와 URL을 사용해 자원을 다루는 API 설계 방식이다. 엄격한 하나의 라이브러리가 아니라 설계 스타일에 가깝다.",
        "GET /api/users/{uid}, POST /api/quiz/answer 같은 경로는 REST 스타일에 가깝다.",
    )
    concept_block(
        doc,
        "HTTP 메서드",
        "HTTP 메서드는 요청의 의도를 표현한다. GET은 조회, POST는 생성이나 처리 요청, PATCH는 일부 수정, DELETE는 삭제에 자주 사용된다.",
        "퀴즈 세션 조회는 GET, 정답 제출은 POST, 프로필 이미지는 PATCH, 북마크 삭제는 DELETE가 어울린다.",
    )
    add_table(
        doc,
        ["메서드", "일반적 의미", "프로젝트 예시"],
        [
            ("GET", "데이터 조회", "/api/quiz/session"),
            ("POST", "새 처리 요청 또는 생성", "/translate, /api/quiz/answer"),
            ("PATCH", "일부 수정", "/api/users/{uid}/profile-image"),
            ("DELETE", "삭제 또는 숨김 처리", "/api/users/{uid}/translator-bookmarks/{id}"),
        ],
        [1.0, 2.35, 3.15],
    )
    concept_block(
        doc,
        "요청(Request)",
        "요청은 클라이언트가 서버로 보내는 메시지다. 요청에는 메서드, URL, 헤더, 바디가 들어갈 수 있다.",
        "번역 버튼을 누르면 프런트엔드는 text가 담긴 JSON 바디를 서버로 보낸다.",
    )
    concept_block(
        doc,
        "응답(Response)",
        "응답은 서버가 클라이언트에게 돌려주는 메시지다. 상태 코드, 헤더, 바디가 포함될 수 있다.",
        "번역 API는 normalizedTokens, clips, unknown, noVideoWords 같은 값이 들어 있는 JSON을 응답한다.",
    )
    concept_block(
        doc,
        "상태 코드",
        "상태 코드는 서버가 요청 처리 결과를 숫자로 표현한 것이다. 200대는 성공, 400대는 클라이언트 요청 문제, 500대는 서버 내부 문제에 가깝다.",
        "로그인 시 사용자가 없으면 404와 needsSignup=true를 반환하는 흐름이 있다.",
    )
    concept_block(
        doc,
        "헤더(Header)",
        "헤더는 요청이나 응답에 붙는 부가 정보다. 인증 토큰, 콘텐츠 타입, 캐시 정보 같은 메타데이터가 들어간다.",
        "JWT를 Authorization 헤더에 Bearer 토큰으로 넣는 방식이 일반적이다.",
    )
    concept_block(
        doc,
        "바디(Body)",
        "바디는 요청이나 응답의 실제 데이터 본문이다. POST 요청에서 JSON 데이터를 보낼 때 바디를 사용한다.",
        "POST /translate의 바디에는 사용자가 입력한 text가 들어간다.",
    )
    concept_block(
        doc,
        "JSON",
        "JSON은 키와 값으로 이루어진 데이터 형식이다. 사람이 읽기 쉽고 JavaScript와 Java에서 다루기 쉬워 API 통신에서 많이 사용된다.",
        "프런트엔드와 백엔드는 대부분 JSON으로 요청과 응답을 주고받는다.",
    )
    add_code(
        doc,
        """{
  "text": "오늘 학교에 가요"
}""",
    )
    concept_block(
        doc,
        "CORS",
        "CORS는 브라우저가 서로 다른 출처의 프런트엔드와 백엔드 통신을 제한하는 보안 규칙이다. 서버가 허용한 출처만 브라우저가 요청을 통과시킨다.",
        "프런트엔드와 백엔드가 서로 다른 포트에서 실행되므로 백엔드 WebConfig나 보안 설정에서 CORS를 신경 써야 한다.",
        "CORS 오류는 서버가 죽은 것처럼 보이지만 실제로는 브라우저 보안 정책 때문에 막힌 경우가 많다.",
    )
    concept_block(
        doc,
        "API 계약",
        "API 계약은 프런트엔드와 백엔드가 약속한 입력과 출력 구조다. 계약이 바뀌면 양쪽 코드를 함께 맞춰야 한다.",
        "TranslateResponse에 clips와 items가 들어간다는 사실은 프런트엔드가 번역 결과 UI를 구성하는 데 필요한 계약이다.",
    )

    add_h1(doc, "5. 프런트엔드 개념 완전 정리")
    concept_block(
        doc,
        "React",
        "React는 상태를 기반으로 화면을 선언적으로 그리는 UI 라이브러리다. 개발자는 '지금 상태라면 화면이 어떻게 보여야 하는가'를 코드로 표현한다.",
        "이 프로젝트의 화면은 React 방식으로 작성된 컴포넌트들이다.",
    )
    concept_block(
        doc,
        "React Native",
        "React Native는 React 문법으로 모바일 앱 UI를 만들 수 있게 해 주는 프레임워크다. 웹의 div 대신 View, span 대신 Text 같은 네이티브 컴포넌트를 사용한다.",
        "frontendcodes의 화면은 React Native 컴포넌트로 구성된다.",
    )
    concept_block(
        doc,
        "Expo",
        "Expo는 React Native 앱 개발을 쉽게 해 주는 도구와 플랫폼이다. 개발 서버 실행, 빌드, 네이티브 기능 접근, 라우팅 같은 작업을 편하게 해 준다.",
        "npx expo start로 프런트엔드를 실행한다.",
    )
    concept_block(
        doc,
        "TypeScript",
        "TypeScript는 JavaScript에 타입 시스템을 더한 언어다. 값의 형태를 미리 표현해 실수를 줄이고 코드 이해를 돕는다.",
        "API 응답 타입이나 화면 props 타입을 명확히 하는 데 유용하다.",
    )
    concept_block(
        doc,
        "컴포넌트(Component)",
        "컴포넌트는 화면을 구성하는 조각이다. 버튼, 카드, 입력창, 화면 전체가 모두 컴포넌트가 될 수 있다. 컴포넌트로 나누면 재사용과 유지보수가 쉬워진다.",
        "translator.tsx 자체도 화면 컴포넌트이고, 그 안의 버튼이나 결과 카드도 더 작은 컴포넌트로 볼 수 있다.",
    )
    concept_block(
        doc,
        "JSX",
        "JSX는 JavaScript 또는 TypeScript 안에서 화면 구조를 태그처럼 작성하는 문법이다. HTML과 비슷해 보이지만 실제로는 JavaScript 표현식이다.",
        "React Native에서는 JSX 안에 View, Text, Pressable 같은 컴포넌트를 배치한다.",
    )
    concept_block(
        doc,
        "State",
        "State는 컴포넌트가 현재 기억하는 값이다. 상태가 바뀌면 React는 화면을 다시 계산한다. 화면의 동적인 변화는 대부분 상태 변화로 표현된다.",
        "번역 입력값, 로딩 여부, 번역 결과, 퀴즈 문제 목록이 모두 state로 관리된다.",
    )
    concept_block(
        doc,
        "리렌더링",
        "리렌더링은 상태나 props가 바뀌었을 때 화면을 다시 계산하는 과정이다. 실제 화면 전체를 무조건 새로 그린다는 뜻은 아니고, React가 변경이 필요한 부분을 다시 반영한다는 의미다.",
        "번역 결과 state가 채워지면 결과 UI가 나타난다.",
    )
    concept_block(
        doc,
        "Props",
        "Props는 부모 컴포넌트가 자식 컴포넌트에게 넘기는 입력값이다. 컴포넌트를 재사용하려면 외부에서 필요한 값을 props로 받는 구조가 중요하다.",
        "공통 버튼 컴포넌트가 있다면 제목, 비활성화 여부, 클릭 함수를 props로 받을 수 있다.",
    )
    concept_block(
        doc,
        "이벤트",
        "이벤트는 사용자의 행동이나 시스템 변화다. 버튼 클릭, 입력값 변경, 화면 진입, 영상 재생 종료 등이 이벤트가 될 수 있다.",
        "번역 버튼 클릭은 API 요청을 시작하는 이벤트다.",
    )
    concept_block(
        doc,
        "핸들러",
        "핸들러는 이벤트가 발생했을 때 실행되는 함수다. 사용자 입력을 상태에 저장하거나, 버튼 클릭 시 API를 호출하는 함수가 핸들러다.",
        "handleTranslate 같은 함수는 text state를 읽어 translateText를 호출하는 역할을 한다.",
    )
    concept_block(
        doc,
        "Hook",
        "Hook은 React 기능을 함수형 컴포넌트 안에서 사용할 수 있게 해 주는 함수다. useState, useEffect, useContext 같은 Hook이 대표적이다.",
        "auth-context를 읽거나 화면 상태를 관리할 때 Hook을 사용한다.",
    )
    concept_block(
        doc,
        "Effect",
        "Effect는 화면 그리기 자체가 아니라 외부 세계와 연결되는 부수 작업을 처리한다. API 호출, 저장소 읽기, 구독 설정, 타이머 등이 여기에 해당한다.",
        "화면 진입 시 퀴즈 세션을 불러오는 작업은 effect로 처리하기 쉽다.",
    )
    concept_block(
        doc,
        "Context",
        "Context는 여러 컴포넌트가 공통으로 필요한 값을 공유하는 방식이다. props를 계속 아래로 넘기지 않아도 공통 상태에 접근할 수 있다.",
        "auth-context.tsx가 로그인 사용자, accessToken, guest 여부를 관리한다.",
    )
    concept_block(
        doc,
        "Router",
        "Router는 어떤 경로에서 어떤 화면을 보여 줄지 결정한다. 사용자가 로그인했는지에 따라 login, signup, home으로 이동시키는 흐름도 라우팅과 관련된다.",
        "Expo Router가 app 폴더 구조를 기반으로 화면 경로를 관리한다.",
    )
    concept_block(
        doc,
        "SecureStore",
        "SecureStore는 모바일 환경에서 민감한 값을 비교적 안전하게 저장하는 Expo 저장소다. 토큰 같은 값을 일반 저장소보다 안전하게 보관할 때 사용한다.",
        "모바일에서는 accessToken과 user 정보를 SecureStore에 저장한다.",
    )
    concept_block(
        doc,
        "localStorage",
        "localStorage는 웹 브라우저에서 데이터를 오래 저장하는 간단한 저장소다. 앱을 새로고침해도 값이 남는다.",
        "웹 실행 시에는 SecureStore 대신 localStorage를 사용한다.",
    )
    concept_block(
        doc,
        "비동기 처리",
        "비동기 처리는 시간이 걸리는 작업을 기다리는 동안 프로그램 전체가 멈추지 않도록 하는 방식이다. 네트워크 요청은 대표적인 비동기 작업이다.",
        "translateText나 fetch 요청은 서버 응답을 기다려야 하므로 async/await 흐름으로 처리한다.",
    )
    concept_block(
        doc,
        "Promise와 async/await",
        "Promise는 나중에 완료될 작업의 결과를 표현하는 객체이고, async/await는 Promise를 더 읽기 쉽게 다루는 문법이다.",
        "API 호출 결과를 기다렸다가 state를 업데이트할 때 async/await를 사용한다.",
    )
    concept_block(
        doc,
        "fetch",
        "fetch는 HTTP 요청을 보내는 JavaScript API다. URL, 메서드, 헤더, 바디를 지정해 서버 API를 호출한다.",
        "frontendcodes/lib/api의 함수들이 fetch를 감싸서 백엔드와 통신한다.",
    )
    concept_block(
        doc,
        "로딩 상태",
        "로딩 상태는 요청이 진행 중임을 사용자에게 알려 주는 상태다. 로딩 표시가 없으면 사용자는 버튼을 눌렀는지, 앱이 멈췄는지 알기 어렵다.",
        "번역 요청 중 isLoading을 true로 두면 버튼 비활성화나 로딩 UI를 보여 줄 수 있다.",
    )
    concept_block(
        doc,
        "에러 상태",
        "에러 상태는 요청 실패나 처리 실패를 사용자에게 알리는 상태다. 네트워크 실패, 서버 오류, 잘못된 입력 등이 원인이 될 수 있다.",
        "API 호출 실패 시 화면에 안내 메시지를 보여 주는 데 필요하다.",
    )
    concept_block(
        doc,
        "빈 상태",
        "빈 상태는 아직 보여 줄 데이터가 없는 상태다. 실패와 다르고 로딩과도 다르다. 잘 설계된 화면은 성공, 로딩, 에러, 빈 상태를 모두 구분한다.",
        "북마크가 하나도 없을 때 별도 안내를 보여 주는 것이 빈 상태 처리다.",
    )
    concept_block(
        doc,
        "플랫폼 차이",
        "같은 React Native 코드라도 웹, iOS, Android에서 동작이 조금씩 다를 수 있다. 저장소, 영상 재생, 권한, URL 처리에서 차이가 생길 수 있다.",
        "웹에서는 localStorage를 쓰고 모바일에서는 SecureStore를 쓰는 것이 대표적인 플랫폼 차이다.",
    )

    add_h1(doc, "6. 백엔드 개념 완전 정리")
    concept_block(
        doc,
        "Java",
        "Java는 백엔드 서버 개발에 많이 쓰이는 정적 타입 언어다. 안정성과 생태계가 강하고, Spring Boot와 함께 많이 사용된다.",
        "backend 폴더의 서버 코드는 Java로 작성되어 있다.",
    )
    concept_block(
        doc,
        "JVM",
        "JVM은 Java 프로그램을 실행하는 가상 머신이다. Java 코드는 컴파일된 뒤 JVM 위에서 실행된다.",
        "백엔드를 실행하려면 Java 17 환경이 필요하다.",
    )
    concept_block(
        doc,
        "Gradle",
        "Gradle은 Java 프로젝트의 빌드 도구다. 의존성 설치, 컴파일, 테스트, 실행을 관리한다.",
        "backend/build.gradle이 백엔드 의존성과 빌드 설정을 담고 있다.",
    )
    concept_block(
        doc,
        "Gradle Wrapper",
        "Gradle Wrapper는 프로젝트가 지정한 Gradle 버전을 자동으로 사용하게 해 주는 실행 파일이다. 개발자마다 전역 Gradle 설치 상태가 달라도 같은 방식으로 실행할 수 있다.",
        "backend 폴더에서 .\\gradlew.bat bootRun으로 실행한다.",
    )
    concept_block(
        doc,
        "Spring Boot",
        "Spring Boot는 Spring 기반 서버를 빠르게 만들 수 있게 해 주는 프레임워크다. 서버 시작, 설정, 의존성 주입, 웹 API, 보안 같은 기능을 제공한다.",
        "SignLanguageApplication이 Spring Boot 애플리케이션의 시작점이다.",
    )
    concept_block(
        doc,
        "애플리케이션 시작점",
        "시작점은 프로그램이 처음 실행되는 위치다. Spring Boot에서는 보통 main 메서드가 있는 클래스가 시작점이다.",
        "SignLanguageApplication이 백엔드 서버를 시작한다.",
    )
    concept_block(
        doc,
        "application.yml",
        "application.yml은 Spring Boot 설정 파일이다. 서버 포트, 외부 API 키, Firebase 설정, JWT 설정 같은 값을 선언한다.",
        "backend/src/main/resources/application.yml에 주요 설정이 있다.",
    )
    concept_block(
        doc,
        "Controller",
        "Controller는 HTTP 요청을 처음 받는 계층이다. 요청 경로와 메서드를 보고 어떤 Java 메서드를 실행할지 연결한다.",
        "TranslateController, QuizController, UserController, AuthController가 대표적이다.",
    )
    concept_block(
        doc,
        "Service",
        "Service는 실제 비즈니스 로직을 처리하는 계층이다. Controller가 요청을 받으면 Service가 계산, 조회, 저장 같은 핵심 일을 한다.",
        "TranslationService, QuizService, UserService가 대표적이다.",
    )
    concept_block(
        doc,
        "DTO",
        "DTO는 요청과 응답 데이터를 담는 객체다. API 경계에서 어떤 데이터가 들어오고 나가는지 명확하게 해 준다.",
        "TranslateRequest, TranslateResponse 같은 객체가 DTO 역할을 한다.",
    )
    concept_block(
        doc,
        "Record",
        "Java record는 데이터를 담는 불변 객체를 간결하게 만들 수 있는 문법이다. DTO를 만들 때 자주 쓰기 좋다.",
        "요청과 응답 구조가 단순할 때 record를 사용하기 좋다.",
    )
    concept_block(
        doc,
        "Bean",
        "Bean은 Spring이 생성하고 관리하는 객체다. 개발자가 직접 new 하지 않아도 Spring 컨테이너가 만들고 필요한 곳에 연결해 준다.",
        "Service, Config, Controller는 대부분 Bean으로 관리된다.",
    )
    concept_block(
        doc,
        "Dependency Injection",
        "Dependency Injection은 필요한 객체를 외부에서 주입받는 방식이다. 객체를 직접 만들지 않고 생성자 등을 통해 받아 쓰면 테스트와 유지보수가 쉬워진다.",
        "TranslationService 생성자에 여러 서비스가 들어가면 Spring이 알아서 연결한다.",
    )
    concept_block(
        doc,
        "Configuration",
        "Configuration은 애플리케이션 전역 설정을 담는 코드다. Firebase 연결, CORS, 보안 정책 같은 기능은 개별 API보다 설정 계층에 가깝다.",
        "FirebaseConfig, SecurityConfig, WebConfig가 여기에 해당한다.",
    )
    concept_block(
        doc,
        "비즈니스 로직",
        "비즈니스 로직은 서비스가 지켜야 하는 실제 규칙이다. 단순 기술 코드가 아니라 앱의 동작 의미를 결정한다.",
        "번역 후보를 사전 hit 수로 고르는 것, 퀴즈 choices가 4개여야 하는 것, 오답 통계를 누적하는 것이 비즈니스 로직이다.",
    )
    concept_block(
        doc,
        "검증(Validation)",
        "검증은 입력값이 올바른지 확인하는 과정이다. 잘못된 요청을 그대로 처리하면 오류나 보안 문제가 생긴다.",
        "빈 questionText, choices 개수 오류, 짧은 JWT_SECRET 같은 값은 검증 대상이다.",
    )
    concept_block(
        doc,
        "예외(Exception)",
        "예외는 프로그램 실행 중 정상 흐름으로 처리하기 어려운 문제가 생겼다는 신호다. 서버는 예외를 적절한 HTTP 오류 응답으로 바꿔 주어야 한다.",
        "Firebase 설정 누락이나 잘못된 로그인 토큰은 예외 상황이 될 수 있다.",
    )
    concept_block(
        doc,
        "로그(Log)",
        "로그는 서버가 실행 중 어떤 일이 일어났는지 기록하는 메시지다. 장애를 분석하거나 요청 흐름을 확인할 때 필요하다.",
        "외부 API 실패, Storage URL 조회 실패 같은 상황은 로그로 남기면 디버깅이 쉬워진다.",
    )
    concept_block(
        doc,
        "무상태(Stateless)",
        "무상태는 서버가 각 요청을 독립적으로 처리하는 성질이다. 요청마다 필요한 인증 정보나 입력이 함께 와야 한다.",
        "JWT를 요청에 담아 보내면 서버가 별도 세션 메모리 없이 사용자를 확인할 수 있다.",
    )

    add_h1(doc, "7. 인증과 보안 개념 완전 정리")
    concept_block(
        doc,
        "인증(Authentication)",
        "인증은 사용자가 누구인지 확인하는 과정이다. 로그인은 대표적인 인증 절차다.",
        "Google 로그인과 idToken 검증이 인증에 해당한다.",
    )
    concept_block(
        doc,
        "인가(Authorization)",
        "인가는 인증된 사용자가 어떤 행동을 할 수 있는지 판단하는 과정이다. 본인 데이터만 수정 가능하게 하는 것이 인가다.",
        "users/{uid} API는 실제 운영에서는 JWT의 사용자와 path의 uid가 같은지 확인해야 한다.",
    )
    concept_block(
        doc,
        "OAuth",
        "OAuth는 외부 서비스가 사용자 인증을 도와주는 표준 흐름이다. 사용자는 Google 계정으로 로그인하고, 우리 앱은 Google이 발급한 토큰을 검증해 신원을 확인한다.",
        "프런트엔드는 Google 로그인으로 idToken을 얻고 백엔드에 보낸다.",
    )
    concept_block(
        doc,
        "idToken",
        "idToken은 사용자가 Google 계정으로 로그인했다는 정보를 담은 토큰이다. 서버는 이 토큰이 진짜인지, 우리 앱을 대상으로 발급됐는지 검증한다.",
        "GoogleAuthService가 idToken 검증을 담당한다.",
    )
    concept_block(
        doc,
        "JWT",
        "JWT는 JSON Web Token의 약자로, 서버가 발급하는 서명된 토큰이다. 사용자가 로그인한 뒤 API 요청에 이 토큰을 보내면 서버가 사용자를 식별할 수 있다.",
        "JwtService가 accessToken을 만든다.",
    )
    concept_block(
        doc,
        "JWT Secret",
        "JWT Secret은 토큰 서명에 쓰는 비밀 키다. 이 값이 노출되면 공격자가 가짜 토큰을 만들 수 있으므로 매우 중요하다.",
        "JWT_SECRET 환경 변수는 32자 이상이어야 한다.",
    )
    concept_block(
        doc,
        "토큰 만료",
        "토큰은 영원히 유효하면 위험하다. 만료 시간을 두면 탈취된 토큰의 피해 범위를 줄일 수 있다.",
        "application.yml에는 JWT 만료 시간이 설정된다.",
    )
    concept_block(
        doc,
        "Bearer Token",
        "Bearer Token은 HTTP Authorization 헤더에 토큰을 넣어 보내는 방식이다. 일반적으로 'Authorization: Bearer 토큰값' 형태를 쓴다.",
        "운영 보강 시 프런트엔드는 accessToken을 Authorization 헤더에 붙이고 백엔드는 이를 검증하는 구조가 필요하다.",
    )
    concept_block(
        doc,
        "서비스 계정",
        "서비스 계정은 사람이 아니라 서버 프로그램이 외부 서비스에 접근할 때 쓰는 계정이다. Firebase Admin이나 Google Cloud Storage 접근에 필요하다.",
        "Firestore와 Storage 연결에는 Firebase 서비스 계정 정보가 필요하다.",
    )
    concept_block(
        doc,
        "최소 권한 원칙",
        "최소 권한 원칙은 필요한 권한만 부여해야 한다는 보안 원칙이다. 너무 넓은 권한을 주면 문제가 생겼을 때 피해가 커진다.",
        "Storage나 Firestore 서비스 계정 권한은 필요한 범위로 제한하는 것이 좋다.",
    )
    concept_block(
        doc,
        "permitAll",
        "permitAll은 Spring Security에서 인증 없이 접근을 허용한다는 뜻이다. 개발 중에는 편하지만 사용자 데이터 API에 넓게 적용하면 위험하다.",
        "현재 여러 API가 permitAll로 열려 있으므로 운영 전 보안 강화가 필요하다.",
    )
    concept_block(
        doc,
        "CSRF",
        "CSRF는 사용자가 의도하지 않은 요청을 인증 상태를 악용해 보내게 만드는 공격이다. 인증 방식과 클라이언트 구조에 따라 대응 방식이 달라진다.",
        "Spring Security 설정에서 CSRF를 어떻게 처리할지 결정해야 한다.",
    )

    add_h1(doc, "8. 데이터베이스와 Firestore 개념 완전 정리")
    concept_block(
        doc,
        "데이터베이스",
        "데이터베이스는 앱이 오래 기억해야 하는 정보를 저장하는 시스템이다. 서버가 재시작되어도 데이터가 남아야 한다면 데이터베이스가 필요하다.",
        "사용자 정보, 퀴즈 문제, 오답 기록, 북마크는 Firestore에 저장된다.",
    )
    concept_block(
        doc,
        "관계형 데이터베이스",
        "관계형 데이터베이스는 데이터를 표, 행, 열로 나누어 저장한다. SQL을 사용하며 엄격한 스키마를 갖는 경우가 많다.",
        "이 프로젝트는 관계형 데이터베이스 대신 Firestore를 사용한다.",
    )
    concept_block(
        doc,
        "문서형 데이터베이스",
        "문서형 데이터베이스는 JSON과 비슷한 문서 단위로 데이터를 저장한다. 유연한 구조가 장점이지만 문서 구조를 잘 관리하지 않으면 혼란이 생긴다.",
        "Firestore가 문서형 데이터베이스다.",
    )
    concept_block(
        doc,
        "Collection",
        "Collection은 Firestore에서 문서들의 묶음이다. 관계형 데이터베이스의 테이블과 비슷하게 생각할 수 있지만 완전히 같지는 않다.",
        "users, quiz_items가 collection이다.",
    )
    concept_block(
        doc,
        "Document",
        "Document는 Firestore에서 하나의 데이터 항목이다. 문서 안에는 여러 field가 들어간다.",
        "users/{uid}는 특정 사용자 한 명의 문서다.",
    )
    concept_block(
        doc,
        "Field",
        "Field는 문서 안의 개별 값이다. 문자열, 숫자, boolean, 배열, map, timestamp 같은 값을 가질 수 있다.",
        "name, email, dailySolvedGoal, isRegistered가 field다.",
    )
    concept_block(
        doc,
        "Subcollection",
        "Subcollection은 문서 아래에 다시 붙는 collection이다. 사용자별 하위 데이터를 묶을 때 편리하다.",
        "users/{uid}/translator_bookmarks가 subcollection이다.",
    )
    concept_block(
        doc,
        "Map 필드",
        "Map은 key-value 형태로 여러 값을 한 필드 안에 담는 구조다. 날짜별 통계나 문제별 카운트를 저장할 때 유용하다.",
        "dailySolvedCounts와 incorrectQuestionCounts가 map 형태로 쓰인다.",
    )
    concept_block(
        doc,
        "Timestamp",
        "Timestamp는 특정 시각을 저장하는 타입이다. 생성일, 수정일, 저장일, 오답 시각 같은 값을 기록할 때 쓴다.",
        "savedAt, wrongAt 같은 필드가 시간 기록 역할을 한다.",
    )
    concept_block(
        doc,
        "Query",
        "Query는 데이터베이스에서 조건에 맞는 데이터를 찾는 요청이다. 전체를 다 가져오지 않고 필요한 문서만 가져오는 것이 중요하다.",
        "퀴즈 세션을 만들 때 active 문제를 조회하는 과정이 query에 해당한다.",
    )
    concept_block(
        doc,
        "Index",
        "Index는 데이터베이스가 조건 검색을 빠르게 하도록 도와주는 구조다. Firestore에서는 특정 복합 조건에 index가 필요할 수 있다.",
        "category와 active 조건을 조합해 많이 조회한다면 index가 필요할 수 있다.",
    )
    concept_block(
        doc,
        "Transaction",
        "Transaction은 여러 읽기와 쓰기를 하나의 안전한 작업으로 묶는 방식이다. 동시에 여러 사용자가 같은 데이터를 수정할 때 일관성을 지키는 데 도움이 된다.",
        "퀴즈 통계처럼 attempt_count, correct_count, wrong_count를 안정적으로 갱신할 때 고려할 수 있다.",
    )
    concept_block(
        doc,
        "스키마 문서화",
        "Firestore는 유연하지만, 유연하다는 이유로 구조를 문서화하지 않으면 나중에 어떤 필드가 필수인지 알기 어려워진다.",
        "users와 quiz_items의 필드 구조는 별도 문서로 고정해 두는 것이 좋다.",
    )

    add_h1(doc, "9. Firebase Storage와 영상 파일 개념")
    concept_block(
        doc,
        "파일 저장소",
        "파일 저장소는 동영상, 이미지, 음성처럼 큰 파일을 저장하는 공간이다. 데이터베이스는 메타데이터를 저장하고, 파일 저장소는 실제 파일을 저장하는 식으로 역할을 나누는 것이 일반적이다.",
        "수어 영상 파일은 Firebase Storage에 저장된다.",
    )
    concept_block(
        doc,
        "Bucket",
        "Bucket은 Storage의 가장 큰 저장 공간 단위다. 하나의 프로젝트 안에서도 여러 bucket을 가질 수 있다.",
        "FIREBASE_STORAGE_BUCKET 환경 변수가 어떤 bucket을 사용할지 알려 준다.",
    )
    concept_block(
        doc,
        "Object",
        "Object는 Storage 안의 실제 파일 하나다. 영상 파일 하나가 object라고 생각하면 된다.",
        "각 수어 영상 파일이 Storage object다.",
    )
    concept_block(
        doc,
        "Prefix",
        "Prefix는 파일 경로 앞부분이다. 폴더처럼 보이지만 Storage에서는 객체 이름의 앞부분으로 이해하는 편이 정확하다.",
        "StorageVideoCache는 configured prefix, 기본 prefix, alternate prefix를 시도한다.",
    )
    concept_block(
        doc,
        "파일명 규칙",
        "파일명 규칙은 단어와 영상 파일을 연결하는 기준이다. 파일명이 불규칙하면 코드가 영상을 찾기 어렵다.",
        "StorageVideoCache는 파일명에서 확장자를 제거한 값을 단어 key로 사용할 수 있다.",
    )
    concept_block(
        doc,
        "다운로드 URL",
        "다운로드 URL은 앱이 파일을 실제로 가져오거나 재생할 수 있는 주소다. 영상 재생을 위해서는 단어를 URL로 바꾸는 과정이 필요하다.",
        "번역 응답의 clips 안에는 영상 URL이 포함된다.",
    )
    concept_block(
        doc,
        "Signed URL",
        "Signed URL은 일정 시간 동안만 유효한 서명된 접근 주소다. 파일을 완전히 공개하지 않고 필요한 동안만 접근을 허용할 수 있다.",
        "다운로드 토큰이 없을 때 signed URL 생성 권한이 필요할 수 있다.",
    )
    concept_block(
        doc,
        "캐시",
        "캐시는 반복 조회를 줄이기 위해 이전 결과를 저장해 두는 구조다. 같은 단어의 영상을 매번 Storage에서 찾으면 느리고 비용이 들 수 있다.",
        "StorageVideoCache가 찾은 URL과 실패한 단어를 기억한다.",
        "캐시는 원본 파일이 바뀌었을 때 갱신 정책을 고민해야 한다.",
    )

    add_h1(doc, "10. 번역과 자연어 처리 개념 완전 정리")
    concept_block(
        doc,
        "자연어",
        "자연어는 사람이 일상적으로 사용하는 언어다. 한국어, 영어처럼 사람이 쓰는 언어는 규칙이 복잡하고 예외가 많다.",
        "사용자가 번역기에 입력하는 한국어 문장이 자연어다.",
    )
    concept_block(
        doc,
        "자연어 처리",
        "자연어 처리는 컴퓨터가 사람의 언어를 분석하고 변환하도록 만드는 기술 분야다. 문장 분리, 형태소 분석, 정규화, 번역, 요약 등이 포함된다.",
        "TranslationService는 간단한 자연어 처리 파이프라인을 구성한다.",
    )
    concept_block(
        doc,
        "토큰",
        "토큰은 문장을 처리하기 위해 잘라 낸 단위다. 단어와 비슷하지만 항상 같은 것은 아니다. 어떤 기준으로 자르느냐에 따라 번역 결과가 달라진다.",
        "normalizedTokens는 최종적으로 선택된 수어 재생용 토큰 목록이다.",
    )
    concept_block(
        doc,
        "토큰화",
        "토큰화는 문장을 토큰으로 나누는 과정이다. 공백만으로 나눌 수도 있지만 한국어는 조사와 어미가 붙기 때문에 더 복잡한 처리가 필요하다.",
        "TextNormalizer와 SignSentenceSimplifier가 토큰화 전후의 정리를 돕는다.",
    )
    concept_block(
        doc,
        "형태소",
        "형태소는 의미를 가진 가장 작은 언어 단위다. 한국어에서는 조사, 어미, 어근 등이 형태소 분석 대상이 될 수 있다.",
        "ETRI 형태소 분석은 문장을 더 세밀하게 나누는 후보를 제공한다.",
    )
    concept_block(
        doc,
        "정규화",
        "정규화는 다양한 입력 표현을 표준 형태로 정리하는 과정이다. 같은 뜻의 여러 표현을 사전과 맞기 쉬운 형태로 바꾼다.",
        "OpenAiMorphologyNormalizerService는 정규화 후보를 만드는 보조 역할을 한다.",
    )
    concept_block(
        doc,
        "기본형",
        "기본형은 활용된 단어를 사전에서 찾기 쉬운 형태로 되돌린 것이다. 예를 들어 '가요'를 '가다'에 가깝게 보는 식이다.",
        "수어 사전이 기본형 중심이라면 기본형 복원이 중요하다.",
    )
    concept_block(
        doc,
        "조사",
        "조사는 한국어에서 단어 뒤에 붙어 문법 관계를 나타낸다. 은, 는, 이, 가, 을, 를 같은 요소가 있다.",
        "수어 영상과 직접 대응되지 않는 조사는 제거하거나 약화하는 후보가 된다.",
    )
    concept_block(
        doc,
        "시제",
        "시제는 과거, 현재, 미래 같은 시간 정보를 나타낸다. 문장 번역에서는 단어 자체와 별도로 metadata로 남길 수 있다.",
        "SignSentenceSimplifier는 과거/미래 같은 정보를 metadata에 남길 수 있다.",
    )
    concept_block(
        doc,
        "부정",
        "부정은 아니다, 안, 못처럼 어떤 행동이나 상태를 부정하는 표현이다. 단순 단어 매칭만 하면 부정 의미가 사라질 수 있어 따로 처리해야 한다.",
        "부정 여부는 번역 metadata나 토큰 흐름에 영향을 줄 수 있다.",
    )
    concept_block(
        doc,
        "의문",
        "의문은 질문 문장인지 여부다. 질문 표지나 어미가 있으면 일반 평서문과 다르게 표시하거나 순서를 바꿔야 할 수 있다.",
        "질문 여부는 metadata에 남길 수 있다.",
    )
    concept_block(
        doc,
        "문장 단순화",
        "문장 단순화는 한국어 원문을 수어 영상 재생에 더 적합한 구조로 정리하는 과정이다. 불필요한 문법 요소를 줄이고 핵심 의미 단위를 남긴다.",
        "SignSentenceSimplifier가 문장 단순화의 핵심이다.",
    )
    concept_block(
        doc,
        "수어식 순서",
        "수어식 순서는 한국어 문장 순서를 그대로 따르기보다 수어 표현에 더 어울리는 순서로 재배열하는 생각이다.",
        "시간, 장소, 주체, 행동 같은 순서로 재구성하는 규칙이 들어갈 수 있다.",
    )
    concept_block(
        doc,
        "후보 스트림",
        "후보 스트림은 하나의 입력 문장에서 가능한 토큰 목록 후보를 뜻한다. 여러 분석 방법이 서로 다른 후보를 만들 수 있다.",
        "rule, etri, openai 후보 스트림을 비교한다.",
    )
    concept_block(
        doc,
        "사전 매칭",
        "사전 매칭은 토큰이 sign_dictionary.json 안의 단어와 맞는지 확인하는 과정이다. 최종 영상 재생 가능성과 직접 연결된다.",
        "DictionaryLoader가 사전 데이터를 읽고 매칭 기준을 제공한다.",
    )
    concept_block(
        doc,
        "사전 hit 수",
        "사전 hit 수는 후보 토큰 중 사전에 존재하는 단어가 얼마나 많은지 나타낸다. 이 값이 높으면 실제 영상으로 연결될 가능성이 높다.",
        "TranslationService는 사전 hit 수를 기준으로 후보를 선택한다.",
    )
    concept_block(
        doc,
        "unknown",
        "unknown은 사전에서 찾지 못한 단어다. 언어 처리나 사전 데이터 보강이 필요한 신호다.",
        "TranslateResponse.unknown에 남는다.",
    )
    concept_block(
        doc,
        "noVideoWords",
        "noVideoWords는 단어는 알지만 연결할 영상이 없는 경우다. 이 경우는 사전보다 영상 자산 보강이 필요하다.",
        "TranslateResponse.noVideoWords에 남는다.",
    )
    concept_block(
        doc,
        "Fallback",
        "Fallback은 주 방식이 실패했을 때 사용하는 예비 방식이다. 외부 API가 실패해도 규칙 기반 후보로 계속 진행하면 서비스가 멈추지 않는다.",
        "ETRI나 OpenAI가 실패해도 rule 기반 처리가 남는다.",
    )
    concept_block(
        doc,
        "Metadata",
        "Metadata는 본 데이터에 대한 추가 설명 정보다. 번역 결과에서는 시제, 부정, 의문, 사용된 후보 출처 같은 정보를 metadata로 담을 수 있다.",
        "TranslateResponse.metadata는 프런트엔드가 더 자세한 설명 UI를 만들 때 유용하다.",
    )
    concept_block(
        doc,
        "외부 API",
        "외부 API는 우리 서버 밖의 서비스가 제공하는 기능이다. 강력하지만 네트워크 실패, 키 누락, 요금, 속도 제한 같은 변수가 있다.",
        "ETRI, 우리말샘, 한국어기초사전, OpenAI가 외부 API 후보에 해당한다.",
    )

    add_h1(doc, "11. 프로젝트 흐름으로 모든 개념 연결하기")
    add_h2(doc, "11.1 번역 요청")
    add_flow(
        doc,
        "번역 요청 전체 연결",
        [
            "사용자 입력",
            "  ↓ 프런트엔드 state",
            "번역 버튼 이벤트",
            "  ↓ 핸들러 실행",
            "fetch POST /translate",
            "  ↓ HTTP JSON 요청",
            "TranslateController",
            "  ↓",
            "TranslationService",
            "  ├─ 정규화",
            "  ├─ 토큰화",
            "  ├─ 후보 스트림 생성",
            "  ├─ 사전 hit 비교",
            "  ├─ unknown 보정",
            "  └─ Storage URL 조회",
            "  ↓ JSON 응답",
            "프런트엔드 result state",
            "  ↓ 리렌더링",
            "영상 재생 UI 표시",
        ],
    )
    add_h2(doc, "11.2 로그인 요청")
    add_flow(
        doc,
        "로그인 요청 전체 연결",
        [
            "Google 로그인 버튼",
            "  ↓ 이벤트와 핸들러",
            "Google idToken 획득",
            "  ↓ POST /api/auth/login",
            "AuthController",
            "  ↓",
            "GoogleAuthService 검증",
            "  ↓",
            "UserService 사용자 조회",
            "  ├─ 기존 사용자: JwtService로 JWT 발급",
            "  └─ 신규 사용자: needsSignup 반환",
            "  ↓",
            "프런트엔드 auth context 저장",
            "  ↓",
            "SecureStore 또는 localStorage 저장",
        ],
    )
    add_h2(doc, "11.3 퀴즈 요청")
    add_flow(
        doc,
        "퀴즈 요청 전체 연결",
        [
            "quiz 화면 진입",
            "  ↓ effect 실행",
            "GET /api/quiz/session",
            "  ↓",
            "QuizController",
            "  ↓",
            "QuizService",
            "  ├─ Firestore quiz_items 조회",
            "  ├─ isActive, choices 검증",
            "  ├─ 영상 URL 보정",
            "  └─ 세션 응답 생성",
            "  ↓",
            "프런트엔드 questions state",
            "  ↓",
            "사용자 선택지 클릭",
            "  ↓ POST /api/quiz/answer",
            "백엔드 정답 판정과 통계 갱신",
        ],
    )

    add_h1(doc, "12. 개념과 실제 코드 위치 연결표")
    add_table(
        doc,
        ["개념", "코드 위치", "무엇을 보면 되는가"],
        [
            ("라우팅", "frontendcodes/app/_layout.tsx", "앱 진입과 화면 이동 구조"),
            ("인증 상태", "frontendcodes/context/auth-context.tsx", "user, accessToken, guest 저장과 복원"),
            ("API 호출", "frontendcodes/lib/api", "fetch 래퍼와 백엔드 URL 처리"),
            ("번역 화면", "frontendcodes/app/translator.tsx", "입력, 로딩, 결과, 영상 재생"),
            ("퀴즈 화면", "frontendcodes/app/quiz.tsx", "세션 호출, 답 제출, 결과 표시"),
            ("백엔드 설정", "backend/src/main/resources/application.yml", "환경 변수와 외부 서비스 설정"),
            ("보안 설정", "backend/src/main/java/.../config/SecurityConfig.java", "허용 경로와 보안 정책"),
            ("번역 API", "TranslateController, TranslationService", "문장 처리와 응답 생성"),
            ("문장 단순화", "SignSentenceSimplifier", "조사 제거, 역할 분류, 순서 재배열"),
            ("사전 로딩", "DictionaryLoader", "sign_dictionary.json 읽기"),
            ("영상 URL", "StorageVideoCache", "Firebase Storage 조회와 캐시"),
            ("퀴즈 로직", "QuizController, QuizService", "문제 조회와 정답 판정"),
            ("사용자 데이터", "UserController, UserService", "학습량, 오답, 북마크 저장"),
            ("로그인", "AuthController, GoogleAuthService, JwtService", "Google 검증과 JWT 발급"),
        ],
        [1.35, 2.6, 2.55],
    )

    add_h1(doc, "13. 디버깅을 위한 개념 지도")
    add_table(
        doc,
        ["증상", "관련 개념", "먼저 볼 곳"],
        [
            ("프런트가 서버를 못 찾음", "URL, localhost, port, CORS", "base-url.ts, WebConfig, 서버 실행 상태"),
            ("로그인이 실패함", "OAuth, idToken, JWT, 환경 변수", "GoogleAuthService, AuthController, .env"),
            ("번역 결과가 비어 있음", "토큰화, 정규화, 사전 매칭", "TranslationService, DictionaryLoader"),
            ("영상이 안 나옴", "Storage, URL, signed URL, prefix", "StorageVideoCache, Firebase Storage"),
            ("퀴즈가 안 나옴", "Firestore query, validation", "QuizService, quiz_items 데이터"),
            ("오답 기록이 안 쌓임", "Firestore update, map field, user id", "UserService.tryQuestion"),
            ("화면이 안 바뀜", "state, render, async", "해당 화면 state 업데이트 코드"),
            ("앱 재실행 후 로그인이 풀림", "local storage, SecureStore", "auth-context.tsx"),
        ],
        [1.7, 2.3, 2.5],
    )
    add_callout(
        doc,
        "디버깅의 기본 순서",
        "먼저 문제가 프런트엔드인지 백엔드인지 나눈다. 그 다음 네트워크 요청이 실제로 갔는지 확인하고, 서버 응답이 정상인지 본다. 응답이 정상인데 화면이 틀리면 state와 렌더링을 보고, 응답 자체가 틀리면 Controller와 Service를 따라간다.",
        FILL_WARN,
    )

    add_h1(doc, "14. 꼭 기억할 핵심 개념 요약")
    numbered(
        doc,
        [
            "프런트엔드는 화면과 상태를 담당한다.",
            "백엔드는 규칙, 검증, 저장, 외부 연동을 담당한다.",
            "API는 프런트엔드와 백엔드의 계약이다.",
            "HTTP 요청은 메서드, URL, 헤더, 바디로 구성된다.",
            "JSON은 API 데이터 교환에 쓰이는 대표 형식이다.",
            "React에서는 state가 바뀌면 화면이 다시 계산된다.",
            "Spring Boot에서는 Controller가 입구이고 Service가 실제 로직을 처리한다.",
            "JWT는 서버가 발급한 로그인 상태 표현 토큰이다.",
            "Firestore는 문서형 데이터베이스이고 Storage는 파일 저장소다.",
            "번역 파이프라인은 입력 정리, 토큰화, 후보 비교, 사전 매칭, 영상 연결로 이어진다.",
            "unknown은 사전 문제이고 noVideoWords는 영상 자산 문제다.",
            "운영 환경에서는 permitAll 범위를 줄이고 서버에서 권한을 강제해야 한다.",
        ],
    )

    add_h1(doc, "15. 마지막 정리")
    add_para(
        doc,
        "개발 개념은 서로 따로 떨어져 있지 않다. 사용자가 버튼을 누르는 순간 이벤트, state, API, HTTP, Controller, Service, 데이터베이스, Storage, 응답, 리렌더링이 한 번에 이어진다. 그래서 하나씩 외우는 것보다 실제 요청 흐름 안에서 연결해서 이해하는 편이 오래 간다."
    )
    add_para(
        doc,
        "이 문서는 수어지교 프로젝트를 읽기 위한 개념 지도다. 코드에서 모르는 단어가 나오면 해당 장으로 돌아와 개념을 확인하고, 다시 코드 위치 연결표를 통해 실제 파일로 돌아가면 된다."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
