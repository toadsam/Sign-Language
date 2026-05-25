from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "sign-language-tech-guide-ko-detailed.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "64748B"
FILL_BLUE = "E8EEF5"
FILL_GRAY = "F4F6F9"
FILL_WARN = "FFF7E6"
FILL_OK = "ECFDF5"
BORDER = "CBD5E1"
FONT = "Malgun Gothic"
MONO = "Consolas"


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)


def set_table_indent(table, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("수어지교 기술 문서")
    set_run_font(run, size=9, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.1)
    run = p.add_run(title)
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=10)
    run = p.add_run(subtitle)
    set_run_font(run, size=11, color=MUTED)


def add_h1(doc, text):
    return doc.add_heading(text, level=1)


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def add_h3(doc, text):
    return doc.add_heading(text, level=3)


def add_para(doc, text="", bold_lead=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_lead:
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=INK)
        if text:
            run = p.add_run(text)
            set_run_font(run, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4)
    run = p.add_run(text)
    set_run_font(run, color=INK)
    return p


def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(tabs)
    p_pr.append(ind)
    for node in (start, num_fmt, lvl_text, lvl_jc, p_pr):
        lvl.append(node)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_list(doc, items):
    num_id = new_numbering_id(doc)
    for text in items:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p_pr.append(num_pr)
        run = p.add_run(text)
        set_run_font(run, color=INK)


def add_code(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=8, line=1.1)
    for line in text.strip("\n").split("\n"):
        run = p.add_run(line)
        set_run_font(run, MONO, size=9.2, color="0F172A")
        run.add_break()
    return p


def add_callout(doc, title, body, fill=FILL_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_indent(table, 120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_indent(table, 120)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, FILL_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        r = p.add_run(h)
        set_run_font(r, size=9.3, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.18)
            r = p.add_run(value)
            set_run_font(r, size=9.2, color=INK)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            set_cell_borders(cells[i])
    doc.add_paragraph()
    return table


def add_flow(doc, title, lines):
    add_h3(doc, title)
    add_code(doc, "\n".join(lines))


def build_doc():
    doc = Document()
    configure_document(doc)

    add_title(
        doc,
        "수어지교 기술 문서",
        "기초 개념부터 프런트엔드와 백엔드의 전체 흐름까지 초보자 기준으로 다시 풀어 쓴 상세 가이드",
    )
    add_callout(
        doc,
        "문서의 기준",
        "이 문서는 현재 저장소의 실제 코드 구조를 기준으로 작성했다. frontendcodes의 Expo/React Native 앱, backend의 Spring Boot 서버, Firebase Firestore와 Storage, Google 로그인, JWT, 번역 파이프라인, 퀴즈와 사용자 기록 흐름을 설명한다.",
        FILL_OK,
    )
    add_para(
        doc,
        "이번 문서는 먼저 아주 기초적인 개념을 설명하고, 그 다음 백엔드와 프런트엔드를 완전히 나누어 설명한 뒤, 마지막에 두 영역이 실제 요청 하나 안에서 어떻게 연결되는지 다시 합쳐서 설명한다."
    )

    add_h1(doc, "1. 프로젝트를 한 문장으로 이해하기")
    add_para(
        doc,
        "수어지교는 수어 영상을 활용해 단어를 학습하고, 한국어 문장을 수어 영상 재생 순서로 변환해 보여 주며, 퀴즈와 오답노트로 반복 학습까지 지원하는 앱이다."
    )
    add_para(
        doc,
        "기술적으로는 하나의 앱처럼 보이지만 내부에는 서로 다른 역할의 시스템이 함께 있다. 화면을 보여 주는 프런트엔드, 요청을 처리하는 백엔드, 데이터를 저장하는 Firestore, 영상을 저장하는 Storage, 로그인 확인을 도와주는 Google, 번역 후보를 보완하는 외부 API가 함께 움직인다."
    )
    for text in [
        "프런트엔드는 사용자가 직접 보는 화면과 상호작용을 담당한다.",
        "백엔드는 요청을 받아 번역, 채점, 저장, 로그인 검증을 수행한다.",
        "Firestore는 사용자와 퀴즈 같은 구조화된 데이터를 저장한다.",
        "Firebase Storage는 실제 수어 영상 파일을 저장한다.",
        "외부 API는 형태소 후보나 단어 보정을 도와 번역 정확도를 보완한다.",
    ]:
        add_bullet(doc, text)

    add_h1(doc, "2. 초보자를 위한 아주 기초적인 개념")
    add_table(
        doc,
        ["개념", "쉽게 말하면"],
        [
            ("앱", "휴대폰이나 웹에서 사용하는 프로그램이다."),
            ("프런트엔드", "사용자가 보는 앞부분이다. 화면과 버튼과 영상 재생이 여기에 속한다."),
            ("백엔드", "요청을 처리하는 뒷부분이다. 계산과 저장과 검증을 맡는다."),
            ("서버", "백엔드가 실제로 실행되는 대상이다."),
            ("API", "프런트엔드와 백엔드가 대화하는 약속이다."),
            ("HTTP 요청", "프런트엔드가 백엔드에게 일을 부탁하는 메시지다."),
            ("JSON", "데이터를 주고받는 글자 형식이다."),
            ("상태", "프로그램이 현재 기억하고 있는 값이다."),
            ("토큰", "문장을 처리하기 위해 잘라 낸 단어 단위다."),
            ("정규화", "여러 표현을 사전과 맞기 쉬운 형태로 정리하는 일이다."),
            ("인증", "사용자가 누구인지 확인하는 과정이다."),
            ("데이터베이스", "앱을 껐다 켜도 남아 있는 데이터를 저장하는 공간이다."),
        ],
        [1.45, 5.05],
    )
    add_para(
        doc,
        "처음 보는 사람은 프런트엔드를 접수 창구, 백엔드를 실제 작업 사무실, Firestore를 문서 보관실, Storage를 영상 창고라고 생각하면 이해가 쉽다."
    )
    add_flow(
        doc,
        "사용자 기준 기본 흐름",
        [
            "사용자",
            "  ↓ 입력 또는 버튼 클릭",
            "프런트엔드 화면",
            "  ↓ API 요청",
            "백엔드 서버",
            "  ↓ 계산 / 저장 / 조회",
            "Firestore / Storage / 외부 API",
            "  ↓ 응답",
            "프런트엔드 상태 변경",
            "  ↓ 화면 다시 그리기",
        ],
    )

    add_h1(doc, "3. 저장소 구조를 먼저 머리에 넣기")
    add_table(
        doc,
        ["영역", "경로", "역할"],
        [
            ("프런트엔드 화면", "frontendcodes/app", "login, signup, home, translator, quiz, mypage 등 실제 화면 파일이 있다."),
            ("프런트엔드 공통 로직", "frontendcodes/lib, frontendcodes/context", "API 호출, 인증 상태, 공통 유틸을 담당한다."),
            ("백엔드 컨트롤러", "backend/src/main/java/.../api, auth, quiz, user", "HTTP 요청을 처음 받는 입구다."),
            ("백엔드 서비스", "backend/src/main/java/.../service", "실제 번역, 퀴즈, 사용자 처리 로직이 있다."),
            ("백엔드 설정", "backend/src/main/resources/application.yml, config", "포트, 보안, Firebase, 외부 API 설정이 있다."),
            ("수어 사전", "backend/src/main/resources/sign_dictionary.json", "단어와 영상 파일명을 연결하는 로컬 사전이다."),
            ("영속 데이터", "Firestore", "users, quiz_items, 오답노트, 북마크 등 문서 데이터를 저장한다."),
            ("파일 저장소", "Firebase Storage", "수어 영상 파일을 저장한다."),
        ],
        [1.35, 2.35, 2.85],
    )

    add_h1(doc, "4. 실행 준비와 환경 변수")
    add_para(
        doc,
        "이 프로젝트는 프런트엔드와 백엔드를 각각 따로 실행해야 한다. 화면이 한 프로그램처럼 보여도 실제로는 두 프로그램이 동시에 돌아가며 서로 통신한다."
    )
    add_h2(doc, "4.1 백엔드 실행")
    add_code(
        doc,
        r"""
cd backend
.\gradlew.bat bootRun
""",
    )
    add_h2(doc, "4.2 프런트엔드 실행")
    add_code(
        doc,
        r"""
cd frontendcodes
npm install
npx expo start
""",
    )
    add_h2(doc, "4.3 주요 환경 변수")
    add_table(
        doc,
        ["변수", "필수 여부", "설명"],
        [
            ("JWT_SECRET", "필수", "백엔드가 JWT를 서명하는 비밀 키다."),
            ("GOOGLE_CLIENT_ID", "로그인 사용 시 필요", "Google idToken이 우리 앱용인지 검증할 때 사용한다."),
            ("FIREBASE_SERVICE_ACCOUNT_PATH", "Firestore 사용 시 필요", "Firestore 접근을 위한 인증 정보다."),
            ("FIREBASE_STORAGE_BUCKET", "영상 사용 시 필요", "수어 영상이 들어 있는 버킷 이름이다."),
            ("FIREBASE_STORAGE_SERVICE_ACCOUNT_PATH", "영상 사용 시 필요", "Storage 접근을 위한 인증 정보다."),
            ("OPENAI_API_KEY", "선택", "OpenAI 후보 정규화를 쓸 때 필요하다."),
            ("ETRI_ACCESS_KEY", "선택", "ETRI 형태소 후보를 쓸 때 필요하다."),
            ("URIMALSAM_API_KEY / KRDICT_API_KEY", "선택", "단어 보정 후보를 더 얻을 때 사용한다."),
            ("EXPO_PUBLIC_API_BASE_URL", "프런트엔드 설정", "프런트엔드가 연결할 백엔드 주소다."),
        ],
        [2.0, 1.2, 3.3],
    )
    add_callout(
        doc,
        "왜 환경 변수가 많은가",
        "비밀 키나 외부 서비스 연결 정보는 코드에 직접 넣지 않고 실행 환경에서 받아야 안전하다. 또한 배포 환경마다 값을 바꾸기 쉽다.",
        FILL_WARN,
    )

    add_h1(doc, "5. 백엔드 설명")
    add_callout(
        doc,
        "백엔드를 한 줄로 요약하면",
        "프런트엔드가 보낸 요청을 받아 실제 일을 처리하는 작업장이다. 번역, 퀴즈, 로그인, 사용자 기록 저장이 핵심 역할이다.",
    )
    add_h2(doc, "5.1 백엔드가 맡는 책임")
    for text in [
        "문장을 수어 재생에 맞는 토큰 목록으로 바꾼다.",
        "토큰별 수어 영상 URL을 찾는다.",
        "퀴즈 문제를 꺼내고 정답을 판정한다.",
        "사용자 학습 기록과 오답노트와 북마크를 저장한다.",
        "Google 로그인 결과를 검증하고 JWT를 발급한다.",
        "Firebase와 외부 API를 연결한다.",
    ]:
        add_bullet(doc, text)

    add_h2(doc, "5.2 Spring Boot를 처음 보는 사람을 위한 개념")
    add_table(
        doc,
        ["용어", "쉽게 이해하기"],
        [
            ("Controller", "요청을 처음 받는 접수 창구다."),
            ("Service", "실제 작업을 수행하는 계층이다."),
            ("Config", "보안, CORS, Firebase 같은 운영 규칙을 넣는 곳이다."),
            ("Bean", "스프링이 자동으로 만들고 관리하는 객체다."),
            ("Dependency Injection", "필요한 객체를 직접 만들지 않고 스프링이 넣어 주는 방식이다."),
            ("DTO", "요청과 응답 데이터를 정리해서 담는 객체다."),
        ],
        [1.55, 4.95],
    )
    add_para(
        doc,
        "이 프로젝트에서 TranslateController는 접수 창구이고, TranslationService는 번역을 실제로 수행하는 작업자다. QuizController와 QuizService, AuthController와 관련 서비스도 같은 구조로 이해하면 된다."
    )

    add_h2(doc, "5.3 백엔드 시작 시 내부에서 일어나는 일")
    add_numbered_list(
        doc,
        [
            "Spring Boot가 애플리케이션 컨테이너를 만든다.",
            "application.yml과 환경 변수를 읽어 설정 값을 준비한다.",
            "FirebaseConfig가 Firestore와 Storage 접근 객체를 만든다.",
            "DictionaryLoader가 sign_dictionary.json을 메모리에 올린다.",
            "StorageVideoCache가 영상 조회용 캐시 구조를 준비한다.",
            "SecurityConfig가 허용 경로와 보안 규칙을 설정한다.",
            "WebConfig가 CORS와 정적 리소스 관련 설정을 잡는다.",
        ],
    )

    add_h2(doc, "5.4 번역 기능 백엔드 상세 흐름")
    add_para(
        doc,
        "번역 기능은 가장 복합적인 영역이다. 단순히 문장을 나누는 수준이 아니라, 문장을 정리하고, 후보를 만들고, 사전과 비교하고, 영상을 연결하는 여러 단계가 이어진다."
    )
    add_flow(
        doc,
        "번역 서비스 내부 단계",
        [
            "TranslateController.translate()",
            "  ↓",
            "TranslationService.translate()",
            "  ├─ TextNormalizer",
            "  ├─ SignSentenceSimplifier",
            "  ├─ ExternalLexiconApiClient",
            "  ├─ OpenAiMorphologyNormalizerService",
            "  ├─ chooseTokenStream",
            "  ├─ UnknownTokenResolverService",
            "  ├─ DictionaryLoader",
            "  └─ StorageVideoCache",
            "  ↓",
            "TranslateResponse 반환",
        ],
    )
    add_table(
        doc,
        ["단계", "무슨 일을 하나", "왜 필요한가"],
        [
            ("입력 정리", "문장을 다루기 쉬운 형태로 청소한다.", "사용자 입력은 공백과 기호가 일정하지 않기 때문이다."),
            ("문장 단순화", "조사와 문장 역할을 정리해 수어식 순서에 가깝게 만든다.", "한국어 원문 그대로는 사전 매칭이 잘 안 되기 때문이다."),
            ("후보 생성", "규칙 기반, ETRI, OpenAI 결과를 각각 만든다.", "한 방식만 쓰면 놓치는 표현이 생길 수 있기 때문이다."),
            ("후보 선택", "사전과 가장 잘 맞는 토큰열을 고른다.", "최종 목표가 실제 영상 재생 가능한 결과를 만드는 것이기 때문이다."),
            ("미확인 단어 보정", "사전에 없는 단어를 다시 변형하거나 외부 후보를 참고한다.", "처음 실패한 단어를 최대한 살리기 위해서다."),
            ("영상 연결", "선택된 토큰마다 영상 URL을 찾는다.", "토큰만 있고 영상이 없으면 사용자 경험이 완성되지 않는다."),
        ],
        [1.05, 2.35, 3.1],
    )
    add_h3(doc, "TextNormalizer")
    add_para(
        doc,
        "이 단계는 입력 청소 단계다. 사용자가 입력한 문자열을 뒤 단계가 안정적으로 다룰 수 있도록 공백과 기호, 표현 흔들림을 정리한다."
    )
    add_h3(doc, "SignSentenceSimplifier")
    add_para(
        doc,
        "이 단계는 규칙 기반 핵심이다. 시간, 장소, 주어, 목적어, 서술어, 질문 여부를 구분하고 수어 영상 재생에 더 맞는 순서로 재배열한다."
    )
    for text in [
        "조사 제거 또는 약화",
        "시간, 장소, 행동 역할 분류",
        "질문, 부정, 시제 메타데이터 생성",
        "수어식 순서 재배열",
    ]:
        add_bullet(doc, text)
    add_h3(doc, "후보 토큰 스트림 선택")
    add_para(
        doc,
        "TranslationService는 rule, etri, openai 후보를 비교한다. 이 프로젝트의 기준은 언어학적으로 가장 근사한 분석보다, 지금 가진 사전과 영상 자산에 가장 잘 연결되는 분석이다."
    )
    add_table(
        doc,
        ["후보", "출처", "의미"],
        [
            ("rule", "SignSentenceSimplifier", "항상 만들 수 있는 기본 후보"),
            ("etri", "ETRI 형태소 분석", "외부 분석이 성공했을 때 얻는 후보"),
            ("openai", "OpenAI 정규화", "OpenAI 호출이 성공했을 때 얻는 후보"),
        ],
        [1.0, 2.0, 3.5],
    )
    add_para(
        doc,
        "사전 hit 수가 같을 때는 openai, etri, rule 순으로 우선한다. 그러나 외부 API가 실패해도 rule 후보가 남아 있기 때문에 번역 기능 전체는 계속 동작한다."
    )
    add_h3(doc, "UnknownTokenResolverService")
    add_para(
        doc,
        "최종 후보를 골라도 사전에 없는 단어가 남을 수 있다. 이 서비스는 다시 한 번 단어를 보정해 보며, 그래도 못 찾으면 unknown에 남긴다. 사전은 있지만 영상이 없으면 noVideoWords에 남긴다."
    )
    add_h3(doc, "StorageVideoCache")
    add_para(
        doc,
        "이 서비스는 단어를 실제 영상으로 연결하는 마지막 단계다. 여러 prefix와 파일명 후보를 시도하고, 이미 찾은 결과는 캐시에 보관해 반복 조회를 줄인다."
    )
    for text in [
        "configured prefix와 기본 prefix를 차례로 시도한다.",
        "파일명에서 확장자를 제거한 값을 단어 key와 비교한다.",
        "가능하면 direct URL을 만들고 필요하면 signed URL도 시도한다.",
        "실패 결과도 캐시에 남겨 불필요한 반복 조회를 줄인다.",
    ]:
        add_bullet(doc, text)

    add_h2(doc, "5.5 퀴즈 기능 백엔드 상세 흐름")
    add_para(
        doc,
        "퀴즈는 Firestore의 quiz_items 데이터를 사용해 세션을 만들고, 사용자의 답을 서버가 판정하며, 통계와 사용자 기록을 갱신하는 구조다."
    )
    add_flow(
        doc,
        "퀴즈 백엔드 흐름",
        [
            "QuizController",
            "  ├─ GET /api/quiz/session",
            "  ├─ GET /api/quiz/session/wrong",
            "  └─ POST /api/quiz/answer",
            "  ↓",
            "QuizService",
            "  ├─ Firestore에서 문제 조회",
            "  ├─ 활성 문제 필터링",
            "  ├─ SessionQuestion 변환",
            "  ├─ 영상 URL 보정",
            "  ├─ 정답 판정",
            "  └─ 통계 갱신",
        ],
    )
    for text in [
        "isActive=true인 문제만 일반 세션에 포함된다.",
        "choices는 4개가 아니면 제외된다.",
        "videoUrl이 없으면 correctChoiceText 기준으로 다시 영상을 찾는다.",
        "오답 기반 세션은 incorrectQuestionCounts와 연결된다.",
    ]:
        add_bullet(doc, text)

    add_h2(doc, "5.6 로그인과 회원가입 백엔드 상세 흐름")
    add_para(
        doc,
        "사용자는 Google로 로그인하지만, 백엔드는 그 결과를 바탕으로 우리 서비스용 JWT를 발급한다. 즉 외부 로그인과 내부 세션 상태 관리를 분리한 구조다."
    )
    add_flow(
        doc,
        "인증 백엔드 흐름",
        [
            "AuthController.login()",
            "  ↓",
            "GoogleAuthService.verifyIdToken()",
            "  ↓",
            "사용자 조회",
            "  ├─ 있으면 JWT 발급",
            "  └─ 없으면 needsSignup 반환",
            "  ↓",
            "signup / signup complete로 사용자 문서 완성",
        ],
    )

    add_h2(doc, "5.7 사용자 데이터와 Firestore 구조")
    add_table(
        doc,
        ["경로", "주요 필드", "용도"],
        [
            ("users/{uid}", "name, email, profileImage, dailySolvedGoal, isRegistered", "사용자 기본 정보와 설정"),
            ("users/{uid}", "incorrectQuestionCounts", "문제별 오답 횟수 맵"),
            ("users/{uid}", "dailySolvedCounts", "날짜별 풀이 수 맵"),
            ("users/{uid}/translator_bookmarks/{id}", "questionText, word, videoUrl, savedAt", "번역 북마크"),
            ("users/{uid}/wrong_note_saved/{quizId}", "questionText, word, videoUrl, isHidden", "저장한 오답노트"),
            ("quiz_items/{quizId}", "questionText, choices, correctChoiceId, videoUrl, isActive", "퀴즈 문제 원본"),
            ("quiz_items/{quizId}", "attempt_count, correct_count, wrong_count, difficulty_level", "문제별 통계"),
        ],
        [1.85, 2.25, 2.4],
    )
    add_callout(
        doc,
        "날짜 처리",
        "사용자 일일 학습량은 Asia/Seoul 기준 날짜 키로 누적된다. 즉 통계 기준 시간은 한국 시간으로 고정되어 있다.",
    )

    add_h2(doc, "5.8 백엔드 API 요약")
    add_table(
        doc,
        ["메서드", "경로", "담당", "설명"],
        [
            ("GET", "/api/health", "HealthController", "서버 상태 확인"),
            ("POST", "/translate", "TranslateController", "문장을 번역 토큰과 영상 정보로 변환"),
            ("POST", "/api/auth/login", "AuthController", "기존 사용자 로그인"),
            ("POST", "/api/auth/signup", "AuthController", "신규 사용자 기본 생성"),
            ("POST", "/api/auth/signup/complete", "AuthController", "회원가입 완료"),
            ("GET", "/api/quiz/session", "QuizController", "일반 퀴즈 세션"),
            ("GET", "/api/quiz/session/wrong", "QuizController", "오답 기반 세션"),
            ("POST", "/api/quiz/answer", "QuizController", "정답 판정"),
            ("GET", "/api/users/{uid}", "UserController", "사용자 조회"),
            ("PATCH", "/api/users/{uid}/tryQuestion", "UserController", "풀이 기록 반영"),
            ("POST", "/api/users/{uid}/translator-bookmarks", "UserController", "번역 북마크 저장"),
            ("GET", "/api/users/{uid}/top-wrong-words", "UserController", "많이 틀린 단어 조회"),
        ],
        [0.8, 2.15, 1.25, 2.3],
    )

    add_h1(doc, "6. 프런트엔드 설명")
    add_callout(
        doc,
        "프런트엔드를 한 줄로 요약하면",
        "사용자가 앱을 실제로 경험하는 부분이며, 입력을 받고 상태를 바꾸고 API를 호출하고 결과를 화면으로 다시 만드는 계층이다.",
    )
    add_h2(doc, "6.1 프런트엔드가 맡는 책임")
    for text in [
        "로그인, 회원가입, 홈, 번역기, 퀴즈, 마이페이지 화면을 보여 준다.",
        "사용자 입력과 현재 화면 상태를 메모리에 유지한다.",
        "백엔드 API를 호출하고 로딩, 실패, 성공 상태를 처리한다.",
        "응답 데이터를 영상 재생과 카드 UI에 맞게 다시 조립한다.",
        "로그인 토큰과 사용자 정보를 로컬 저장소에 보관한다.",
    ]:
        add_bullet(doc, text)

    add_h2(doc, "6.2 React Native와 Expo를 처음 보는 사람을 위한 개념")
    add_table(
        doc,
        ["개념", "쉽게 이해하기"],
        [
            ("Component", "화면을 이루는 조각이다."),
            ("State", "현재 화면이 기억하는 값이다."),
            ("Effect", "화면이 열리거나 값이 바뀔 때 실행되는 작업이다."),
            ("Router", "어떤 경로에서 어떤 화면을 보여 줄지 정하는 장치다."),
            ("Context", "여러 화면이 공유해야 하는 상태를 담는 곳이다."),
            ("Hook", "React 기능을 재사용 가능한 함수 형태로 묶은 것이다."),
            ("Props", "부모가 자식 컴포넌트에 넘겨 주는 값이다."),
        ],
        [1.4, 5.1],
    )
    add_para(
        doc,
        "프런트엔드는 상태가 바뀌면 화면을 다시 그리는 프로그램이라고 생각하면 된다. 예를 들어 번역 버튼을 누르면 isLoading이 true가 되고, 응답이 오면 result가 채워지고, 화면은 그 값에 맞춰 다시 바뀐다."
    )

    add_h2(doc, "6.3 프런트엔드 폴더 읽는 순서")
    add_table(
        doc,
        ["파일", "왜 먼저 보나", "설명"],
        [
            ("app/_layout.tsx", "라우팅의 시작점", "앱의 전체 라우트 구조를 이해할 수 있다."),
            ("context/auth-context.tsx", "인증의 중심", "로그인 상태와 토큰 저장 방식이 여기에 있다."),
            ("lib/api/*", "화면과 서버 연결", "프런트가 어떤 API를 부르는지 알 수 있다."),
            ("app/login.tsx, signup.tsx", "인증 흐름 이해", "로그인과 가입 시작점을 볼 수 있다."),
            ("app/translator.tsx", "핵심 기능 이해", "번역 화면의 전체 흐름이 있다."),
            ("app/quiz.tsx", "학습 흐름 이해", "퀴즈 세션과 채점 반영 흐름이 있다."),
        ],
        [1.55, 1.35, 3.6],
    )

    add_h2(doc, "6.4 프런트엔드 시작 과정")
    add_numbered_list(
        doc,
        [
            "expo-router/entry가 앱을 시작한다.",
            "app/_layout.tsx가 전체 라우트 구조를 설정한다.",
            "AuthProvider가 저장된 인증 상태를 복원한다.",
            "AuthGate가 로그인 여부를 보고 적절한 화면으로 보낸다.",
            "각 화면은 필요할 때 lib/api 함수를 호출한다.",
            "응답이 오면 state를 바꾸고 화면이 다시 렌더링된다.",
        ],
    )

    add_h2(doc, "6.5 인증 상태 관리")
    add_para(
        doc,
        "auth-context는 프런트엔드의 공통 기억 장치다. accessToken, user, guest 여부를 저장하고 복원한다. 모바일에서는 SecureStore, 웹에서는 localStorage를 사용한다."
    )
    for text in [
        "앱 시작 시 저장된 값을 읽어 현재 세션을 복원한다.",
        "로그인 성공 시 토큰과 user를 저장한다.",
        "게스트 모드도 별도 상태로 관리한다.",
        "로그아웃 시 저장 값과 메모리 상태를 함께 비운다.",
    ]:
        add_bullet(doc, text)

    add_h2(doc, "6.6 lib/api 계층의 역할")
    add_para(
        doc,
        "화면마다 fetch를 직접 작성하면 중복이 많아지고 수정 범위가 커진다. 그래서 auth.ts, translate.ts, quiz.ts, users.ts 같은 파일로 API 호출을 따로 분리했다. 화면은 무슨 요청을 할지만 알고, 네트워크 세부 구현은 API 계층이 맡는다."
    )

    add_h2(doc, "6.7 번역기 화면 프런트엔드 상세 흐름")
    add_flow(
        doc,
        "translator.tsx 흐름",
        [
            "문장 입력",
            "  ↓ text state 변경",
            "번역 버튼 클릭",
            "  ↓ isLoading = true",
            "translateText(text) 호출",
            "  ↓ 응답 수신",
            "result, playbackItems, unknown, noVideoWords 갱신",
            "  ↓ 화면 재렌더링",
            "영상 목록과 결과 상세 표시",
        ],
    )
    add_para(
        doc,
        "이 화면은 입력창 하나만 있는 단순 화면이 아니다. 로딩 상태, 에러 상태, 결과 상태, 영상 재생 상태, 북마크 저장, 웹과 네이티브 환경별 재생 차이까지 함께 관리해야 한다."
    )
    for text in [
        "입력 텍스트는 state로 관리된다.",
        "번역 중에는 버튼 비활성화나 로딩 표시가 필요하다.",
        "응답의 clips와 items를 재생 UI에 맞게 변환해야 한다.",
        "unknown과 noVideoWords를 따로 보여 주면 실패 원인을 이해하기 쉽다.",
        "로그인 사용자라면 결과를 북마크로 저장할 수 있다.",
    ]:
        add_bullet(doc, text)

    add_h2(doc, "6.8 퀴즈 화면 프런트엔드 상세 흐름")
    add_flow(
        doc,
        "quiz.tsx 흐름",
        [
            "화면 진입",
            "  ↓ 세션 API 호출",
            "questions state 채움",
            "  ↓ 현재 문제 표시",
            "선택지 클릭",
            "  ↓ answer API 호출",
            "정답 여부 상태 갱신",
            "  ↓ 필요 시 오답 저장 및 통계 반영",
            "다음 문제 이동",
        ],
    )
    add_para(
        doc,
        "퀴즈 화면은 문제 배열과 현재 인덱스를 관리하는 구조라고 생각하면 쉽다. 사용자가 선택지를 누르면 프런트엔드가 스스로 채점하지 않고 서버 응답을 받아 화면을 바꾼다."
    )

    add_h2(doc, "6.9 다른 화면들의 역할")
    add_table(
        doc,
        ["화면 파일", "사용자가 보는 기능", "주요 API"],
        [
            ("login.tsx", "Google 로그인과 게스트 시작", "/api/auth/login, /api/auth/signup"),
            ("signup.tsx", "추가 정보 입력과 가입 완료", "/api/auth/signup/complete"),
            ("home.tsx", "학습 진입과 요약 정보", "/api/users/{uid}"),
            ("translator.tsx", "문장 번역과 영상 재생", "/translate"),
            ("quiz.tsx", "문제 풀이와 채점 반영", "/api/quiz/session, /api/quiz/answer"),
            ("bookmark.tsx", "번역 북마크 조회와 삭제", "/api/users/{uid}/translator-bookmarks"),
            ("wrongnote.tsx", "오답노트 조회와 관리", "/api/users/{uid}/wrong-note-saved"),
            ("mypage.tsx", "프로필과 통계", "/api/users/{uid}, /api/users/{uid}/top-wrong-words"),
        ],
        [1.6, 2.45, 2.45],
    )

    add_h1(doc, "7. 프런트엔드와 백엔드가 실제로 연결되는 방식")
    add_para(
        doc,
        "두 영역을 따로 본 뒤에는 다시 합쳐 봐야 한다. 사용자가 버튼 하나를 눌렀을 때 프런트엔드와 백엔드가 어떻게 이어지는지 이해해야 전체 흐름이 완성된다."
    )
    add_h2(doc, "7.1 번역 요청 한 번의 전체 흐름")
    add_numbered_list(
        doc,
        [
            "사용자가 translator 화면에 문장을 입력한다.",
            "프런트엔드 state가 입력값을 기억한다.",
            "번역 버튼을 누르면 translateText가 POST /translate를 호출한다.",
            "TranslateController가 요청을 받고 TranslationService로 넘긴다.",
            "백엔드가 문장을 정리하고 토큰 후보를 만들고 영상을 찾는다.",
            "TranslateResponse JSON이 프런트엔드로 돌아온다.",
            "프런트엔드는 결과 state를 갱신한다.",
            "React가 상태 변화를 감지해 화면을 다시 그린다.",
            "사용자는 영상 목록, 미확인 단어, 상세 결과를 본다.",
        ],
    )
    add_code(
        doc,
        """POST /translate
{
  "text": "오늘 학교에 가요"
}

응답의 핵심 개념
{
  "input": "...",
  "simplifiedSentence": "...",
  "normalizedTokens": ["오늘", "학교", "가다"],
  "clips": [...],
  "items": [...],
  "unknown": [],
  "noVideoWords": []
}""",
    )
    add_h2(doc, "7.2 로그인 요청 한 번의 전체 흐름")
    add_numbered_list(
        doc,
        [
            "사용자가 Google 로그인 버튼을 누른다.",
            "프런트엔드가 Google에서 idToken을 받는다.",
            "login.tsx가 이 값을 /api/auth/login으로 보낸다.",
            "백엔드가 Google 토큰을 검증하고 사용자 존재 여부를 확인한다.",
            "기존 사용자면 JWT를 발급하고, 신규 사용자면 needsSignup을 반환한다.",
            "프런트엔드는 home 또는 signup으로 이동한다.",
            "성공 시 토큰과 사용자 정보를 로컬 저장소에 기록한다.",
        ],
    )
    add_h2(doc, "7.3 퀴즈 요청 한 번의 전체 흐름")
    add_numbered_list(
        doc,
        [
            "사용자가 quiz 화면에 들어간다.",
            "프런트엔드가 세션 API를 호출한다.",
            "백엔드가 Firestore에서 문제를 가져와 세션용 JSON을 만든다.",
            "프런트엔드는 문제 배열 state를 채운다.",
            "사용자가 답을 누르면 answer API를 호출한다.",
            "백엔드가 정답을 판정하고 통계를 갱신한다.",
            "프런트엔드는 결과를 보여 주고 다음 문제로 이동한다.",
        ],
    )

    add_h1(doc, "8. 기능을 추가하거나 수정하는 실제 순서")
    add_h2(doc, "8.1 새 수어 단어와 영상을 추가하는 방법")
    add_numbered_list(
        doc,
        [
            "영상 파일 이름 규칙을 먼저 확인한다.",
            "Firebase Storage의 맞는 prefix 아래에 영상 파일을 올린다.",
            "sign_dictionary.json에 id, word, file 항목을 추가한다.",
            "서버를 다시 시작해 새 사전을 로드한다.",
            "번역기로 해당 단어를 넣어 /translate 응답과 영상 재생을 확인한다.",
            "실패하면 prefix, 파일명, 버킷, 권한 순서로 점검한다.",
        ],
    )
    add_h2(doc, "8.2 새 퀴즈 문제를 추가하는 방법")
    add_numbered_list(
        doc,
        [
            "Firestore quiz_items에 새 문서를 만든다.",
            "questionText를 넣고 choices를 정확히 4개 배열로 맞춘다.",
            "correctChoiceId를 정확히 지정한다.",
            "videoUrl 또는 영상 연결 가능한 correctChoiceText를 준비한다.",
            "isActive=true로 설정해 일반 세션에 포함되게 한다.",
            "필요하면 category와 difficulty_level을 함께 저장한다.",
        ],
    )
    add_h2(doc, "8.3 새 API를 추가하는 방법")
    add_numbered_list(
        doc,
        [
            "요청 JSON과 응답 JSON을 먼저 설계한다.",
            "Controller에서 받을 DTO와 응답 DTO를 정한다.",
            "실제 비즈니스 규칙은 Service에 넣는다.",
            "필요하면 Firestore, Storage, 외부 API 연동을 추가한다.",
            "프런트엔드 lib/api에 호출 함수를 만든다.",
            "화면에서는 loading, error, empty 상태까지 함께 처리한다.",
        ],
    )

    add_h1(doc, "9. 디버깅 체크리스트")
    add_h2(doc, "9.1 백엔드가 안 켜질 때")
    for text in [
        "JWT_SECRET 길이를 확인한다.",
        "GOOGLE_CLIENT_ID 누락 여부를 확인한다.",
        "Firebase 서비스 계정 경로나 JSON이 유효한지 본다.",
        "FIREBASE_STORAGE_BUCKET이 비어 있지 않은지 확인한다.",
        "backend 폴더에서 gradlew wrapper를 통해 실행했는지 확인한다.",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "9.2 프런트엔드가 백엔드를 못 찾을 때")
    for text in [
        "EXPO_PUBLIC_API_BASE_URL이 실제 서버 주소와 맞는지 확인한다.",
        "모바일 기기 테스트에서는 localhost 대신 PC의 LAN IP가 필요할 수 있다.",
        "백엔드 CORS 설정이 현재 프런트 origin을 허용하는지 확인한다.",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "9.3 번역 결과에 영상이 없을 때")
    for text in [
        "unknown에 있으면 사전에 없는 단어다.",
        "noVideoWords에 있으면 사전은 있지만 영상이 없는 경우다.",
        "Storage prefix와 파일명이 단어와 맞는지 본다.",
        "다운로드 토큰 또는 signed URL 권한 문제인지 점검한다.",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "9.4 퀴즈가 안 나올 때")
    for text in [
        "isActive=true인지 확인한다.",
        "choices가 4개인지 확인한다.",
        "questionText와 videoUrl이 비어 있지 않은지 확인한다.",
        "category 필터 조건이 실제 데이터와 맞는지 본다.",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "9.5 로그인 문제가 있을 때")
    for text in [
        "프런트엔드와 백엔드의 Google 클라이언트 ID가 같은 앱을 가리키는지 확인한다.",
        "redirect URI가 실행 환경과 맞는지 확인한다.",
        "needsSignup 응답은 신규 사용자 흐름일 수 있으므로 오류와 구분한다.",
    ]:
        add_bullet(doc, text)

    add_h1(doc, "10. 보안과 운영 관점")
    add_table(
        doc,
        ["위험", "왜 문제인가", "개선 방향"],
        [
            ("사용자 API 공개 범위", "uid만 알면 접근 범위가 넓어질 수 있다.", "JWT subject와 path uid를 대조한다."),
            ("JWT 발급 중심 구조", "발급은 하지만 검증 범위가 약할 수 있다.", "검증 필터를 명시적으로 넣는다."),
            ("환경 변수 누락", "배포 시 서버 시작 실패나 Storage 연결 실패가 난다.", "필수 변수 검증을 추가한다."),
            ("외부 API 실패 불투명", "정확도 저하 원인을 파악하기 어렵다.", "metadata에 사용된 분석 경로를 남긴다."),
            ("문자 인코딩 문제", "깨진 문자열이 사용자 경험과 데이터 품질을 해친다.", "UTF-8 정리와 검사 절차를 둔다."),
        ],
        [1.4, 2.45, 2.65],
    )

    add_h1(doc, "11. 테스트 전략")
    add_h2(doc, "11.1 수동 테스트")
    add_numbered_list(
        doc,
        [
            "/api/health가 200인지 확인한다.",
            "게스트 진입이 정상인지 확인한다.",
            "신규 Google 로그인 시 signup 흐름이 이어지는지 본다.",
            "기존 사용자는 바로 home으로 이동하는지 본다.",
            "짧은 단어, 짧은 문장, 알 수 없는 단어 번역을 각각 확인한다.",
            "번역 결과의 clips, unknown, noVideoWords 표시를 확인한다.",
            "퀴즈 세션 생성과 정답 판정이 정상인지 확인한다.",
            "오답노트 저장과 조회가 되는지 확인한다.",
            "북마크 저장, 검색, 삭제가 되는지 확인한다.",
            "일일 학습량과 7일 통계가 갱신되는지 확인한다.",
        ],
    )
    add_h2(doc, "11.2 자동 테스트로 보강하면 좋은 것")
    for text in [
        "TextNormalizer 단위 테스트",
        "SignSentenceSimplifier 단위 테스트",
        "TranslationService 후보 선택 테스트",
        "QuizService 세션 생성과 정답 판정 테스트",
        "UserService 통계 업데이트 테스트",
        "프런트엔드 API URL 보정 테스트",
    ]:
        add_bullet(doc, text)

    add_h1(doc, "12. 처음 유지보수할 때 읽는 순서")
    add_numbered_list(
        doc,
        [
            "application.yml로 전체 설정과 의존성을 본다.",
            "frontendcodes/app/_layout.tsx와 auth-context.tsx를 본다.",
            "frontendcodes/lib/api를 보고 프런트가 호출하는 API를 파악한다.",
            "TranslateController와 TranslationService로 번역 핵심 흐름을 본다.",
            "QuizController와 QuizService로 학습 흐름을 본다.",
            "UserController와 UserService로 사용자 데이터 흐름을 본다.",
            "AuthController와 JwtService로 인증 흐름을 본다.",
            "마지막으로 각 화면 파일을 열어 UI 상태와 API 연결을 본다.",
        ],
    )

    add_h1(doc, "13. 자주 헷갈리는 질문")
    add_table(
        doc,
        ["질문", "답"],
        [
            ("왜 clips와 items가 둘 다 필요한가?", "clips는 영상 중심 재생 목록이고, items는 영상이 없는 단어까지 포함한 전체 순서를 보존하기 위해서다."),
            ("왜 OpenAI만 쓰지 않나?", "키가 없거나 실패해도 서비스가 계속 돌아야 하고, 최종 기준은 현재 사전과 영상 자산이기 때문이다."),
            ("왜 퀴즈 채점을 서버가 하나?", "정답과 통계의 권위를 서버가 가져야 조작을 막기 쉽기 때문이다."),
            ("왜 SecureStore와 localStorage가 둘 다 나오나?", "모바일과 웹 저장 방식이 다르기 때문이다."),
            ("왜 사전 hit 수가 중요한가?", "최종 목표가 영상과 연결되는 번역이기 때문에 실제 사전에 잘 맞는 후보가 유리하다."),
        ],
        [2.2, 4.3],
    )

    add_h1(doc, "14. 개선 로드맵")
    add_h2(doc, "14.1 우선순위 높음")
    for text in [
        "JWT 검증 필터 추가",
        "문자 인코딩 정리",
        "TranslationService 테스트 확장",
        "Firestore 스키마 문서화",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "14.2 우선순위 중간")
    for text in [
        "번역 응답 metadata 확장",
        "관리자용 데이터 등록 도구 추가",
        "Storage cache 갱신 정책 추가",
        "오답 기반 추천 학습 강화",
    ]:
        add_bullet(doc, text)

    add_h1(doc, "15. 용어 사전")
    add_table(
        doc,
        ["용어", "설명"],
        [
            ("Controller", "HTTP 요청을 받아 적절한 서비스로 넘기는 클래스"),
            ("Service", "실제 비즈니스 로직이 들어 있는 클래스"),
            ("DTO", "요청과 응답 데이터를 담는 객체"),
            ("Bean", "Spring이 생성하고 관리하는 객체"),
            ("Dependency Injection", "필요한 객체를 외부에서 주입하는 방식"),
            ("CORS", "다른 주소의 프런트와 백엔드 통신을 허용하는 브라우저 보안 설정"),
            ("JWT", "로그인 상태를 표현하는 서명된 토큰"),
            ("Morpheme", "의미를 가진 가장 작은 말 단위"),
            ("Fallback", "주 방식이 실패했을 때 대신 쓰는 예비 방식"),
            ("Cache", "반복 조회를 줄이기 위해 결과를 잠시 저장하는 구조"),
        ],
        [1.55, 4.95],
    )

    add_h1(doc, "16. 마지막 요약")
    add_para(
        doc,
        "이 프로젝트의 핵심은 프런트엔드가 입력을 받고, 백엔드가 번역과 채점과 저장을 처리하고, Firestore와 Storage가 그 결과를 유지하는 연결 구조다."
    )
    add_para(
        doc,
        "처음 보는 사람은 화면부터 뜯어보기보다, 어떤 화면이 어떤 API를 부르고 그 API가 어떤 서비스와 저장소를 거치는지 따라가는 편이 훨씬 빠르다. 이번 문서는 그 흐름을 백엔드와 프런트엔드로 분리해 이해한 뒤 다시 하나로 합치는 데 초점을 맞췄다."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
