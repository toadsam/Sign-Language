from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "sign-language-real-code-deep-dive-ko.docx"

FONT = "Malgun Gothic"
CODE_FONT = "Consolas"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x55, 0x65, 0x76)
CODE_BG = "F4F6F9"
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "AAB7C4"


def set_run_font(run, font=FONT, size=None, color=None, bold=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style_font(style, font=FONT, size=None, color=None, bold=None):
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def add_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=TABLE_BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    for kind, text in (("begin", None), ("instrText", "PAGE"), ("end", None)):
        if kind == "instrText":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = text
        else:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        run._r.append(node)


def configure_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=10.5, color=INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    set_style_font(styles["Title"], size=23, color=RGBColor(0x0B, 0x25, 0x45), bold=True)
    styles["Title"].paragraph_format.space_after = Pt(8)

    set_style_font(styles["Subtitle"], size=10.5, color=MUTED)
    styles["Subtitle"].paragraph_format.space_after = Pt(16)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, size=10.5, color=INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "CodeBlock" not in [s.name for s in styles]:
        code_style = styles.add_style("CodeBlock", 1)
    else:
        code_style = styles["CodeBlock"]
    set_style_font(code_style, font=CODE_FONT, size=8.3, color=RGBColor(0x1F, 0x29, 0x37))
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = "Sign-Language 실제 코드 상세 해설"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        set_run_font(run, size=8.5, color=MUTED)

    return doc


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = widths or [6.5 / len(headers)] * len(headers)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=8.8)
    doc.add_paragraph()
    return table


def read_lines(relative_path):
    path = ROOT / relative_path
    return path.read_text(encoding="utf-8", errors="replace").splitlines()


def find_line(lines, pattern, start=0):
    for idx in range(start, len(lines)):
        if pattern in lines[idx]:
            return idx
    raise ValueError(f"pattern not found: {pattern}")


def extract_range(relative_path, start_pattern, end_pattern=None, max_lines=80, occurrence=1):
    lines = read_lines(relative_path)
    start = 0
    for _ in range(occurrence):
        start = find_line(lines, start_pattern, start)
        found = start
        start += 1
    start = found
    if end_pattern:
        end = find_line(lines, end_pattern, start + 1) + 1
    else:
        end = min(len(lines), start + max_lines)
    return start + 1, "\n".join(lines[start:end])


def extract_block(relative_path, pattern, occurrence=1, max_lines=180):
    lines = read_lines(relative_path)
    search = 0
    for _ in range(occurrence):
        idx = find_line(lines, pattern, search)
        search = idx + 1
    start = idx
    brace = 0
    seen = False
    end = min(len(lines), start + max_lines)
    for j in range(start, min(len(lines), start + max_lines)):
        brace += lines[j].count("{")
        brace -= lines[j].count("}")
        if "{" in lines[j]:
            seen = True
        if seen and brace == 0 and j > start:
            end = j + 1
            break
    return start + 1, "\n".join(lines[start:end])


def add_code(doc, relative_path, code, start_line):
    label = doc.add_paragraph()
    run = label.add_run(f"파일: {relative_path} / 시작 줄: {start_line}")
    set_run_font(run, size=8.8, color=MUTED, bold=True)
    label.paragraph_format.space_after = Pt(2)

    numbered = []
    for offset, line in enumerate(code.splitlines()):
        numbered.append(f"{start_line + offset:>4} | {line}")
    p = doc.add_paragraph(style="CodeBlock")
    add_shading(p, CODE_BG)
    run = p.add_run("\n".join(numbered))
    set_run_font(run, font=CODE_FONT, size=8.0, color=RGBColor(0x1F, 0x29, 0x37))


def add_code_block_from_range(doc, relative_path, start_pattern, title, explanation, end_pattern=None, max_lines=80, occurrence=1):
    doc.add_paragraph(title, style="Heading 3")
    start_line, code = extract_range(relative_path, start_pattern, end_pattern, max_lines, occurrence)
    add_code(doc, relative_path, code, start_line)
    add_para(doc, "코드 해설", style="Heading 3")
    add_bullets(doc, explanation)


def add_code_block_from_method(doc, relative_path, pattern, title, explanation, occurrence=1, max_lines=180):
    doc.add_paragraph(title, style="Heading 3")
    start_line, code = extract_block(relative_path, pattern, occurrence, max_lines)
    add_code(doc, relative_path, code, start_line)
    add_para(doc, "코드 해설", style="Heading 3")
    add_bullets(doc, explanation)


def add_concept_box(doc, title, lines):
    p = doc.add_paragraph()
    add_shading(p, "EEF6FF")
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        add_shading(p, "EEF6FF")
        run = p.add_run(line)
        set_run_font(run, size=9.8)


def build():
    doc = configure_doc()

    doc.add_paragraph("Sign-Language 실제 코드 기반 초상세 해설서", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    run = subtitle.add_run(
        "코드를 직접 보여주고, 그 코드가 어떤 역할을 하는지, 왜 필요한지, 다른 파일과 어떻게 연결되는지 강의 노트처럼 풀어쓴 문서"
    )
    set_run_font(run)

    add_para(
        doc,
        "이 문서는 단순한 프로젝트 소개 문서가 아니다. 실제 파일 안의 코드를 가져와서 코드 블록으로 보여준 뒤, 그 아래에서 실행 시점, 줄별 의미, 연결되는 기능, 사용된 개념을 설명한다. 그래서 코드를 아직 완전히 이해하지 못하는 사람도 프론트엔드 버튼 클릭이 백엔드 API와 Firestore, Firebase Storage, 영상 재생까지 어떻게 이어지는지 따라갈 수 있도록 구성했다.",
    )
    add_concept_box(
        doc,
        "읽는 방법",
        [
            "먼저 코드 블록의 파일 경로와 시작 줄을 본다. 이 정보는 실제 프로젝트에서 해당 코드를 찾을 때 사용한다.",
            "코드 블록 아래의 설명은 '무엇을 한다'에서 끝나지 않고 '왜 이렇게 해야 하는지'와 '다른 파일과 어떻게 이어지는지'까지 설명한다.",
            "Java 백엔드 코드는 Controller -> Service -> Firestore/Storage/외부 API 순서로 읽으면 이해하기 쉽다.",
            "React Native 프론트엔드 코드는 화면 컴포넌트 -> 상태 useState -> 부수 효과 useEffect -> API 클라이언트 순서로 읽으면 이해하기 쉽다.",
        ],
    )

    doc.add_paragraph("1. 전체 구조를 코드 관점에서 먼저 잡기", style="Heading 1")
    add_table(
        doc,
        ["영역", "주요 파일", "코드가 맡는 일"],
        [
            ["앱 시작/라우팅", "frontendcodes/app/_layout.tsx", "로그인 여부를 보고 /login 또는 실제 화면으로 이동시킨다."],
            ["인증 상태", "frontendcodes/context/auth-context.tsx", "토큰과 사용자 정보를 저장하고 모든 화면에서 공유한다."],
            ["번역 화면", "frontendcodes/app/translator.tsx", "문장 입력, /translate 호출, 여러 수어 영상 순차 재생을 담당한다."],
            ["번역 API", "backend/.../api/TranslateController.java", "프론트 요청을 받아 TranslationService로 넘긴다."],
            ["번역 로직", "backend/.../service/TranslationService.java", "정규화, 사전 매칭, 영상 URL 조회를 모두 연결한다."],
            ["퀴즈 화면", "frontendcodes/app/quiz.tsx", "문제 영상 재생, 보기 선택, 답 제출, 오답 저장을 담당한다."],
            ["퀴즈 API", "backend/.../quiz/QuizService.java", "Firestore 문제 조회, 정답 판정, 문제 통계 갱신을 담당한다."],
            ["사용자 데이터", "backend/.../user/UserService.java", "학습 기록, 일일 목표, 오답노트, 북마크, 프로필을 저장한다."],
        ],
        widths=[1.2, 2.2, 3.1],
    )

    add_para(
        doc,
        "전체 코드를 이해할 때 가장 중요한 관점은 '화면에서 일어난 이벤트가 API 요청이 되고, API 요청은 서비스 로직으로 들어가며, 서비스 로직은 데이터베이스나 Storage와 연결된다'는 것이다. 이 프로젝트는 프론트엔드와 백엔드가 분리되어 있기 때문에 한 파일만 봐서는 전체 기능이 보이지 않는다. 예를 들어 번역 버튼 하나를 눌렀을 때도 translator.tsx, translate.ts, TranslateController.java, TranslationService.java, DictionaryLoader.java, StorageVideoCache.java가 연속적으로 사용된다.",
    )

    doc.add_paragraph("2. 백엔드 서버는 어떻게 시작되는가", style="Heading 1")
    add_code_block_from_range(
        doc,
        "backend/src/main/java/com/wow/signlanguage/SignLanguageApplication.java",
        "public class SignLanguageApplication",
        "Spring Boot 시작 코드",
        [
            "`@SpringBootApplication`은 이 클래스가 Spring Boot 앱의 시작점이라는 뜻이다. 이 어노테이션 하나 안에는 컴포넌트 스캔, 자동 설정, 설정 클래스 등록 기능이 묶여 있다.",
            "`main` 메서드는 Java 프로그램이 처음 실행되는 입구이다. 백엔드 서버를 실행하면 가장 먼저 이 메서드가 호출된다.",
            "`SpringApplication.run(...)`은 내장 Tomcat 서버를 띄우고, `@RestController`, `@Service`, `@Component`, `@Configuration`이 붙은 클래스들을 찾아 객체로 등록한다.",
            "이 코드가 실행된 뒤에야 `TranslateController`, `QuizController`, `UserController` 같은 API 클래스가 요청을 받을 수 있다.",
            "즉 이 파일은 기능 로직을 직접 담지는 않지만, 모든 백엔드 코드가 작동할 수 있게 시동을 거는 엔진 스위치 역할을 한다.",
        ],
        max_lines=20,
    )

    doc.add_paragraph("3. 프론트엔드 API 주소는 어떻게 결정되는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/base-url.ts",
        "export function getBaseUrl()",
        "getBaseUrl: 프론트가 백엔드를 찾는 방식",
        [
            "`DEFAULT_BASE_URL`은 개발 중 기본 백엔드 주소이다. 보통 Spring Boot 서버가 8080 포트에서 실행되므로 기본값이 `http://localhost:8080`이다.",
            "`process.env.EXPO_PUBLIC_API_BASE_URL`은 배포 환경이나 다른 PC에서 백엔드 주소를 바꿀 수 있게 해 주는 환경변수이다.",
            "`replace(/\\/+$/, '')`는 주소 끝의 `/`를 제거한다. `http://localhost:8080/`와 `http://localhost:8080`이 섞이면 URL이 `//api`처럼 깨질 수 있기 때문이다.",
            "`withoutAppSuffix`는 프론트엔드가 `/app` 아래에 배포되어도 API는 루트에서 호출해야 하는 상황을 처리한다.",
            "웹에서 localhost로 실행 중이고 별도 환경변수가 없으면 무조건 로컬 백엔드를 보도록 한다. 이 덕분에 Expo Web 개발 중에는 설정 없이 로컬 Spring Boot와 붙는다.",
            "이 함수는 `translate.ts`, `quiz.ts`, `users.ts`, `auth.ts` 등 거의 모든 API 클라이언트의 기반이다. 여기서 주소가 틀리면 앱 전체 API 호출이 실패한다.",
        ],
        max_lines=70,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/base-url.ts",
        "export function resolveBackendUrl",
        "resolveBackendUrl: 상대 경로와 로컬 URL을 실제 재생 URL로 바꾸기",
        [
            "백엔드가 `/clips/q001.mp4`처럼 상대 경로를 반환하면 프론트는 그 앞에 백엔드 origin을 붙여야 영상을 재생할 수 있다.",
            "`if (!rawUrl)`와 `if (!trimmed)`는 빈 URL을 안전하게 처리한다. 영상이 없는 단어는 빈 문자열로 남겨 화면에서 별도로 처리할 수 있다.",
            "`trimmed.startsWith('/')`이면 `baseUrl + trimmed`로 절대 URL을 만든다. 예를 들어 `/clips/a.mp4`가 `http://localhost:8080/clips/a.mp4`가 된다.",
            "`new URL(trimmed)`로 이미 절대 URL인지 검사한다. Firebase Storage URL처럼 완전한 URL이면 그대로 쓸 수 있다.",
            "이미 절대 URL이지만 origin이 로컬 백엔드라면 현재 설정된 `baseUrl`로 바꾼다. 이 부분은 개발/배포 주소가 섞였을 때 영상을 올바른 서버에서 가져오게 하는 보정 로직이다.",
            "`catch`에서는 URL 파싱이 안 되는 일반 문자열을 백엔드 상대 경로로 보고 붙인다. 즉 입력이 조금 이상해도 최대한 작동하도록 만든 방어 코드이다.",
        ],
        max_lines=80,
    )

    doc.add_paragraph("4. 로그인 상태는 앱 전체에서 어떻게 공유되는가", style="Heading 1")
    add_code_block_from_range(
        doc,
        "frontendcodes/context/auth-context.tsx",
        "const ACCESS_TOKEN_KEY",
        "AuthContext에서 사용하는 저장 key와 플랫폼별 저장소",
        [
            "`ACCESS_TOKEN_KEY`, `USER_KEY`, `GUEST_MODE_KEY`는 앱 저장소에 값을 넣을 때 사용하는 이름표이다. 이 key가 바뀌면 기존 저장된 로그인 정보와 호환되지 않는다.",
            "`storage` 객체는 웹과 모바일의 저장 방식을 하나의 인터페이스로 감싼다. 화면 코드는 `storage.get`, `storage.set`, `storage.remove`만 쓰면 되고, 내부가 localStorage인지 SecureStore인지 신경 쓰지 않는다.",
            "`Platform.OS === 'web'`이면 브라우저의 `window.localStorage`를 사용한다. 웹은 Expo SecureStore를 사용할 수 없기 때문이다.",
            "모바일에서는 `expo-secure-store`를 사용한다. accessToken은 민감한 값이므로 일반 AsyncStorage보다 SecureStore가 더 적합하다.",
            "각 저장 함수가 try/catch를 갖는 이유는 저장소 오류가 나도 앱의 메모리 상태는 계속 갱신할 수 있게 하기 위해서이다.",
        ],
        end_pattern="type AuthContextValue",
        max_lines=70,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/context/auth-context.tsx",
        "export function AuthProvider",
        "AuthProvider: 토큰, 사용자, 게스트 모드를 앱 전체에 공급",
        [
            "`useState` 네 개가 인증 상태의 핵심이다. `isLoading`은 초기 복원 중인지, `accessToken`은 로그인 토큰, `user`는 사용자 프로필, `isGuest`는 게스트 사용 여부를 뜻한다.",
            "첫 번째 `useEffect`의 `bootstrap()`은 앱이 켜질 때 저장소에서 이전 로그인 정보를 복원한다. 이 과정이 끝나기 전에는 라우팅 판단을 하면 안 되므로 `isLoading`을 사용한다.",
            "`Promise.all`로 token, user, guest 값을 동시에 읽는다. 순서대로 기다리는 것보다 빠르고, 세 값이 모두 있어야 인증 상태를 정확히 판단할 수 있다.",
            "`signInWithGoogleIdToken`은 백엔드에 Google idToken을 보내고 accessToken과 user를 받은 뒤 저장소에 저장한다.",
            "`signInWithBackendUser`는 login.tsx가 이미 백엔드에서 토큰과 user를 받은 경우 사용한다. 같은 인증 상태 저장 로직을 재사용한다.",
            "`continueAsGuest`는 accessToken 없이도 앱을 둘러볼 수 있도록 한다. 다만 user.id가 필요한 API는 사용할 수 없거나 제한된다.",
            "`value`를 `useMemo`로 감싼 이유는 Context value 객체가 매 렌더마다 새로 만들어져 불필요한 리렌더가 생기는 것을 줄이기 위해서이다.",
            "`useAuth()`는 Context를 쉽게 꺼내는 custom hook이다. AuthProvider 밖에서 쓰면 에러를 던져 잘못된 사용을 빠르게 발견하게 한다.",
        ],
        max_lines=130,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/app/_layout.tsx",
        "function AuthGate()",
        "AuthGate: 로그인하지 않은 사용자를 /login으로 보내는 라우팅 문지기",
        [
            "`useSegments()`는 현재 Expo Router 경로 조각을 가져온다. 예를 들어 `/home`이면 첫 조각이 `home`이다.",
            "`useAuth()`에서 `isLoading`, `accessToken`, `isGuest`를 가져온다. 화면 접근 가능 여부는 이 세 값으로 결정된다.",
            "`if (isLoading) return;`은 저장소 복원이 끝나기 전에 화면을 이동시키지 않기 위한 보호 코드이다.",
            "`inLoginScreen`, `inSignupScreen`은 인증이 없어도 접근 가능한 화면인지 판단한다.",
            "`canUseApp`은 토큰이 있거나 게스트 모드이면 true이다. 즉 정식 로그인과 게스트 접근을 모두 허용한다.",
            "인증이 없고 로그인/회원가입 화면도 아니라면 `router.replace('/login')`으로 보낸다. `replace`를 쓰면 뒤로가기로 보호 화면에 다시 돌아가는 것을 줄일 수 있다.",
            "이 코드 덕분에 home, quiz, translator, mypage 같은 개별 화면이 매번 로그인 검사를 반복하지 않아도 된다.",
        ],
        max_lines=100,
    )

    doc.add_paragraph("5. Google 로그인은 프론트와 백엔드에서 어떻게 이어지는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/auth.ts",
        "export async function loginWithGoogleToken",
        "프론트 API 클라이언트: Google idToken을 백엔드로 전달",
        [
            "`loginWithGoogleToken`은 화면 컴포넌트가 직접 fetch 세부사항을 알지 않도록 API 호출을 함수로 감싼 것이다.",
            "`getBaseUrl()`을 사용하므로 개발/배포 환경에 따라 백엔드 주소가 자동으로 달라진다.",
            "`method: 'POST'`는 로그인처럼 데이터를 서버에 보내 처리하는 요청에 적합하다.",
            "`Content-Type: application/json`은 request body가 JSON임을 백엔드에 알려준다.",
            "`body: JSON.stringify({ idToken })`에서 Google이 발급한 idToken을 백엔드로 보낸다. 백엔드는 이 토큰이 진짜 Google 토큰인지 검증한다.",
            "`if (!response.ok)`는 200번대 응답이 아니면 예외를 던진다. 화면에서는 이 예외를 잡아 로그인 실패 메시지를 보여줄 수 있다.",
            "성공하면 accessToken, refreshToken, user 정보가 들어 있는 JSON을 반환한다.",
        ],
        max_lines=50,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/auth/AuthController.java",
        "@PostMapping(\"/google\")",
        "백엔드 인증 컨트롤러: Google 토큰 검증 후 JWT 발급",
        [
            "`@PostMapping(\"/google\")` 때문에 이 메서드는 `POST /api/auth/google` 요청을 처리한다. 클래스 위의 `@RequestMapping(\"/api/auth\")`와 합쳐진 결과이다.",
            "`@RequestBody GoogleAuthRequest request`는 프론트가 보낸 JSON `{ idToken: ... }`을 Java record로 변환해 받는다.",
            "`googleAuthService.verify(request.idToken())`는 Google 토큰이 유효한지 검증하고, Google 사용자 정보인 sub, email, name, picture를 꺼낸다.",
            "`jwtService.generateAccessToken(user)`는 앱이 이후 API 요청에서 사용할 access token을 만든다.",
            "`jwtService.generateRefreshToken(user.sub())`는 access token보다 긴 수명의 refresh token을 만든다.",
            "`AuthResponse`는 프론트가 받기 쉬운 JSON 구조이다. tokenType은 보통 `Bearer`이고, expiresIn은 access token 만료 시간이다.",
            "이 메서드는 로그인 상태를 직접 저장하지 않는다. 백엔드는 토큰을 발급만 하고, 저장은 프론트의 AuthContext가 담당한다.",
        ],
        max_lines=60,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/auth/JwtService.java",
        "public String generateAccessToken",
        "JwtService: accessToken 안에 사용자 정보를 넣고 서명하기",
        [
            "`Instant now`는 토큰 발급 시각이다. 토큰에는 보통 발급 시각과 만료 시각을 함께 넣는다.",
            "`expiresAt`은 현재 시각에 access token 유효 시간을 더한 값이다. 기본 설정은 `jwt.access-exp-min`으로 제어된다.",
            "`Jwts.builder()`는 JWT를 만들기 위한 builder이다.",
            "`.subject(user.sub())`는 JWT의 주체를 Google 사용자 고유 id로 설정한다. 나중에 토큰 검증 시 '이 토큰이 누구의 것인지' 판단하는 기준이 된다.",
            "`.claim(\"email\", ...)`, `.claim(\"name\", ...)`, `.claim(\"picture\", ...)`는 토큰 안에 추가 사용자 정보를 넣는다.",
            "`.issuedAt`과 `.expiration`은 토큰의 시간 정보를 넣는다. 만료 시간이 지나면 토큰은 더 이상 유효하지 않아야 한다.",
            "`.signWith(key)`는 서버 secret key로 토큰에 서명한다. 서명이 있어야 클라이언트가 토큰 내용을 마음대로 바꿨는지 검출할 수 있다.",
            "`compact()`가 최종 문자열 형태의 JWT를 만든다.",
        ],
        max_lines=70,
    )

    doc.add_paragraph("6. 번역 버튼을 누르면 어떤 코드가 실행되는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/translate.ts",
        "export async function translateText",
        "프론트 번역 API 함수: 입력 문장을 /translate로 전송",
        [
            "`translateText(text)`는 translator 화면이 백엔드 번역 API를 호출할 때 쓰는 함수이다.",
            "`fetch(`${getBaseUrl()}/translate`...)`는 Spring Boot의 `TranslateController`로 요청을 보낸다.",
            "body에는 `{ text }`만 들어간다. 즉 프론트는 문장을 보내고, 백엔드는 그 문장을 분석해 수어 영상 목록을 만들어 준다.",
            "응답이 실패하면 예외를 던진다. translator 화면의 `handleTranslate`는 이 예외를 잡아 에러 메시지를 표시한다.",
            "`clips`와 `items`의 URL을 `resolveOptionalVideoUrl`로 보정한다. 백엔드가 상대 경로를 보내든 Firebase 절대 URL을 보내든 프론트에서 재생 가능한 형태로 맞춘다.",
            "`items`가 없으면 `clips`를 기반으로 기본 items를 만든다. 이것은 백엔드 응답 구조가 확장되기 전 데이터와도 호환되게 하는 코드이다.",
            "`hasVideo: item.hasVideo && !!item.url`은 hasVideo가 true라도 실제 URL이 비어 있으면 재생 불가능으로 처리한다.",
        ],
        max_lines=95,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/api/TranslateController.java",
        "public TranslateResponse translate",
        "백엔드 번역 컨트롤러: 요청을 서비스로 넘기는 얇은 계층",
        [
            "`@PostMapping(\"/translate\")`는 이 메서드가 `POST /translate` 요청을 받는다는 뜻이다.",
            "`@RequestBody TranslateRequest request`는 JSON body를 Java 객체로 변환한다. 프론트에서 `{ text: '...' }`를 보내면 `request.text()`로 꺼낼 수 있다.",
            "`request == null ? \"\" : request.text()`는 request 자체가 null인 예외 상황을 막는다. API는 외부 입력을 받기 때문에 항상 null 방어가 필요하다.",
            "`translationService.translate(text)`가 실제 번역 로직이다. 컨트롤러는 로직을 직접 구현하지 않고 서비스에 위임한다.",
            "이 구조를 쓰면 API 경로가 바뀌거나 요청/응답 형식이 바뀔 때는 컨트롤러를 보고, 번역 알고리즘을 바꿀 때는 서비스만 보면 된다.",
        ],
        max_lines=40,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/TranslationService.java",
        "public TranslateResponse translate",
        "TranslationService.translate: 번역 파이프라인 전체를 연결하는 핵심 메서드",
        [
            "`safeInput`은 입력이 null이어도 빈 문자열로 처리한다. 서버 로직에서 null이 그대로 들어가면 trim, stream 처리 중 예외가 날 수 있다.",
            "`signSentenceSimplifier.simplify(safeInput)`는 규칙 기반으로 문장을 수어 검색용 단어 목록으로 바꾸는 첫 번째 분석 경로이다.",
            "`ruleTokens`는 규칙 기반 결과를 다시 `TextNormalizer`로 정규화한 목록이다. 단어 끝의 활용형이나 문장부호를 정리해 사전 검색률을 높인다.",
            "`etriTokens`는 ETRI 형태소 분석 결과에서 의미 있는 lemma를 뽑은 목록이다. 규칙 기반 분석이 놓친 형태소를 보완한다.",
            "`openAiResult`는 OpenAI를 이용한 선택적 보정 결과이다. API key가 없거나 실패하면 Optional.empty가 되며, 전체 번역은 계속 진행된다.",
            "`chooseTokenStream`은 rule, etri, openai 후보 중 내부 사전에 가장 많이 맞는 목록을 고른다. 이 프로젝트에서는 '영상으로 보여줄 수 있는 단어가 많은가'가 좋은 분석 결과의 기준이다.",
            "`resolveTokens`는 선택된 토큰 중 사전에 없는 단어를 외부 사전 후보나 문맥 후보로 다시 보정한다.",
            "`clips`는 실제 재생 가능한 영상만 담고, `items`는 영상이 없는 단어까지 순서 보존용으로 담는다.",
            "`storageVideoCache.findUrl(token)`이 Firebase Storage에서 실제 영상 URL을 찾는다. 사전에는 있어도 영상 파일이 없을 수 있으므로 별도 분기가 필요하다.",
            "`unknown`은 사전에도 없는 단어이고, `noVideoWords`는 사전에는 있지만 영상이 없는 단어이다. 둘을 구분해야 나중에 사전을 늘릴지 영상을 업로드할지 판단할 수 있다.",
            "마지막 `TranslateResponse`는 프론트가 영상 재생, 결과 표시, 저장 기능을 모두 처리할 수 있도록 원본 입력, 단순화 문장, 토큰, 규칙, 메타데이터, 영상 목록, 누락 단어를 한 번에 반환한다.",
        ],
        max_lines=120,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/TranslationService.java",
        "private TokenStreamChoice chooseTokenStream",
        "chooseTokenStream: rule/etri/openai 중 가장 좋은 토큰 목록 고르기",
        [
            "`choices` 리스트에는 비어 있지 않은 후보만 들어간다. 빈 후보를 비교 대상에 넣으면 점수 계산이 의미 없어지기 때문이다.",
            "ruleTokens는 직접 만든 규칙 엔진 결과, etriTokens는 외부 형태소 분석 결과, openAiTokens는 LLM 기반 보정 결과이다.",
            "`countDictionaryHits`는 후보 토큰 중 내부 사전에 존재하는 단어 수를 센다.",
            "사전 hit 수가 가장 큰 후보가 선택된다. 이 서비스의 목적은 결국 영상을 재생하는 것이므로, 내부 사전에 많이 맞을수록 실제 출력 가능성이 높다.",
            "hit 수가 같으면 `shouldPreferTie`로 우선순위를 적용한다. 현재 우선순위는 openai > etri > rule이다.",
            "외부 분석 결과가 항상 맞는 것은 아니기 때문에 무조건 OpenAI를 쓰지 않고, 사전 hit 수를 기준으로 실제 프로젝트 데이터에 더 잘 맞는 결과를 고른다.",
        ],
        max_lines=90,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/TranslationService.java",
        "private List<String> resolveTokens",
        "resolveTokens: 사전에 없는 토큰을 다시 살리는 후처리",
        [
            "`contextDictionaryWords`는 ETRI/OpenAI 문맥 후보 중 내부 사전에 실제로 있는 단어만 모아 둔 목록이다.",
            "반복문은 최종 선택된 토큰을 하나씩 본다. 이미 사전에 있으면 그대로 결과에 넣는다.",
            "사전에 없는 경우 `unknownTokenResolverService.resolveToken(token)`을 호출해 외부 사전 후보와 내부 사전을 비교한다.",
            "그래도 못 찾으면 `pickBestContextMatch`를 사용한다. 예를 들어 선택된 토큰은 약간 틀렸지만 문맥 후보에 비슷한 사전 단어가 있으면 그쪽으로 교체한다.",
            "마지막까지 해결하지 못하면 원래 token을 넣는다. 그래야 프론트에서 unknown으로 표시하거나 사용자에게 빠진 단어를 알려줄 수 있다.",
        ],
        max_lines=90,
    )

    doc.add_paragraph("7. 문장을 수어 검색용 단어로 바꾸는 규칙 엔진", style="Heading 1")
    add_code_block_from_range(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/SignSentenceSimplifier.java",
        "private static final Set<String> TIME_WORDS",
        "역할 분류용 단어 집합과 조사 목록",
        [
            "`TIME_WORDS`는 '오늘', '내일', '지금'처럼 시간 역할을 할 가능성이 높은 단어를 모아 둔 집합이다.",
            "`QUESTION_WORDS`는 '어디', '무엇', '언제' 같은 의문 단어이다. 이 단어가 발견되면 metadata의 question이 true가 되고, 의문어 재배치 규칙이 적용된다.",
            "`PLACE_WORDS`는 장소 또는 방향 표현이다. 수어 표현에서는 장소 정보가 문장 앞쪽에 배치되는 경우가 많기 때문에 별도로 분류한다.",
            "`PREDICATE_WORDS`는 동작이나 상태를 나타내는 단어이다. 문장의 서술어 역할로 분류된다.",
            "`PARTICLES`, `SUBJECT_PARTICLES`, `OBJECT_PARTICLES`, `PLACE_PARTICLES`는 조사를 제거하거나 역할을 추정할 때 사용한다.",
            "이 코드는 완전한 한국어 형태소 분석기는 아니지만, 프로젝트 목적에 맞는 '작고 빠른 규칙 기반 분석기' 역할을 한다.",
        ],
        end_pattern="private final TextNormalizer textNormalizer",
        max_lines=70,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/SignSentenceSimplifier.java",
        "public SimplificationResult simplify",
        "simplify: 토큰을 역할별로 나누고 수어식 순서로 재조립",
        [
            "`safeInput`은 null과 앞뒤 공백을 정리한 입력이다. 빈 문장이면 빈 결과를 바로 반환한다.",
            "`appliedRules`는 어떤 변환 규칙이 적용되었는지 기록하는 집합이다. LinkedHashSet이라 중복 없이 적용 순서를 유지한다.",
            "`timeTokens`, `placeTokens`, `subjectTokens`, `objectTokens`, `predicateTokens`, `questionTokens`는 문장 성분별 바구니이다.",
            "`question`, `negative`, `tense`는 문장의 메타 정보이다. 프론트에서 직접 크게 쓰지 않더라도 디버깅과 설명 가능한 응답에 도움이 된다.",
            "`for (String rawToken : tokenize(safeInput))`에서 문장을 토큰 단위로 분석한다. 각 토큰은 `analyze` 메서드로 들어간다.",
            "`analysis.removedParticle()`이 true이면 조사가 제거되었다는 뜻이고 `particle_removal` 규칙이 기록된다.",
            "부정 표현이 standalone marker이면 실제 토큰 목록에는 넣지 않고 negative만 true로 만든다. 예를 들어 '안 가요'의 '안'은 부정 의미를 만드는 표지이다.",
            "과거/미래가 감지되면 tense를 갱신하고 `tense_separation` 규칙을 기록한다.",
            "의문어는 별도 questionTokens에 넣고 마지막 쪽에 배치한다. 수어 표현에서 의문 요소가 문장 끝에 오는 경우를 반영한 구조이다.",
            "마지막에 시간 -> 장소 -> 주어 -> 목적어 -> 서술어 -> 부정 -> 시제 -> 의문어 순서로 합친다. 이것이 이 프로젝트의 수어식 단순화 결과이다.",
            "`SimplificationResult`는 단순화된 문장, 토큰 목록, 적용 규칙, 메타데이터를 모두 담는다. TranslationService는 이 결과를 다음 단계에서 사전 매칭에 사용한다.",
        ],
        max_lines=130,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/SignSentenceSimplifier.java",
        "private TokenAnalysis analyze",
        "analyze: 단어 하나를 정규화하고 역할을 추정",
        [
            "이 메서드는 문장 전체가 아니라 토큰 하나를 분석한다. 문장 처리에서 가장 작은 분석 단위라고 보면 된다.",
            "`stripPunctuation`은 앞뒤 문장부호를 제거한다. 사용자는 '어디?'처럼 입력할 수 있으므로 문장부호가 사전 검색을 방해하면 안 된다.",
            "`standaloneNegativeMarker`는 '안', '못' 같은 독립 부정 표지를 감지한다. 이런 표지는 단어 영상으로 직접 찾기보다 부정 metadata로 처리한다.",
            "`detectParticleRole(token)`은 토큰 끝의 조사를 보고 주어/목적어/장소 역할을 추정한다. 예를 들어 '학교에서'는 장소 역할일 가능성이 높다.",
            "`stripParticles`와 `stripCommonParticles`는 조사 제거 단계이다. 사전에는 보통 '학교'가 있지 '학교에서'가 있지는 않다.",
            "`resolveKnownWord(stripped)`는 정규화 후 내부 사전에 있는 단어로 맞추려 한다.",
            "`questionWord`, `negativeWord`, `past`, `future`는 해당 토큰이 문장 전체 metadata에 영향을 주는지 판단한다.",
            "`inferRole`은 최종적으로 TIME, PLACE, SUBJECT, OBJECT, PREDICATE 중 하나를 고른다.",
            "결과는 `TokenAnalysis` record로 반환된다. record를 쓰면 분석 결과를 불변 데이터 묶음처럼 안전하게 전달할 수 있다.",
        ],
        max_lines=85,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/normalizer/TextNormalizer.java",
        "public List<String> normalizeTokens",
        "TextNormalizer.normalizeTokens: 문장 전체를 사전 검색용 토큰으로 정리",
        [
            "`input == null || input.isBlank()`는 빈 입력을 안전하게 빈 리스트로 처리한다.",
            "`replaceAll`은 문장부호를 공백으로 바꾼다. 토큰 사이가 붙는 것을 막기 위해 제거가 아니라 공백 치환을 사용한다.",
            "`split(\"\\\\s+\")`는 공백이 하나든 여러 개든 토큰을 나눈다.",
            "각 raw token은 `normalizeToken`을 거친다. 즉 문장 전체 정규화의 핵심은 단어 하나 정규화 함수를 반복 적용하는 것이다.",
            "정규화 결과가 blank이면 결과 목록에 넣지 않는다. 불필요한 빈 토큰이 뒤쪽 사전 검색에 들어가면 오류나 노이즈가 생긴다.",
        ],
        max_lines=45,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/normalizer/TextNormalizer.java",
        "public String normalizeToken",
        "TextNormalizer.normalizeToken: 단어 하나를 대표형으로 변환",
        [
            "`rawToken.trim()`은 앞뒤 공백을 제거한다. 공백이 붙어 있으면 Map 검색이 실패한다.",
            "`REPLACEMENTS.get(token)`은 직접 치환 사전이다. 자주 나오는 표현은 규칙보다 명시적 매핑이 더 정확할 때가 있다.",
            "`normalizeCommonKoreanEnding`은 흔한 한국어 어미를 기본형에 가깝게 바꾼다.",
            "`normalizePredicateEnding`은 동사/형용사 표현을 사전형 후보로 바꾼다.",
            "마지막 `REPLACEMENTS.getOrDefault(token, token)`은 앞 단계에서 안 바뀐 단어도 혹시 치환 사전에 있으면 바꾸고, 아니면 원래 단어를 유지한다.",
            "이 메서드는 번역 품질에 직접 영향을 준다. 정규화가 잘되면 내부 사전 hit가 올라가고, 영상으로 표현 가능한 단어가 늘어난다.",
        ],
        max_lines=45,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/dictionary/DictionaryLoader.java",
        "public void load",
        "DictionaryLoader.load: sign_dictionary.json을 메모리 사전으로 만들기",
        [
            "`@PostConstruct`가 붙은 `load()`는 Spring이 이 컴포넌트를 만든 직후 자동 실행된다. 즉 첫 번역 요청이 오기 전에 사전이 준비된다.",
            "`ClassPathResource(\"sign_dictionary.json\")`는 `src/main/resources` 안의 JSON 파일을 읽는다.",
            "`objectMapper.readValue`는 JSON 배열을 `List<SignDictionaryEntry>`로 변환한다.",
            "`mutableMap.put(entry.word(), entry)`는 단어를 key로 하는 Map을 만든다. 단어 검색을 빠르게 하기 위한 핵심이다.",
            "`Collections.unmodifiableList`와 `unmodifiableMap`은 로딩 후 사전이 실수로 변경되지 않게 한다.",
            "실패하면 `IllegalStateException`을 던진다. 사전이 없으면 번역 기능이 제대로 작동할 수 없으므로 서버 시작 단계에서 빨리 실패하는 편이 낫다.",
        ],
        max_lines=70,
    )

    doc.add_paragraph("8. 영상 URL은 어떻게 찾아오는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/storage/StorageVideoCache.java",
        "public String findUrl",
        "StorageVideoCache.findUrl: 단어를 Firebase Storage 영상 URL로 바꾸기",
        [
            "`word == null || word.isBlank()`이면 검색할 단어가 없으므로 null을 반환한다.",
            "`normalizedWord`는 앞뒤 공백을 제거한 실제 검색 키이다.",
            "`cache.get(normalizedWord)`는 이전에 성공적으로 찾은 URL을 바로 반환한다. 같은 단어를 다시 찾을 때 Storage API를 호출하지 않아도 된다.",
            "`missingWords.contains(normalizedWord)`는 이전에 못 찾은 단어를 다시 검색하지 않게 한다. 실패 결과도 캐싱하는 셈이다.",
            "`loadUrl(normalizedWord)`가 실제 Storage 조회를 수행한다. 이 작업은 느릴 수 있으므로 캐시 앞단이 중요하다.",
            "URL을 찾지 못하면 `missingWords`에 넣고 null을 반환한다. TranslationService는 이 null을 보고 noVideoWords 또는 unknown 처리를 한다.",
            "`cache.putIfAbsent`는 동시에 여러 요청이 같은 단어를 찾을 때도 안전하게 캐시를 갱신한다. 서버는 멀티스레드로 요청을 처리하므로 ConcurrentHashMap 사용이 적절하다.",
        ],
        max_lines=70,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/storage/StorageVideoCache.java",
        "private String loadUrl",
        "loadUrl: Storage prefix를 돌며 파일명에서 단어를 매칭",
        [
            "`folderPrefixes`를 반복하는 이유는 Storage 폴더 경로가 하나로 고정되어 있지 않기 때문이다. 코드에는 기존 오타 경로와 올바른 경로를 모두 지원하는 호환 로직이 있다.",
            "`searchPrefix = folderPrefix + word`는 단어로 시작하는 파일만 검색하기 위한 prefix이다.",
            "`storage.list(... prefix(searchPrefix))`는 bucket 안에서 해당 prefix로 시작하는 객체들을 가져온다.",
            "`blobName.endsWith(\"/\")`는 폴더 자체나 빈 객체를 건너뛰기 위한 조건이다.",
            "`fileName = blobName.substring(folderPrefix.length())`는 폴더 경로를 떼고 실제 파일명만 남긴다.",
            "`extractWord(fileName)`은 파일명에서 확장자와 언더스코어 뒤쪽을 제거해 단어만 뽑는다. 예를 들어 `학교_001.mp4`는 `학교`가 된다.",
            "입력 단어와 추출 단어가 같으면 `buildDownloadUrl(blob)`으로 재생 가능한 URL을 만든다.",
            "Storage 조회 실패는 로그로 남기고 다음 prefix를 시도한다. 하나의 경로 문제가 전체 번역을 멈추지 않게 하는 방어적 구조이다.",
        ],
        max_lines=90,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/storage/StorageVideoCache.java",
        "private String buildDownloadUrl",
        "buildDownloadUrl: Firebase 다운로드 URL 또는 signed URL 만들기",
        [
            "`urlEncode(blob.getName())`은 Storage 객체 경로에 한글이나 공백이 있어도 URL에서 안전하게 쓰기 위한 인코딩이다.",
            "`https://firebasestorage.googleapis.com/v0/b/.../o/...` 형태는 Firebase Storage의 직접 media URL 형식이다.",
            "`findDownloadToken(blob)`은 파일 metadata에 저장된 Firebase 다운로드 토큰을 찾는다.",
            "토큰이 있으면 `&token=...`을 붙여 일반 클라이언트가 접근 가능한 URL을 만든다.",
            "토큰이 없으면 `buildSignedUrlOrFallback`을 시도한다. signed URL은 서버 인증으로 일정 기간 접근 가능한 URL을 만들어 주는 방식이다.",
            "signed URL 생성도 실패하면 fallback direct media URL을 반환한다. 이 경우 Storage 보안 규칙에 따라 재생이 안 될 수 있으므로 로그 확인이 필요하다.",
        ],
        max_lines=80,
    )

    doc.add_paragraph("9. 외부 형태소 분석과 OpenAI 보정은 왜 들어가는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/ExternalLexiconApiClient.java",
        "public List<String> fetchCandidates",
        "fetchCandidates: 사전에 없는 단어를 살리기 위한 외부 후보 수집",
        [
            "`LinkedHashSet`을 쓰는 이유는 중복 제거와 순서 유지가 동시에 필요하기 때문이다.",
            "ETRI, 우리말샘, 한국어기초사전 결과를 모두 합친다. 한 API가 못 찾은 단어를 다른 API가 찾을 수 있기 때문이다.",
            "`candidates.removeIf(String::isBlank)`는 빈 문자열 후보를 제거한다. 빈 후보가 뒤쪽 사전 매칭에 들어가면 의미 없는 검색이 된다.",
            "이 메서드의 결과는 `UnknownTokenResolverService`에서 내부 사전 단어와 비교된다.",
            "중요한 점은 외부 API 후보를 그대로 사용자에게 보여주기보다, 내부 사전과 맞는 후보만 실제 번역 토큰으로 사용한다는 것이다.",
        ],
        max_lines=35,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/ExternalLexiconApiClient.java",
        "private List<String> fetchEtriLemmas",
        "fetchEtriLemmas: ETRI 형태소 분석 응답에서 lemma 추출",
        [
            "`etriEnabled`와 `etriAccessKey` 검사로 API 사용 가능 여부를 먼저 확인한다. key가 없으면 빈 리스트를 반환해 전체 번역이 계속 진행되게 한다.",
            "`analysis_code`가 `morp`인 요청을 보내 형태소 분석을 받는다.",
            "`objectMapper.readTree(response)`로 JSON 응답을 트리 구조로 읽는다.",
            "`return_object.sentence[].morp[]` 안에 형태소 분석 결과가 들어 있다고 보고 순회한다.",
            "`contentOnly`가 true이면 조사, 어미 같은 기능어를 제외하고 명사/동사/형용사 같은 내용어만 남긴다.",
            "`expandLemmaByPos`는 동사/형용사 lemma를 사전형 후보로 확장한다. 한국어 동사는 어간만으로는 내부 사전과 안 맞을 수 있기 때문이다.",
            "예외가 나면 빈 리스트를 반환한다. 외부 API 실패가 번역 전체 실패로 이어지지 않도록 하는 설계이다.",
        ],
        max_lines=90,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/service/OpenAiMorphologyNormalizerService.java",
        "public Optional<MorphologyNormalizationResult> normalize",
        "OpenAI normalize: 선택적으로 문장 토큰화를 보정",
        [
            "`Optional`을 반환하는 이유는 OpenAI 보정이 항상 가능한 기능이 아니기 때문이다. API key가 없거나 호출이 실패할 수 있다.",
            "`!enabled || apiKey == null || apiKey.isBlank()` 조건이면 즉시 empty를 반환한다. 설정만으로 OpenAI 기능을 끌 수 있다.",
            "`postForObject`는 Chat Completions API에 요청을 보낸다.",
            "`choices[0].message.content`에서 모델 응답 텍스트를 꺼낸다. 이 프로젝트는 JSON 응답을 기대하므로 다시 `objectMapper.readTree(content)`를 호출한다.",
            "`extractTokens` 결과가 비어 있으면 사용할 수 없는 응답으로 보고 empty를 반환한다.",
            "`MorphologyNormalizationResult`에는 simplifiedSentence, tokens, tense가 들어간다. TranslationService는 이 tokens를 rule/etri 결과와 비교한다.",
            "이 코드의 핵심은 OpenAI가 실패해도 번역 기능 전체가 죽지 않는 것이다. 그래서 모든 실패는 Optional.empty로 흡수된다.",
        ],
        max_lines=75,
    )

    doc.add_paragraph("10. 퀴즈 화면은 어떤 코드로 움직이는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/quiz.ts",
        "export async function fetchQuizSession",
        "fetchQuizSession: 일반 퀴즈 문제 목록 가져오기",
        [
            "`count`는 가져올 문제 수이다. 기본값은 10이다.",
            "`URLSearchParams`를 사용하면 query string을 안전하게 만들 수 있다.",
            "`category`가 있으면 query에 넣는다. learn 화면에서 기초 단어, 일상 회화 같은 코스를 구분할 때 사용된다.",
            "`GET /api/quiz/session`은 백엔드 QuizController의 session 메서드로 연결된다.",
            "응답의 각 문제에 대해 `videoUrl`을 `resolveBackendUrl`로 보정한다. 영상 URL은 프론트에서 재생 가능한 절대 URL이어야 한다.",
        ],
        max_lines=45,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/quiz.ts",
        "export async function submitQuizAnswer",
        "submitQuizAnswer: 사용자가 고른 답을 서버에 제출",
        [
            "프론트는 정답을 직접 판단하지 않고 백엔드로 `quizId`와 `selectedChoiceId`를 보낸다.",
            "body는 `{ quizId, selectedChoiceId }` 형태이다. selectedChoiceId는 A/B/C/D 중 하나이다.",
            "백엔드는 Firestore에서 해당 문제를 다시 읽고 correctChoiceId와 비교한다.",
            "응답에는 `isCorrect`, `correctChoiceId`, `correctChoiceText`가 포함된다. 화면은 이 값으로 정답/오답 피드백을 보여준다.",
            "프론트에 문제를 내려줄 때 정답을 포함하지 않고, 제출 후 서버가 판정하는 방식은 기본적인 클라이언트 조작 방지 구조이다.",
        ],
        max_lines=45,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/quiz/QuizController.java",
        "public QuizSessionResponse session",
        "QuizController.session: 일반 퀴즈 세션 API",
        [
            "`@GetMapping(\"/session\")`은 `GET /api/quiz/session` 요청을 처리한다.",
            "`@RequestParam(defaultValue = \"10\") int count`는 query string에 count가 없으면 10을 사용한다.",
            "`@RequestParam(required = false) String category`는 category가 없어도 요청을 허용한다.",
            "컨트롤러는 바로 `quizService.getSession(count, category)`를 호출한다. 문제 선택과 Firestore 조회는 서비스 책임이다.",
            "이처럼 Controller는 HTTP 입구, Service는 실제 로직이라는 역할 분리가 유지된다.",
        ],
        max_lines=35,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/quiz/QuizService.java",
        "public QuizSessionResponse getSession",
        "QuizService.getSession: 활성 문제를 섞어서 세션 만들기",
        [
            "`normalizeCount(count)`는 요청된 문제 수가 0 이하이거나 너무 클 때 안전한 값으로 보정한다.",
            "`getActiveQuestions(category)`는 Firestore에서 활성화된 문제를 읽고, category 조건이 있으면 필터링한다.",
            "`Collections.shuffle(questions)`는 문제 순서를 매번 섞는다. 사용자가 매번 같은 순서의 문제를 보는 것을 막는다.",
            "문제가 요청 수보다 많으면 `subList`로 앞에서 safeCount개만 사용한다.",
            "`QuizSessionResponse(questions.size(), questions)`는 실제 반환 문제 수와 문제 목록을 함께 보낸다.",
        ],
        max_lines=45,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/quiz/QuizService.java",
        "private QuizSessionQuestionResponse toSessionQuestion",
        "toSessionQuestion: Firestore 문서를 프론트가 쓸 수 있는 문제 DTO로 변환",
        [
            "`isActive`가 false이면 null을 반환해 비활성 문제를 제외한다.",
            "`questionText`, `videoUrl`, `choices`, `correctChoiceId`를 Firestore 문서에서 꺼낸다.",
            "`choiceTextById(choices, correctChoiceId)`는 A/B/C/D 정답 id를 실제 단어 텍스트로 바꾼다.",
            "`choices.size() != 4`이면 문제 형식이 깨진 것으로 보고 제외한다. 프론트 퀴즈 UI는 보기 4개를 전제로 만들어져 있다.",
            "`storageVideoCache.findUrlOrFallback(correctChoiceText, firestoreVideoUrl)`은 정답 단어 기준으로 Storage 영상을 우선 찾고, 없으면 Firestore의 기존 videoUrl을 사용한다.",
            "영상 URL이 비어 있으면 문제를 반환하지 않는다. 퀴즈는 영상을 보고 답을 맞히는 구조이므로 영상이 없으면 문제로 사용할 수 없다.",
            "`attemptCount`, `correctCount`, `correctRate`, `difficultyLevel`은 화면에서 난이도와 정답률 정보를 보여주기 위한 보조 데이터이다.",
        ],
        max_lines=95,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/quiz/QuizService.java",
        "public QuizAnswerResponse checkAnswer",
        "checkAnswer: 정답 판정과 문제 통계 증가",
        [
            "request, quizId, selectedChoiceId가 비어 있으면 BAD_REQUEST를 던진다. API 입력 검증은 서버에서 반드시 해야 한다.",
            "`getQuestionById(request.quizId())`로 Firestore의 실제 문제 문서를 읽는다. 클라이언트가 보내온 정보만 믿지 않는다.",
            "`correctChoiceId`를 문서에서 꺼내고, 대문자로 정규화한다.",
            "`selectedChoiceId`도 normalizeChoiceId로 대문자 정규화한다. 프론트가 `a`를 보내도 `A`와 비교할 수 있다.",
            "`isCorrect = correctChoiceId.equals(selectedChoiceId)`가 실제 정답 판정이다.",
            "`updateQuizStats`는 문제별 attempt_count, correct_count, wrong_count를 증가시킨다. 이 통계는 난이도나 정답률 표시 자료가 된다.",
            "응답에는 correctChoiceText까지 포함된다. 프론트는 오답일 때 정답 단어를 사용자에게 보여줄 수 있다.",
        ],
        max_lines=70,
    )

    doc.add_paragraph("11. 사용자의 학습 기록은 어디서 저장되는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/users.ts",
        "export async function recordQuizAttempt",
        "recordQuizAttempt: 프론트에서 사용자 풀이 기록 저장 요청",
        [
            "퀴즈 정답 판정은 `/api/quiz/answer`가 하지만, 사용자 개인 통계 저장은 `/api/users/{uid}/tryQuestion`이 담당한다.",
            "`uid`는 현재 로그인한 사용자 id이다. AuthContext의 user.id에서 온다.",
            "`questionId`와 `isCorrect`를 body로 보낸다. 서버는 이 정보로 총 풀이 수, 정답 수, 오답 목록, 일일 풀이 수를 갱신한다.",
            "응답으로 갱신된 UserInfo를 받는다. 필요하면 프론트 상태를 최신 사용자 통계로 바꿀 수 있다.",
        ],
        max_lines=45,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/user/controller/UserController.java",
        "public ResponseEntity<?> tryQuestion",
        "UserController.tryQuestion: 풀이 기록 API 입구",
        [
            "`@PatchMapping(\"/{uid}/tryQuestion\")`은 특정 사용자 문서 일부를 수정하는 API이다.",
            "body에서 `questionId`와 `isCorrect`를 꺼낸다. 둘 중 하나라도 없으면 bad request를 반환한다.",
            "`userService.tryQuestion(uid, questionId, isCorrect)`가 실제 Firestore 업데이트를 수행한다.",
            "사용자를 찾지 못하면 404, Firestore 오류가 나면 500을 반환한다.",
            "컨트롤러는 HTTP 상태 코드와 입력 검증을 담당하고, 통계 계산은 UserService에 맡긴다.",
        ],
        max_lines=55,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/user/service/UserService.java",
        "public UserInfo tryQuestion",
        "UserService.tryQuestion: 총 풀이 수, 정답 수, 오답 수, 일일 목표를 실제로 갱신",
        [
            "`DocumentReference docRef`는 `users/{uid}` 문서를 가리킨다.",
            "`document.exists()`로 사용자가 실제로 존재하는지 확인한다. 없는 사용자에게 기록을 남길 수 없기 때문이다.",
            "`document.toObject(UserInfo.class)`로 Firestore 문서를 Java 객체로 변환한다.",
            "`totalQuestionNum`은 문제를 풀 때마다 1 증가한다.",
            "`totalQuestions`는 사용자가 푼 문제 id 목록이다. 같은 문제를 여러 번 풀어도 목록에는 한 번만 넣는 구조이다.",
            "`todayKst`는 한국 시간 기준 날짜 key이다. 일일 목표는 한국 날짜 기준으로 계산해야 사용자가 기대하는 날짜와 맞다.",
            "`dailySolvedCounts`에서 오늘 값을 꺼내 1 증가시킨다. home과 learn 화면의 오늘 목표 진행률이 이 값으로 계산된다.",
            "정답이면 `correctQuestionNum`을 증가시킨다.",
            "오답이면 `incorrectQuestionCounts`에서 해당 문제의 오답 횟수를 증가시키고, `incorrectQuestions` 목록에 없으면 추가한다.",
            "`incorrectQuestionDates`는 처음 틀린 시각을 저장한다. 오답노트에서 언제 틀렸는지 표시할 수 있다.",
            "마지막에 updateData를 Firestore에 업데이트하고, 다시 userInfo를 반환한다.",
        ],
        max_lines=120,
    )

    doc.add_paragraph("12. 오답노트와 북마크는 어떻게 저장되는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/wrong-note-saved.ts",
        "export async function saveWrongNote",
        "saveWrongNote: 퀴즈 화면에서 오답노트 저장 요청",
        [
            "`uid`와 `quizId`가 있어야 어떤 사용자의 어떤 문제를 저장할지 알 수 있다.",
            "`POST /api/users/{uid}/wrong-note-saved`는 사용자의 저장된 오답노트 하위 컬렉션에 문서를 만든다.",
            "응답으로 받은 `videoUrl`은 다시 `resolveBackendUrl`로 보정한다. 오답노트 화면에서 바로 영상을 재생해야 하기 때문이다.",
            "프론트의 quiz.tsx는 이 함수를 호출한 뒤 저장 완료 애니메이션을 보여준다.",
        ],
        max_lines=45,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/user/service/UserService.java",
        "public WrongNoteSavedResponse saveWrongNote",
        "UserService.saveWrongNote: quiz_items에서 문제 정보를 복사해 사용자 오답노트에 저장",
        [
            "먼저 uid와 quizId가 비었는지 검사한다. 잘못된 요청을 Firestore까지 보내지 않기 위해서이다.",
            "`userRef`는 users/{uid} 문서이다. 사용자 문서가 없으면 예외를 던진다.",
            "`quizDoc`은 quiz_items/{quizId} 문서이다. 오답노트에 저장할 문제 내용은 원본 퀴즈 문서에서 가져온다.",
            "`questionText`, `choices`, `correctChoiceId`를 읽고, 정답 id로 실제 정답 단어를 구한다.",
            "`resolveQuizVideoUrl`로 오답노트에서 재생할 영상 URL을 구한다. Storage URL이 우선이고 기존 fallback URL도 고려된다.",
            "`wrongAt`은 사용자가 실제로 틀렸던 날짜를 우선 사용하고, 없으면 현재 날짜를 사용한다.",
            "`wrong_note_saved` 하위 컬렉션의 문서 id를 quizId로 사용한다. 같은 문제를 중복 저장해도 같은 문서를 덮어쓰는 구조가 된다.",
            "저장 후 다시 문서를 읽어 `WrongNoteSavedResponse`로 변환한다. 서버 timestamp는 set 직후 클라이언트가 알기 어려우므로 다시 읽는 방식이 안전하다.",
        ],
        max_lines=100,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/api/translator-bookmarks.ts",
        "export async function saveTranslatorBookmark",
        "saveTranslatorBookmark: 번역 문장 북마크 저장",
        [
            "번역기 북마크는 퀴즈 문제가 아니라 사용자가 입력한 문장을 저장한다.",
            "`payload`에는 sentence, word, videoUrl이 들어갈 수 있다. sentence는 필수이고 word/videoUrl은 대표 단어나 대표 영상이다.",
            "`POST /api/users/{uid}/translator-bookmarks`로 요청한다.",
            "응답으로 받은 videoUrl은 절대 URL로 보정한다. 북마크 목록에서 다시 재생하거나 미리보기할 수 있기 때문이다.",
            "translator.tsx는 서버 저장이 실패하면 로컬 저장소 fallback을 사용한다. 로그인하지 않은 사용자도 임시 저장 경험을 가질 수 있다.",
        ],
        max_lines=50,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/user/service/UserService.java",
        "public TranslatorBookmarkResponse saveTranslatorBookmark",
        "UserService.saveTranslatorBookmark: 사용자가 입력한 번역 문장을 Firestore에 저장",
        [
            "`sentence`가 비어 있으면 저장할 핵심 데이터가 없으므로 예외를 던진다.",
            "`userRef`로 사용자 문서를 확인한다. 존재하지 않는 사용자 하위에 북마크를 만들면 데이터가 고아 상태가 될 수 있다.",
            "`bookmarkId = \"txt_\" + UUID.randomUUID()`는 번역기 문장 북마크 id를 새로 만든다. 퀴즈 id와 충돌하지 않도록 `txt_` prefix를 붙인다.",
            "`questionText` 필드에 sentence를 넣는 것은 기존 북마크 응답 구조와 호환하기 위한 설계로 보인다.",
            "`source = \"translator\"`는 이 문서가 번역기에서 온 북마크임을 표시한다.",
            "`savedAt = FieldValue.serverTimestamp()`는 클라이언트 시간이 아니라 서버 시간을 저장한다.",
            "저장 후 문서를 다시 읽고 response DTO로 변환한다.",
        ],
        max_lines=90,
    )

    doc.add_paragraph("13. 번역기 화면의 영상 재생은 왜 복잡한가", style="Heading 1")
    add_code_block_from_range(
        doc,
        "frontendcodes/app/translator.tsx",
        "const [inputText",
        "translator.tsx의 핵심 상태들",
        [
            "`inputText`는 사용자가 입력한 원문이다.",
            "`loading`, `errorMessage`, `feedbackMessage`는 비동기 요청과 사용자 피드백을 위한 상태이다.",
            "`result`는 백엔드 `/translate` 응답 전체이다. 여기 안에 simplifiedSentence, tokens, clips, items가 들어 있다.",
            "`playbackUrlMap`은 동일 URL 중복 처리와 재생 소스 관리를 위해 사용된다.",
            "`currentClipIndex`는 현재 몇 번째 수어 영상을 재생 중인지 나타낸다.",
            "`webActiveSlot`과 `nativeActiveSlot`은 더블 버퍼링 구조에서 현재 화면에 보이는 비디오 슬롯이다.",
            "`webVideoReady`, `nativeVideoReady`는 각 슬롯의 영상이 재생 준비되었는지 나타낸다.",
            "이 상태들이 많은 이유는 번역기가 단일 영상 하나만 재생하는 것이 아니라 여러 단어 영상을 끊김 없이 이어 재생해야 하기 때문이다.",
        ],
        end_pattern="const filteredBookmarks",
        max_lines=80,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/app/translator.tsx",
        "async function handleTranslate",
        "handleTranslate: 입력 문장을 번역하고 재생 목록을 새로 세팅",
        [
            "`overrideText`가 있으면 그 값을 사용하고, 없으면 현재 inputText를 사용한다. 북마크 문장을 다시 번역할 때 재사용하기 쉬운 구조이다.",
            "`trim()` 후 빈 문자열이면 요청하지 않는다. 빈 번역 요청은 서버 리소스 낭비이고 사용자에게도 의미가 없다.",
            "`setLoading(true)`와 `setErrorMessage(null)`로 요청 시작 상태를 만든다.",
            "`translateText(text)`가 실제 백엔드 API 호출이다.",
            "`setResult(next)`로 응답 전체를 저장한다. 화면의 결과 패널과 영상 재생 영역이 이 상태를 바라본다.",
            "`(next.items ?? next.clips)`는 백엔드가 items를 주면 items를 쓰고, 없으면 clips로 fallback한다.",
            "`buildPlaybackUrlMap`은 중복 URL을 정리해 재생 URL 맵을 만든다.",
            "`setCurrentClipIndex(0)`은 새 번역 결과가 오면 첫 번째 영상부터 재생하도록 초기화한다.",
            "오류가 나면 사용자에게 실패 메시지를 보여주고, 마지막에 loading을 false로 되돌린다.",
        ],
        max_lines=95,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/app/translator.tsx",
        "function preloadNextNativeClip",
        "preloadNextNativeClip: 다음 네이티브 영상을 미리 준비",
        [
            "`fromIndex + 1`로 다음 클립 번호를 계산한다.",
            "`getPlaybackUrlAt(nextIndex)`가 다음 영상 URL을 가져온다. 없으면 더 이상 미리 로딩할 것이 없다.",
            "`preloadSlot`은 현재 활성 슬롯의 반대편 슬롯이다. 현재 A를 보여주면 B에 다음 영상을 준비하고, 현재 B를 보여주면 A에 준비한다.",
            "`replaceNativeSlot(preloadSlot, nextUrl, false)`는 다음 영상을 자동 재생하지 않고 미리 로딩한다.",
            "이 구조는 더블 버퍼링이다. 다음 영상을 미리 준비해 두면 현재 영상이 끝났을 때 화면 전환이 더 부드럽다.",
        ],
        max_lines=55,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/app/translator.tsx",
        "function switchToPreparedNativeClip",
        "switchToPreparedNativeClip: 준비된 다음 영상으로 슬롯 전환",
        [
            "`currentIndex`와 `nextIndex`를 계산해 다음 클립이 존재하는지 확인한다.",
            "`nextSlot`은 현재 활성 슬롯의 반대편 슬롯이다.",
            "준비된 슬롯에 URL이 있으면 그 슬롯을 활성 슬롯으로 바꾸고 재생한다.",
            "`runVideoSlideTransition(previousSlot)`은 이전 슬롯에서 다음 슬롯로 넘어가는 시각 효과를 만든다.",
            "`setCurrentClipIndex(nextIndex)`가 현재 단어 표시와 진행률을 다음 클립으로 맞춘다.",
            "전환 후 다시 `preloadNextNativeClip(nextIndex)`를 호출하면 다음다음 영상이 준비된다. 이런 식으로 A/B 슬롯을 번갈아 사용한다.",
        ],
        max_lines=85,
    )

    doc.add_paragraph("14. 홈, 학습, 마이페이지 통계는 어떤 데이터로 계산되는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/daily-goal.ts",
        "export function getKstDateKey",
        "getKstDateKey: 한국 날짜 기준 key 생성",
        [
            "`Intl.DateTimeFormat`에 `timeZone: 'Asia/Seoul'`을 지정해 한국 날짜 기준으로 year/month/day를 계산한다.",
            "서버나 사용자의 브라우저 시간이 다른 나라 시간대여도 한국 날짜 기준 학습 목표를 계산할 수 있다.",
            "반환 형식은 `YYYY-MM-DD`이다. Firestore의 `dailySolvedCounts` Map key와 맞춰 쓰기 좋다.",
            "이 함수는 home, learn, mypage에서 오늘 풀이 수와 연속 목표 달성일을 계산하는 기반이다.",
        ],
        max_lines=35,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/lib/daily-goal.ts",
        "export function getConsecutiveGoalDays",
        "getConsecutiveGoalDays: 연속 목표 달성일 계산",
        [
            "`cursor`는 오늘 날짜 key에서 시작한다.",
            "`while ((counts[cursor] ?? 0) >= target)`는 해당 날짜의 풀이 수가 목표 이상인 동안 반복한다.",
            "목표를 달성한 날이면 streak를 1 증가시키고, `previousDateKey`로 전날로 이동한다.",
            "목표를 못 채운 날짜를 만나면 반복을 멈춘다. 그래서 '오늘부터 연속으로 며칠 달성했는지'가 계산된다.",
            "이 값은 마이페이지나 홈에서 사용자의 학습 습관을 보여주는 지표가 된다.",
        ],
        max_lines=60,
    )
    add_code_block_from_method(
        doc,
        "frontendcodes/app/mypage.tsx",
        "const loadUserInfo",
        "mypage.tsx: 사용자 정보와 많이 틀린 단어를 같이 불러오기",
        [
            "`useCallback`으로 loadUserInfo를 감싼 이유는 의존성이 바뀌지 않으면 같은 함수를 재사용하기 위해서이다.",
            "`if (!user?.id) return;`은 로그인 사용자 id가 없으면 API를 호출하지 않게 한다.",
            "`Promise.all`로 사용자 정보와 많이 틀린 단어를 동시에 요청한다.",
            "`fetchTopWrongWords(user.id).catch(() => [])`는 많이 틀린 단어 API가 실패해도 마이페이지 전체가 실패하지 않게 한다.",
            "`setUserInfo(info)`와 `setTopWrongWords(wrongWords)`로 화면 상태를 갱신한다.",
            "이 코드가 보여주는 중요한 패턴은 '핵심 데이터는 실패하면 에러 처리, 부가 데이터는 실패해도 빈 값 fallback'이다.",
        ],
        max_lines=70,
    )

    doc.add_paragraph("15. Firebase 설정 코드는 왜 이렇게 나뉘어 있는가", style="Heading 1")
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/config/FirebaseConfig.java",
        "public FirebaseApp firebaseApp",
        "firebaseApp: Firestore용 Firebase 앱 초기화",
        [
            "`FirebaseApp.getApps().isEmpty()`를 확인하는 이유는 FirebaseApp이 중복 초기화되면 예외가 날 수 있기 때문이다.",
            "`FirebaseOptions.builder().setCredentials(...)`로 Firestore 접근 권한을 가진 credential을 넣는다.",
            "`firebaseProjectId`가 있으면 명시적으로 projectId를 설정한다. 환경마다 프로젝트가 다를 수 있기 때문이다.",
            "`FirebaseApp.initializeApp`이 실제 Firebase Admin SDK 앱 인스턴스를 만든다.",
            "이 Bean이 있어야 `FirestoreClient.getFirestore(firebaseApp)`로 Firestore Bean을 만들 수 있다.",
        ],
        max_lines=60,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/config/FirebaseConfig.java",
        "public Storage googleCloudStorage",
        "googleCloudStorage: 수어 영상이 들어 있는 Storage 접근 객체 생성",
        [
            "Firestore와 Storage는 서로 다른 인증 정보를 쓸 수 있어서 별도 credential 해석 메서드를 둔다.",
            "`resolveStorageCredentialStream()`은 환경변수 JSON 또는 로컬 파일 경로에서 credential stream을 가져온다.",
            "`createScoped(\"https://www.googleapis.com/auth/cloud-platform\")`는 Storage 접근에 필요한 권한 범위를 지정한다.",
            "`StorageOptions.newBuilder().setCredentials(credentials).build().getService()`가 실제 Storage API 클라이언트를 만든다.",
            "이 Bean은 `StorageVideoCache` 생성자에 주입되어 단어별 영상 파일 검색에 사용된다.",
        ],
        max_lines=50,
    )
    add_code_block_from_method(
        doc,
        "backend/src/main/java/com/wow/signlanguage/config/SecurityConfig.java",
        "public SecurityFilterChain securityFilterChain",
        "SecurityConfig: 어떤 API를 공개할지 정하는 보안 필터 체인",
        [
            "`http.cors`는 다른 origin의 프론트엔드가 API를 호출할 수 있게 하는 기반 설정이다.",
            "`csrf.ignoringRequestMatchers(\"/api/**\", \"/translate\")`는 JSON API에 대해 CSRF 검사를 끈다. 모바일 앱/SPA API에서는 일반적인 설정이다.",
            "`requestMatchers(...).permitAll()`에 들어간 경로는 로그인 없이 접근 가능하다.",
            "`/api/auth/**`는 로그인 자체를 해야 하므로 공개되어야 한다.",
            "`/api/quiz/**`, `/api/users/**`도 현재는 공개되어 있다. 개발 편의상 열어 둔 구조로 보이며, 운영에서는 JWT 검증을 붙이는 것이 안전하다.",
            "`oauth2Login`은 Spring Security의 OAuth2 로그인 흐름을 활성화한다.",
            "이 파일은 기능 로직이 아니라 '요청이 컨트롤러까지 들어갈 수 있는지'를 결정하는 문지기 역할을 한다.",
        ],
        max_lines=90,
    )

    doc.add_paragraph("16. DTO와 record는 왜 많이 쓰이는가", style="Heading 1")
    add_code_block_from_range(
        doc,
        "backend/src/main/java/com/wow/signlanguage/translate/TranslateResponse.java",
        "public record TranslateResponse",
        "TranslateResponse: 번역 API가 프론트에 돌려주는 전체 데이터 구조",
        [
            "`record`는 Java에서 불변 데이터 전달 객체를 간단히 만들기 위한 문법이다.",
            "TranslateResponse는 단순히 영상 URL만 반환하지 않는다. 원본 입력, 단순화 문장, 토큰, 적용 규칙, 메타데이터, 재생 항목, unknown, noVideoWords를 모두 반환한다.",
            "프론트는 이 응답 하나로 영상 재생, 결과 설명, 누락 단어 표시, 북마크 저장을 모두 처리할 수 있다.",
            "DTO를 명확하게 만들면 API 계약이 분명해진다. 프론트와 백엔드가 어떤 필드를 주고받는지 코드만 봐도 알 수 있다.",
        ],
        max_lines=45,
    )
    add_code_block_from_range(
        doc,
        "backend/src/main/java/com/wow/signlanguage/quiz/dto/QuizSessionQuestionResponse.java",
        "public record QuizSessionQuestionResponse",
        "QuizSessionQuestionResponse: 퀴즈 한 문제를 프론트에 보내는 구조",
        [
            "`quizId`는 답 제출과 오답 저장에 사용되는 문제 식별자이다.",
            "`questionText`는 화면에 표시되는 문제 문장 또는 단어이다.",
            "`choices`는 보기 4개이다. 프론트 퀴즈 UI는 이 배열을 A/B/C/D로 렌더링한다.",
            "`videoUrl`은 문제의 핵심이다. 사용자는 이 영상을 보고 정답을 고른다.",
            "`attemptCount`, `correctCount`, `correctRate`, `difficultyLevel`은 학습 피드백과 난이도 표시를 위한 보조 지표이다.",
            "중요하게도 이 DTO에는 `correctChoiceId`가 없다. 정답은 제출 전에는 클라이언트에 보내지 않는 구조이다.",
        ],
        max_lines=45,
    )

    doc.add_paragraph("17. 기능별로 코드를 추적하는 방법", style="Heading 1")
    add_table(
        doc,
        ["기능", "프론트 시작점", "백엔드 시작점", "핵심 서비스"],
        [
            ["Google 로그인", "app/login.tsx", "auth/AuthController.java", "GoogleAuthService, JwtService, UserService"],
            ["번역", "app/translator.tsx", "api/TranslateController.java", "TranslationService, SignSentenceSimplifier, StorageVideoCache"],
            ["일반 퀴즈", "app/learn.tsx -> app/quiz.tsx", "quiz/QuizController.java", "QuizService"],
            ["오답 복습", "app/quiz.tsx?mode=wrong", "quiz/QuizController.java", "QuizService.getWrongSession, UserService"],
            ["오답노트 저장", "app/quiz.tsx", "user/UserController.java", "UserService.saveWrongNote"],
            ["번역 북마크", "app/translator.tsx", "user/UserController.java", "UserService.saveTranslatorBookmark"],
            ["마이페이지 통계", "app/mypage.tsx", "user/UserController.java", "UserService.getTopWrongWords, getRecentDailySolvedCounts"],
        ],
        widths=[1.2, 1.8, 1.8, 1.7],
    )
    add_para(
        doc,
        "코드를 읽을 때는 파일 이름만 외우는 것보다 '사용자 행동 하나가 어떤 파일들을 통과하는지'를 따라가면 훨씬 잘 이해된다. 예를 들어 '퀴즈에서 오답 저장'은 quiz.tsx의 버튼에서 시작해 wrong-note-saved.ts API 클라이언트, UserController.saveWrongNote, UserService.saveWrongNote, Firestore wrong_note_saved 컬렉션으로 이어진다.",
    )

    doc.add_paragraph("18. 이 프로젝트 코드에서 배울 수 있는 핵심 개념", style="Heading 1")
    concept_rows = [
        ["의존성 주입", "Controller가 Service를 직접 만들지 않고 생성자로 받는다. Spring이 객체 생성과 연결을 대신 관리한다."],
        ["계층 분리", "Controller는 HTTP 요청/응답, Service는 비즈니스 로직, DTO는 데이터 전달 구조를 맡는다."],
        ["정규화", "사용자 입력을 내부 사전 검색이 가능한 표준 단어로 바꾸는 과정이다."],
        ["fallback", "OpenAI 실패, 외부 API 실패, Storage URL 없음 같은 상황에서도 전체 기능이 멈추지 않게 대체 경로를 둔다."],
        ["캐싱", "StorageVideoCache가 성공/실패 검색 결과를 저장해 반복 Storage 조회를 줄인다."],
        ["비동기 처리", "프론트의 fetch와 useEffect는 네트워크 요청 결과가 나중에 도착하는 상황을 다룬다."],
        ["상태 관리", "useState, useMemo, useRef로 화면 데이터, 계산값, 애니메이션 값을 관리한다."],
        ["DTO 계약", "프론트와 백엔드가 어떤 JSON을 주고받는지 record/type으로 명확히 표현한다."],
        ["시간대 처리", "dailySolvedCounts는 KST 날짜 key를 기준으로 오늘 목표를 계산한다."],
        ["더블 버퍼링", "translator.tsx는 두 개의 영상 슬롯을 번갈아 사용해 다음 영상을 미리 준비한다."],
    ]
    add_table(doc, ["개념", "이 프로젝트에서의 의미"], concept_rows, widths=[1.4, 5.1])

    doc.add_paragraph("19. 수정할 때 조심해야 할 코드 지점", style="Heading 1")
    add_bullets(
        doc,
        [
            "TextNormalizer를 수정하면 번역 결과 전체가 바뀔 수 있다. 새 규칙이 기존 단어 매칭을 깨지 않는지 확인해야 한다.",
            "sign_dictionary.json의 word 값과 Firebase Storage 파일명 규칙은 서로 맞아야 한다. 사전에만 있고 영상이 없으면 noVideoWords가 늘어난다.",
            "QuizService.toSessionQuestion은 choices가 4개가 아니면 문제를 제외한다. 퀴즈 데이터 업로드 시 choices 개수를 꼭 맞춰야 한다.",
            "AuthContext의 저장 key를 바꾸면 기존 로그인 세션이 사라진 것처럼 보일 수 있다.",
            "base-url.ts를 잘못 수정하면 모든 API 호출과 영상 URL 재생이 동시에 깨질 수 있다.",
            "UserService.tryQuestion은 통계의 중심이다. totalQuestionNum, correctQuestionNum, dailySolvedCounts, incorrectQuestionCounts가 함께 갱신되므로 일부만 수정하면 마이페이지 수치가 어긋날 수 있다.",
            "SecurityConfig에서 `/api/users/**`가 공개되어 있으므로 운영 배포 전에는 JWT 검증과 uid 소유권 확인을 추가하는 것이 좋다.",
            "translator.tsx의 영상 슬롯 로직은 상태가 서로 연결되어 있다. currentClipIndex, activeSlot, ready 상태를 따로 수정하면 영상 전환이 꼬일 수 있다.",
        ],
    )

    doc.add_paragraph("20. 마무리: 코드를 한 문장으로 이해하기", style="Heading 1")
    add_para(
        doc,
        "이 프로젝트의 핵심은 사용자의 입력과 학습 행동을 'API 요청'으로 바꾸고, 백엔드가 그 요청을 '정규화된 단어, 영상 URL, 학습 기록'으로 변환한 뒤, 프론트엔드가 다시 그것을 '영상 재생과 학습 피드백'으로 보여주는 구조이다. 프론트엔드는 경험을 만들고, 백엔드는 판단과 데이터를 만들며, Firebase는 기억하고, Storage는 수어 영상을 제공한다.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
