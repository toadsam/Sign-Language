from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "sign-language-expected-qa-complete-ko.docx"

FONT = "Malgun Gothic"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x55, 0x65, 0x76)
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "AAB7C4"
CALLOUT = "EEF6FF"


def set_font(run, size=None, color=None, bold=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style(style, size=None, color=None, bold=None):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=TABLE_BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style(normal, size=10.5, color=INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    set_style(styles["Title"], size=23, color=RGBColor(0x0B, 0x25, 0x45), bold=True)
    styles["Title"].paragraph_format.space_after = Pt(8)

    set_style(styles["Subtitle"], size=10.5, color=MUTED)
    styles["Subtitle"].paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        set_style(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style(style, size=10.3, color=INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.text = "Sign-Language 예상 질문과 답변 총정리"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        set_font(run, size=8.5, color=MUTED)

    return doc


def para(doc, text="", style=None, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, bold=bold, color=color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def callout(doc, title, items):
    p = para(doc, title, bold=True, color=DARK_BLUE)
    shade_paragraph(p, CALLOUT)
    for item in items:
        p = bullet(doc, item)
        shade_paragraph(p, CALLOUT)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=9, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=8.8)
    para(doc)


SECTIONS = [
    {
        "title": "1. 프로젝트 개요와 기획 의도",
        "items": [
            (
                "이 프로젝트를 한 문장으로 설명하면 무엇인가요?",
                "수어를 처음 접하는 사용자도 영상 기반 퀴즈와 텍스트 입력 기반 수어 영상 변환을 통해 반복 학습할 수 있도록 만든 수어 학습 및 표현 보조 애플리케이션입니다.",
                "단순히 수어 영상을 나열하는 서비스가 아니라, 사용자가 영상을 보고 뜻을 맞히는 학습 흐름과 문장을 입력하면 단어별 수어 영상을 순서대로 보여주는 표현 보조 흐름을 함께 제공합니다. 그래서 학습 기능과 실사용 보조 기능이 연결된 것이 핵심입니다.",
            ),
            (
                "왜 수어 학습 서비스를 만들게 되었나요?",
                "수어는 청각장애인과 비장애인 사이의 의사소통 장벽을 줄이는 중요한 언어이지만, 초보자가 반복적으로 보고 맞히며 익힐 수 있는 접근성 높은 학습 도구가 부족하다고 보았기 때문입니다.",
                "기존 자료는 영상 사전처럼 단어를 찾아보는 데 초점이 있는 경우가 많습니다. 이 프로젝트는 사용자가 직접 문제를 풀고, 틀린 내용을 다시 보고, 문장 입력 결과를 영상으로 확인하는 방식으로 학습 지속성을 높이려는 목적이 있습니다.",
            ),
            (
                "기존 수어 사전이나 번역 서비스와 어떤 차이가 있나요?",
                "검색형 사전 기능만 제공하는 것이 아니라 퀴즈, 오답노트, 북마크, 일일 목표, 텍스트 기반 영상 재생을 하나의 학습 흐름으로 연결한 점이 다릅니다.",
                "또한 입력 문장을 그대로 단어 검색에 쓰지 않고, 조사 제거와 정규화, 외부 형태소 분석, OpenAI 보정 후보를 결합해 내부 수어 영상 사전에 더 잘 맞는 단어 목록을 찾도록 구성했습니다.",
            ),
            (
                "주요 사용자는 누구인가요?",
                "수어를 처음 배우는 일반 사용자, 수어 기초 단어를 반복 학습하려는 학습자, 간단한 표현을 수어 영상으로 확인하고 싶은 사용자입니다.",
                "초기 타깃은 전문 통역사가 아니라 초급 학습자입니다. 그래서 복잡한 문장 전체를 완벽하게 통역하기보다는 기본 단어와 짧은 문장을 영상으로 연결해 학습 이해를 돕는 방향으로 설계했습니다.",
            ),
            (
                "이 프로젝트의 핵심 가치는 무엇인가요?",
                "접근성, 반복 학습, 시각적 이해, 개인화 기록입니다.",
                "수어는 손 모양과 움직임을 보는 것이 중요합니다. 그래서 텍스트 설명보다 영상을 중심에 두었고, 사용자가 푼 문제와 틀린 문제를 기록해 다시 학습할 수 있도록 했습니다.",
            ),
            (
                "왜 퀴즈 기능과 번역 기능을 같이 넣었나요?",
                "퀴즈는 학습을 위한 기능이고, 번역은 사용자가 궁금한 표현을 바로 확인하기 위한 기능입니다. 둘을 함께 제공하면 학습과 실사용 확인이 자연스럽게 이어집니다.",
                "예를 들어 사용자는 번역기에서 본 표현을 북마크하고, 퀴즈에서 틀린 단어는 오답노트에 저장할 수 있습니다. 이처럼 단순 기능 나열이 아니라 반복 학습 루프를 만들고자 했습니다.",
            ),
            (
                "프로젝트 이름이나 서비스 주제와 기능이 어떻게 연결되나요?",
                "수어 학습이라는 주제에 맞춰 모든 핵심 기능이 '보고, 맞히고, 다시 보고, 저장하는' 흐름으로 연결됩니다.",
                "기술적으로는 영상 재생과 데이터 기록이 중심이고, 기획적으로는 초보자가 수어를 부담 없이 접하도록 만드는 것이 중심입니다.",
            ),
            (
                "서비스의 최소 기능 제품, 즉 MVP는 무엇인가요?",
                "수어 영상 퀴즈, 정답 판정, 사용자 학습 기록, 텍스트 입력 후 수어 영상 목록 재생이 MVP입니다.",
                "이 네 가지가 있으면 사용자는 영상을 보고 학습할 수 있고, 입력한 문장을 영상으로 확인할 수 있습니다. 오답노트, 북마크, 마이페이지 통계는 MVP 이후 학습 지속성을 높이는 확장 기능입니다.",
            ),
        ],
    },
    {
        "title": "2. 사용자 경험과 화면 기획",
        "items": [
            (
                "앱의 전체 사용자 흐름은 어떻게 되나요?",
                "사용자는 로그인 또는 게스트 모드로 앱에 들어온 뒤 홈에서 학습 현황을 보고, 학습 화면에서 퀴즈를 풀거나 번역기에서 문장을 입력하고, 마이페이지에서 기록과 오답, 북마크를 확인합니다.",
                "전체 흐름은 로그인 -> 홈 -> 학습/번역 -> 오답노트/북마크/마이페이지로 이어집니다. 사용자는 한 번 학습하고 끝나는 것이 아니라 저장된 기록을 바탕으로 다시 복습할 수 있습니다.",
            ),
            (
                "왜 게스트 모드를 넣었나요?",
                "로그인 장벽 없이 먼저 서비스를 체험하게 하기 위해서입니다.",
                "학습 서비스는 첫 진입 장벽이 높으면 사용자가 이탈하기 쉽습니다. 게스트 모드는 사용자가 기능을 먼저 경험하게 하고, 개인 기록 저장이 필요할 때 로그인을 유도할 수 있는 구조입니다.",
            ),
            (
                "홈 화면은 어떤 역할을 하나요?",
                "홈 화면은 사용자의 학습 상태를 요약하고 주요 기능으로 빠르게 이동하는 대시보드 역할을 합니다.",
                "오늘 목표, 최근 학습량, 빠른 실행 버튼을 통해 사용자가 다음 행동을 쉽게 선택하도록 돕습니다. 단순 시작 화면이 아니라 학습 동기를 유지하는 화면입니다.",
            ),
            (
                "학습 화면에서는 왜 코스 형태를 사용했나요?",
                "사용자가 어떤 퀴즈를 풀지 명확히 선택할 수 있게 하기 위해서입니다.",
                "기초 단어, 일상 회화, 오답 복습처럼 목적이 다른 학습 흐름을 카드로 구분하면 초보자도 무엇을 눌러야 할지 쉽게 이해할 수 있습니다.",
            ),
            (
                "퀴즈 화면의 핵심 UX는 무엇인가요?",
                "영상 중심 학습입니다. 사용자는 수어 영상을 보고 보기 중 정답을 선택하며, 즉시 정답/오답 피드백을 받습니다.",
                "수어는 텍스트보다 움직임을 보는 것이 중요하므로 영상 영역을 핵심으로 두었습니다. 답 제출 후에는 정답 여부, 정답 텍스트, 오답노트 저장 기능을 제공합니다.",
            ),
            (
                "번역기 화면에서 가장 중요한 UX는 무엇인가요?",
                "입력한 문장이 여러 수어 영상으로 자연스럽게 이어져 재생되는 경험입니다.",
                "사용자는 단어별로 따로 영상을 눌러 보지 않아도 됩니다. 앱이 번역 결과의 영상 목록을 순서대로 재생해 주기 때문에 문장 표현 흐름을 더 쉽게 이해할 수 있습니다.",
            ),
            (
                "오답노트 화면은 왜 필요한가요?",
                "틀린 문제를 다시 복습하는 반복 학습 구조를 만들기 위해서입니다.",
                "사용자가 틀린 문제를 단순히 지나치면 학습 효과가 떨어집니다. 오답노트는 사용자가 취약한 단어를 다시 보고 기억할 수 있게 합니다.",
            ),
            (
                "북마크 화면은 어떤 문제를 해결하나요?",
                "번역기에서 확인한 문장이나 표현을 나중에 다시 볼 수 있게 합니다.",
                "사용자는 궁금한 표현을 한 번 확인하고 끝내는 것이 아니라 저장해 두고 반복해서 확인할 수 있습니다. 이는 학습 지속성과 실사용성을 높입니다.",
            ),
            (
                "마이페이지는 단순 프로필 화면인가요?",
                "아닙니다. 마이페이지는 프로필뿐 아니라 정답률, 목표 달성, 많이 틀린 단어 등 학습 분석 정보를 보여주는 화면입니다.",
                "사용자는 자신의 학습 상태를 확인하고, 어떤 단어를 더 복습해야 하는지 파악할 수 있습니다.",
            ),
        ],
    },
    {
        "title": "3. 기술 스택 선정 이유",
        "items": [
            (
                "왜 프론트엔드에 Expo와 React Native를 사용했나요?",
                "모바일과 웹을 동시에 고려할 수 있고, 영상 재생, 라우팅, 인증, 저장소 같은 기능을 빠르게 구현할 수 있기 때문입니다.",
                "Expo는 개발 환경 구성이 비교적 쉽고, Expo Router를 통해 파일 기반 라우팅을 사용할 수 있습니다. React Native는 하나의 코드베이스로 모바일 UI를 만들 수 있어 학습 앱 구현에 적합합니다.",
            ),
            (
                "왜 백엔드에 Spring Boot를 사용했나요?",
                "REST API 서버를 안정적으로 만들 수 있고, Firebase, Google 인증, JWT, 보안 설정 등 서버 기능을 구조적으로 구현하기 좋기 때문입니다.",
                "Spring Boot는 Controller, Service, Configuration 계층을 명확하게 나눌 수 있어 프로젝트 규모가 커져도 유지보수가 쉽습니다.",
            ),
            (
                "왜 Firebase Firestore를 사용했나요?",
                "사용자 정보, 퀴즈 데이터, 오답노트, 북마크처럼 문서 기반으로 저장하기 좋은 데이터가 많기 때문입니다.",
                "Firestore는 컬렉션과 문서 구조가 직관적이고, 사용자별 하위 컬렉션을 만들기 쉬워 `users/{uid}/wrong_note_saved` 같은 구조에 잘 맞습니다.",
            ),
            (
                "왜 Firebase Storage 또는 Google Cloud Storage를 사용했나요?",
                "수어 영상 파일을 저장하고 URL로 제공해야 하기 때문입니다.",
                "영상은 데이터베이스 문서에 직접 넣기에는 크기가 크므로 Storage에 저장하고, Firestore에는 파일 정보나 fallback URL을 저장하는 구조가 적합합니다.",
            ),
            (
                "왜 JWT를 사용했나요?",
                "로그인 후 프론트엔드가 서버에 인증 상태를 전달할 수 있는 표준적인 토큰 방식이기 때문입니다.",
                "현재 코드에서는 토큰 발급 구조가 구현되어 있고, 추후 API 보호를 강화할 때 JWT 검증 필터를 붙여 사용자별 접근 제어로 확장할 수 있습니다.",
            ),
            (
                "왜 Google 로그인을 사용했나요?",
                "사용자가 별도 비밀번호를 만들지 않고 빠르게 가입하고 로그인할 수 있게 하기 위해서입니다.",
                "Google ID Token을 백엔드에서 검증하면 사용자의 이메일, 이름, 프로필 정보를 신뢰할 수 있는 방식으로 가져올 수 있습니다.",
            ),
            (
                "왜 OpenAI API를 번역 보정에 사용했나요?",
                "규칙 기반 처리와 외부 형태소 분석만으로는 한국어 활용형이나 복잡한 표현을 충분히 정규화하기 어렵기 때문입니다.",
                "다만 OpenAI 결과를 무조건 사용하지 않고 내부 사전 hit 수를 기준으로 rule, ETRI, OpenAI 후보 중 가장 적합한 토큰 흐름을 선택합니다.",
            ),
            (
                "왜 ETRI나 외부 사전을 같이 사용하나요?",
                "내부 사전에 없는 토큰을 보정하거나 문장 형태소 후보를 얻기 위해서입니다.",
                "사용자 입력은 사전 단어와 정확히 일치하지 않는 경우가 많습니다. 외부 사전과 형태소 분석은 단어 후보를 확장해 내부 사전 매칭률을 높이는 역할을 합니다.",
            ),
        ],
    },
    {
        "title": "4. 전체 아키텍처와 데이터 흐름",
        "items": [
            (
                "프로젝트의 전체 아키텍처는 어떻게 구성되어 있나요?",
                "Expo/React Native 프론트엔드, Spring Boot 백엔드, Firestore 데이터베이스, Firebase Storage 영상 저장소, Google/OpenAI/외부 사전 API로 구성됩니다.",
                "프론트엔드는 화면과 사용자 상호작용을 담당하고, 백엔드는 API와 비즈니스 로직을 담당합니다. Firestore는 사용자와 퀴즈 데이터를 저장하고, Storage는 수어 영상을 제공합니다.",
            ),
            (
                "프론트엔드와 백엔드는 어떻게 통신하나요?",
                "HTTP REST API로 통신합니다.",
                "프론트엔드의 `lib/api` 폴더에 API 호출 함수가 있고, 이 함수들이 Spring Boot의 Controller endpoint를 호출합니다. 예를 들어 `translateText()`는 `/translate`, `fetchQuizSession()`은 `/api/quiz/session`을 호출합니다.",
            ),
            (
                "백엔드 내부 계층은 어떻게 나뉘어 있나요?",
                "Controller, Service, DTO, Config, Storage, Dictionary 계층으로 나뉩니다.",
                "Controller는 요청을 받고 응답을 반환합니다. Service는 실제 비즈니스 로직을 처리합니다. DTO는 요청/응답 데이터 구조를 정의합니다. Config는 Firebase와 보안 설정을 담당합니다.",
            ),
            (
                "번역 요청의 전체 흐름은 어떻게 되나요?",
                "사용자가 문장을 입력하면 프론트가 `/translate`를 호출하고, 백엔드는 문장 단순화, 형태소 후보 생성, 내부 사전 매칭, Storage 영상 URL 조회를 거쳐 응답을 반환합니다.",
                "응답에는 원문, 단순화 문장, 토큰 목록, 적용 규칙, 재생 가능한 영상 목록, 영상이 없는 단어, 사전에 없는 단어가 포함됩니다.",
            ),
            (
                "퀴즈 요청의 전체 흐름은 어떻게 되나요?",
                "프론트가 `/api/quiz/session`을 호출하면 백엔드가 Firestore의 `quiz_items`에서 활성 문제를 읽고 영상 URL을 보정해 문제 목록을 반환합니다.",
                "사용자가 답을 제출하면 `/api/quiz/answer`에서 정답을 판정하고, 별도로 `/api/users/{uid}/tryQuestion`에서 사용자 학습 기록을 갱신합니다.",
            ),
            (
                "사용자 데이터는 어디에 저장되나요?",
                "Firestore의 `users/{uid}` 문서와 그 하위 컬렉션에 저장됩니다.",
                "사용자 기본 정보와 학습 통계는 사용자 문서에 저장되고, 번역 북마크와 저장된 오답노트는 사용자 하위 컬렉션에 저장됩니다.",
            ),
            (
                "영상 파일과 문제 데이터는 어떻게 연결되나요?",
                "문제나 번역 토큰의 단어를 기준으로 Storage에서 해당 단어로 시작하는 영상 파일을 찾습니다.",
                "예를 들어 정답 단어가 `학교`이면 StorageVideoCache가 Storage prefix에서 `학교`로 시작하는 파일을 찾고, 다운로드 URL을 만들어 반환합니다.",
            ),
            (
                "왜 백엔드에서 영상 URL을 찾아서 내려주나요?",
                "프론트가 Storage 구조나 인증 방식을 몰라도 되게 하기 위해서입니다.",
                "영상 위치와 URL 생성 방식은 서버가 관리하고, 프론트는 받은 URL을 재생하기만 하면 됩니다. 이렇게 하면 Storage 경로가 바뀌어도 프론트 수정 범위가 줄어듭니다.",
            ),
        ],
    },
    {
        "title": "5. 번역 기능 관련 예상 질문",
        "items": [
            (
                "이 프로젝트의 번역 기능은 실제 수어 통역인가요?",
                "완전한 자연어 수어 통역이라기보다는 입력 문장을 내부 수어 단어 사전에 맞는 토큰으로 단순화하고, 해당 단어들의 수어 영상을 순서대로 보여주는 기능입니다.",
                "현재 구현은 전문 통역 시스템보다 학습 보조와 표현 확인에 가깝습니다. 그래서 복잡한 문법과 문맥을 완벽히 처리하기보다는 짧은 문장과 기본 단어 중심으로 동작합니다.",
            ),
            (
                "문장을 어떤 방식으로 수어 영상으로 바꾸나요?",
                "규칙 기반 단순화, ETRI 형태소 분석, OpenAI 보정 후보를 만들고, 내부 수어 사전에 가장 많이 매칭되는 토큰 흐름을 선택합니다.",
                "선택된 토큰을 다시 내부 사전과 Storage 영상에 매칭해 재생 가능한 영상 목록을 만듭니다.",
            ),
            (
                "왜 문장을 그대로 단어 검색하지 않나요?",
                "한국어는 조사와 어미가 붙기 때문에 원문 단어가 사전 단어와 정확히 일치하지 않는 경우가 많기 때문입니다.",
                "예를 들어 `학교에서`, `갔어요`, `먹었어` 같은 표현은 사전에 `학교`, `가다`, `먹다` 형태로 있을 수 있습니다. 그래서 정규화와 조사 제거가 필요합니다.",
            ),
            (
                "SignSentenceSimplifier는 어떤 역할을 하나요?",
                "입력 문장을 토큰으로 나누고, 조사 제거, 시제 감지, 부정 감지, 의문 감지, 역할 분류, 수어식 순서 재배치를 수행합니다.",
                "시간, 장소, 주어, 목적어, 서술어, 의문어 등을 나눠 수어 영상 검색에 더 적합한 순서의 토큰 목록을 만듭니다.",
            ),
            (
                "TextNormalizer는 어떤 역할을 하나요?",
                "단어의 문장부호, 어미, 활용형, 예외 표현을 사전 검색에 유리한 대표형으로 바꿉니다.",
                "이 정규화가 잘될수록 내부 사전 hit 수가 올라가고, 결과적으로 재생 가능한 영상이 늘어납니다.",
            ),
            (
                "OpenAI 결과는 항상 사용되나요?",
                "아닙니다. OpenAI API key가 없거나 호출이 실패하면 사용되지 않고, 호출에 성공해도 내부 사전 hit 수가 더 높은 후보가 선택됩니다.",
                "TranslationService는 rule, ETRI, OpenAI 후보 중 내부 사전에 가장 잘 맞는 결과를 선택합니다. 따라서 OpenAI는 보조 후보이지 절대적인 기준이 아닙니다.",
            ),
            (
                "ETRI 형태소 분석은 왜 필요한가요?",
                "한국어 문장에서 의미 있는 형태소와 lemma를 추출해 규칙 기반 분석이 놓친 후보를 보완하기 위해서입니다.",
                "특히 동사나 형용사의 기본형 후보를 찾는 데 도움이 됩니다. 다만 API key가 없거나 실패하면 빈 결과를 반환하고 전체 번역은 계속 진행됩니다.",
            ),
            (
                "unknown과 noVideoWords의 차이는 무엇인가요?",
                "`unknown`은 내부 사전에도 없는 단어이고, `noVideoWords`는 내부 사전에는 있지만 Storage에서 영상을 찾지 못한 단어입니다.",
                "둘을 구분해야 개선 방향이 달라집니다. unknown은 사전 확장이 필요하고, noVideoWords는 영상 파일 업로드나 Storage 경로 확인이 필요합니다.",
            ),
            (
                "번역 결과에서 clips와 items는 왜 둘 다 있나요?",
                "`clips`는 실제 재생 가능한 영상만 담고, `items`는 영상 유무와 관계없이 전체 토큰 순서를 보존합니다.",
                "프론트엔드는 items를 사용해 어떤 단어가 영상으로 표현되었고 어떤 단어가 빠졌는지 더 정확히 처리할 수 있습니다.",
            ),
            (
                "문장 전체를 하나의 자연스러운 수어 영상으로 만들 수 있나요?",
                "현재 구조는 단어별 영상을 순서대로 재생하는 방식이므로 완전한 문장 단위 자연스러운 수어 생성과는 다릅니다.",
                "향후에는 문장 단위 수어 문법 변환, 3D 아바타 애니메이션 연결, 동작 보간 등을 추가해야 더 자연스러운 통역에 가까워질 수 있습니다.",
            ),
        ],
    },
    {
        "title": "6. 퀴즈와 학습 기능 관련 예상 질문",
        "items": [
            (
                "퀴즈 문제는 어디에서 가져오나요?",
                "Firestore의 `quiz_items` 컬렉션에서 가져옵니다.",
                "QuizService가 `isActive`가 true인 문제를 조회하고, category 조건이 있으면 해당 카테고리만 필터링한 뒤 랜덤으로 섞어 세션을 구성합니다.",
            ),
            (
                "퀴즈는 어떤 방식으로 랜덤화되나요?",
                "서버에서 활성 문제 목록을 가져온 뒤 `Collections.shuffle`로 순서를 섞고 요청한 개수만큼 잘라 반환합니다.",
                "이 방식은 구현이 단순하고 매번 다른 문제 순서를 제공할 수 있습니다. 다만 고급 추천 알고리즘은 아니므로 추후 난이도/오답률 기반 추천으로 확장할 수 있습니다.",
            ),
            (
                "정답은 프론트엔드에 미리 전달되나요?",
                "아닙니다. 퀴즈 세션 응답에는 정답 id가 포함되지 않고, 사용자가 답을 제출하면 서버가 정답을 판정합니다.",
                "이 구조는 클라이언트에서 정답을 쉽게 확인하거나 조작하는 것을 줄이기 위한 기본적인 설계입니다.",
            ),
            (
                "퀴즈 정답 판정은 어디에서 하나요?",
                "백엔드의 `QuizService.checkAnswer`에서 합니다.",
                "프론트가 보낸 quizId와 selectedChoiceId를 받아 Firestore의 correctChoiceId와 비교하고, 정답 여부와 정답 텍스트를 반환합니다.",
            ),
            (
                "퀴즈 통계는 어떻게 쌓이나요?",
                "문제별 통계는 QuizService에서 attempt_count, correct_count, wrong_count를 갱신하고, 사용자별 통계는 UserService.tryQuestion에서 갱신합니다.",
                "문제 통계와 사용자 통계를 분리하면 전체 문제 난이도 분석과 개인 학습 기록을 따로 관리할 수 있습니다.",
            ),
            (
                "오답 복습 퀴즈는 어떻게 만들어지나요?",
                "사용자 문서의 incorrectQuestionCounts 또는 incorrectQuestions를 기반으로 틀린 문제 id를 가져오고, 해당 문제들을 다시 세션으로 구성합니다.",
                "오답 횟수가 있는 경우 많이 틀린 순서로 우선순위를 줄 수 있어 취약 단어 복습에 도움이 됩니다.",
            ),
            (
                "퀴즈 문제에 영상이 없으면 어떻게 되나요?",
                "QuizService.toSessionQuestion에서 videoUrl을 찾지 못하면 해당 문제는 세션에 포함하지 않습니다.",
                "이 앱의 퀴즈는 영상을 보고 답을 맞히는 구조이므로 영상이 없는 문제는 정상 학습 경험을 제공할 수 없습니다.",
            ),
            (
                "퀴즈 카테고리는 어떻게 적용되나요?",
                "프론트 학습 화면에서 category query를 넘기고, 백엔드가 Firestore 문제 문서의 category와 비교해 필터링합니다.",
                "예를 들어 기초 단어는 `basic`, 일상 회화는 `daily` 같은 식으로 구분할 수 있습니다.",
            ),
            (
                "학습 목표는 어떻게 계산되나요?",
                "사용자 문서의 `dailySolvedCounts`에서 한국 날짜 기준 오늘 풀이 수를 가져와 목표치와 비교합니다.",
                "프론트의 daily-goal 유틸은 KST 날짜 key를 만들고 오늘 풀이 수, 목표 달성률, 연속 달성일을 계산합니다.",
            ),
            (
                "퀴즈 기능의 한계는 무엇인가요?",
                "현재는 문제를 랜덤으로 제공하는 방식이 중심이라 개인별 난이도 적응이나 추천 알고리즘은 제한적입니다.",
                "향후에는 사용자의 오답률, 학습 빈도, 단어 난이도, 최근 학습 시점을 반영한 개인화 추천으로 확장할 수 있습니다.",
            ),
        ],
    },
    {
        "title": "7. 사용자 데이터, 오답노트, 북마크 관련 질문",
        "items": [
            (
                "사용자 정보에는 어떤 값이 저장되나요?",
                "이름, 이메일, 전화번호, 소속, 프로필 이미지, 가입 완료 여부, 총 풀이 수, 정답 수, 틀린 문제 목록, 날짜별 풀이 수, 문제별 오답 횟수 등이 저장됩니다.",
                "이 정보는 홈, 마이페이지, 오답 복습, 일일 목표 계산에 사용됩니다.",
            ),
            (
                "오답노트와 incorrectQuestions는 같은 것인가요?",
                "완전히 같지는 않습니다. incorrectQuestions는 사용자가 틀린 문제 기록이고, wrong_note_saved는 사용자가 저장한 오답노트 항목입니다.",
                "즉 incorrectQuestions는 자동 학습 기록에 가깝고, wrong_note_saved는 사용자가 복습용으로 저장한 별도 목록에 가깝습니다.",
            ),
            (
                "오답노트 삭제는 실제 삭제인가요?",
                "코드상 저장된 오답노트 삭제는 완전 삭제보다 hidden 처리에 가까운 방식입니다.",
                "이 방식은 사용자가 화면에서는 삭제된 것처럼 보이게 하면서도 데이터 이력을 남길 수 있다는 장점이 있습니다.",
            ),
            (
                "번역기 북마크는 무엇을 저장하나요?",
                "사용자가 입력한 문장, 대표 단어, 대표 영상 URL, 저장 시각 등을 저장합니다.",
                "이 기능은 번역기에서 확인한 표현을 나중에 다시 볼 수 있게 해 줍니다.",
            ),
            (
                "게스트 사용자는 북마크를 사용할 수 있나요?",
                "서버 저장은 로그인 사용자 id가 필요하지만, translator 화면에는 로컬 저장 fallback 구조가 있습니다.",
                "따라서 로그인하지 않은 사용자도 일부 북마크 경험을 할 수 있지만, 기기나 브라우저 저장소에 의존하며 계정 간 동기화는 되지 않습니다.",
            ),
            (
                "마이페이지의 많이 틀린 단어는 어떻게 계산하나요?",
                "사용자 문서의 incorrectQuestionCounts를 기준으로 오답 횟수가 높은 문제를 정렬하고, quiz_items에서 해당 문제의 단어 정보를 가져옵니다.",
                "incorrectQuestionCounts가 비어 있으면 incorrectQuestions를 기반으로 기본 1회 오답으로 처리하는 fallback도 있습니다.",
            ),
            (
                "dailySolvedCounts는 왜 Map 형태인가요?",
                "날짜별 풀이 수를 key-value 형태로 저장하기 좋기 때문입니다.",
                "예를 들어 `2026-05-26: 7`처럼 날짜를 key로 두면 특정 날짜의 풀이 수를 빠르게 조회하고 최근 7일 그래프도 쉽게 만들 수 있습니다.",
            ),
            (
                "사용자 데이터 구조에서 개선할 점은 무엇인가요?",
                "퀴즈 북마크와 번역기 북마크 컬렉션 이름을 더 명확히 분리하고, 사용자 API에 인증/권한 검사를 추가하는 것이 좋습니다.",
                "현재는 기능 구현 중심 구조이므로 운영 서비스로 확장할 때 데이터 모델 정리와 접근 제어 강화가 필요합니다.",
            ),
        ],
    },
    {
        "title": "8. 인증과 보안 관련 예상 질문",
        "items": [
            (
                "로그인 과정은 어떻게 동작하나요?",
                "프론트엔드가 Google 로그인으로 idToken을 받고, 백엔드가 이 토큰을 검증한 뒤 JWT accessToken과 refreshToken을 발급합니다.",
                "프론트는 받은 토큰과 사용자 정보를 AuthContext에 저장하고 이후 앱 접근 상태를 판단합니다.",
            ),
            (
                "Google ID Token을 왜 백엔드에서 검증하나요?",
                "클라이언트가 보낸 토큰을 그대로 믿으면 안 되기 때문입니다.",
                "백엔드에서 Google 검증기를 통해 토큰이 실제 Google이 발급한 것인지, 대상 client id가 맞는지 확인해야 신뢰할 수 있습니다.",
            ),
            (
                "JWT에는 어떤 정보가 들어가나요?",
                "subject에는 Google 사용자 고유 id가 들어가고, claim에는 email, name, picture 등이 들어갑니다.",
                "발급 시각과 만료 시각도 포함되며, 서버의 secret key로 서명됩니다.",
            ),
            (
                "현재 보안상 가장 보완해야 할 점은 무엇인가요?",
                "사용자 API가 대부분 공개되어 있으므로 JWT 검증과 uid 소유권 확인을 추가해야 합니다.",
                "예를 들어 `/api/users/{uid}` 요청을 할 때 토큰 subject와 path의 uid가 같은지 확인해야 다른 사용자의 데이터 접근을 막을 수 있습니다.",
            ),
            (
                "CSRF를 끈 이유는 무엇인가요?",
                "모바일 앱과 SPA에서 JSON API를 호출하는 구조라 전통적인 서버 렌더링 폼 기반 CSRF 보호와 맞지 않기 때문입니다.",
                "다만 CSRF를 끈다고 해서 인증이 필요 없다는 뜻은 아니며, 운영 환경에서는 JWT 인증과 CORS 정책을 함께 관리해야 합니다.",
            ),
            (
                "토큰은 프론트에서 어디에 저장하나요?",
                "웹에서는 localStorage, 모바일에서는 Expo SecureStore에 저장합니다.",
                "SecureStore는 모바일에서 민감한 값을 저장하기에 더 적합합니다. 웹 localStorage는 XSS에 취약할 수 있으므로 운영 서비스에서는 추가 보안 대책이 필요합니다.",
            ),
            (
                "refreshToken은 현재 어떻게 사용되나요?",
                "코드상 발급은 되지만 refresh flow 전체가 완성되어 있다고 보기는 어렵습니다.",
                "향후 accessToken 만료 시 refreshToken으로 새 accessToken을 받는 endpoint와 토큰 폐기/재발급 정책을 추가할 수 있습니다.",
            ),
            (
                "개인정보는 어떤 것이 저장되나요?",
                "이름, 이메일, 전화번호, 소속, 프로필 이미지 URL 등이 저장될 수 있습니다.",
                "따라서 운영 서비스라면 개인정보 처리방침, 암호화/접근 제어, 데이터 삭제 요청 처리, 최소 수집 원칙을 고려해야 합니다.",
            ),
        ],
    },
    {
        "title": "9. Firebase와 Storage 관련 질문",
        "items": [
            (
                "Firestore와 Storage는 각각 어떤 역할인가요?",
                "Firestore는 사용자, 퀴즈, 오답노트, 북마크 같은 구조화된 데이터를 저장하고, Storage는 수어 영상 파일을 저장합니다.",
                "영상은 용량이 크기 때문에 문서 DB에 직접 저장하지 않고 Storage에 두고 URL만 사용합니다.",
            ),
            (
                "StorageVideoCache는 왜 필요한가요?",
                "Storage에서 영상 파일을 찾는 작업은 매번 호출하면 느리고 비용이 들 수 있기 때문입니다.",
                "한 번 찾은 단어의 URL은 cache에 저장하고, 못 찾은 단어도 missingWords에 저장해 같은 실패 검색을 반복하지 않습니다.",
            ),
            (
                "영상 파일명은 어떻게 단어와 매칭되나요?",
                "파일명에서 확장자를 제거하고, 언더스코어가 있으면 앞부분을 단어로 사용합니다.",
                "예를 들어 `학교_001.mp4`는 `학교`라는 단어로 매칭됩니다. 따라서 파일명 규칙이 사전 단어와 맞아야 합니다.",
            ),
            (
                "Storage prefix를 여러 개 쓰는 이유는 무엇인가요?",
                "기존 경로 오타나 경로 변경에 대응하기 위해서입니다.",
                "코드에는 `necessory_json_files/Model_videos/`와 `necessary_json_files/Model_videos/`가 모두 들어 있어 기존 데이터와 새 데이터 모두를 찾을 수 있게 합니다.",
            ),
            (
                "Firebase Storage 다운로드 URL은 어떻게 만들어지나요?",
                "파일 metadata에 다운로드 토큰이 있으면 Firebase media URL에 token을 붙이고, 없으면 signed URL을 생성하려고 시도합니다.",
                "토큰 또는 signed URL이 있어야 클라이언트가 영상을 재생할 수 있습니다. Storage 보안 규칙에 따라 direct URL은 실패할 수 있습니다.",
            ),
            (
                "Firestore 문서를 읽고 쓰는 과정에서 동기/비동기 처리는 어떻게 하나요?",
                "Java Firestore SDK는 Future 기반이므로 `.get().get()` 형태로 결과를 기다립니다.",
                "InterruptedException과 ExecutionException을 처리해야 하며, 일부 코드에서는 Spring의 ResponseStatusException으로 HTTP 오류를 변환합니다.",
            ),
            (
                "Firebase 설정은 로컬과 배포에서 어떻게 다르게 처리하나요?",
                "로컬에서는 service account 파일 경로를 사용할 수 있고, 배포에서는 JSON 내용을 환경변수나 Secret Manager로 주입할 수 있습니다.",
                "FirebaseConfig는 key json이 있으면 그것을 우선 사용하고, 없으면 파일 경로나 Application Default Credential을 사용합니다.",
            ),
            (
                "Storage 영상이 안 나올 때 어디를 확인해야 하나요?",
                "Storage bucket 이름, service account 권한, prefix 경로, 파일명 규칙, 다운로드 토큰, 프론트 URL 보정 로직을 확인해야 합니다.",
                "또한 `noVideoWords`에 단어가 잡히는지 보면 사전은 찾았지만 영상 URL을 못 찾은 상황인지 판단할 수 있습니다.",
            ),
        ],
    },
    {
        "title": "10. 프론트엔드 구현 관련 질문",
        "items": [
            (
                "Expo Router를 사용한 이유는 무엇인가요?",
                "파일 기반 라우팅으로 화면 구조를 직관적으로 관리할 수 있기 때문입니다.",
                "`app/home.tsx`는 `/home`, `app/quiz.tsx`는 `/quiz`처럼 파일 이름과 경로가 직접 연결되어 새 화면을 추가하거나 찾기 쉽습니다.",
            ),
            (
                "AuthContext는 왜 필요한가요?",
                "로그인 상태와 사용자 정보를 앱 전체에서 공유하기 위해서입니다.",
                "각 화면마다 토큰을 저장소에서 직접 읽으면 코드가 중복되고 상태가 불안정해집니다. AuthContext를 쓰면 `useAuth()`로 어디서든 동일한 인증 상태를 읽을 수 있습니다.",
            ),
            (
                "API 호출 코드를 lib/api로 분리한 이유는 무엇인가요?",
                "화면 컴포넌트가 fetch 세부사항을 직접 알지 않아도 되게 하기 위해서입니다.",
                "예를 들어 quiz.tsx는 `fetchQuizSession()`만 호출하면 되고, URL 생성, 응답 변환, 에러 처리는 api 파일이 담당합니다.",
            ),
            (
                "영상 재생은 어떤 라이브러리를 사용하나요?",
                "Expo Video를 사용합니다.",
                "퀴즈 화면에서는 현재 문제의 videoUrl을 재생하고, 번역기 화면에서는 여러 클립을 순차 재생하기 위해 두 개의 슬롯을 번갈아 사용하는 구조를 둡니다.",
            ),
            (
                "translator.tsx에서 상태가 많은 이유는 무엇인가요?",
                "번역 요청, 결과 표시, 북마크, 진행률, 웹/네이티브 영상 슬롯, 준비 상태, 애니메이션을 모두 관리해야 하기 때문입니다.",
                "특히 여러 영상을 부드럽게 이어 재생하려면 현재 클립 번호, 다음 영상 preload 상태, active slot 상태가 필요합니다.",
            ),
            (
                "웹과 모바일 저장소를 다르게 처리하는 이유는 무엇인가요?",
                "웹에서는 SecureStore를 사용할 수 없고, 모바일에서는 localStorage가 없기 때문입니다.",
                "AuthContext는 플랫폼별 차이를 storage 객체 안에 숨겨 화면 코드가 동일한 방식으로 저장소를 사용할 수 있게 합니다.",
            ),
            (
                "프론트에서 URL 보정이 필요한 이유는 무엇인가요?",
                "백엔드가 상대 경로나 로컬 URL을 반환할 수 있고, 배포 환경에서는 실제 API origin이 달라질 수 있기 때문입니다.",
                "`resolveBackendUrl`은 `/clips/a.mp4` 같은 경로를 실제 재생 가능한 절대 URL로 바꿉니다.",
            ),
            (
                "애니메이션은 기능적으로 꼭 필요한가요?",
                "핵심 기능은 아니지만 학습 앱의 반응성과 사용감을 높이는 역할을 합니다.",
                "정답 피드백, 로딩, 저장 완료, 화면 진입 애니메이션은 사용자가 앱이 반응하고 있다고 느끼게 해 이탈을 줄일 수 있습니다.",
            ),
        ],
    },
    {
        "title": "11. 백엔드 구현 관련 질문",
        "items": [
            (
                "Controller와 Service를 나눈 이유는 무엇인가요?",
                "HTTP 요청 처리와 실제 비즈니스 로직을 분리하기 위해서입니다.",
                "Controller는 요청을 받고 응답을 반환하는 얇은 계층이고, Service는 번역, 퀴즈 판정, 사용자 기록 갱신 같은 실제 기능을 처리합니다.",
            ),
            (
                "Spring의 의존성 주입은 어디에서 쓰이나요?",
                "Controller가 Service를 생성자로 받고, Service가 DictionaryLoader, StorageVideoCache, Firestore 등을 생성자로 받는 구조에서 사용됩니다.",
                "Spring이 필요한 객체를 자동으로 만들어 연결해 주기 때문에 코드가 느슨하게 결합되고 테스트나 유지보수가 쉬워집니다.",
            ),
            (
                "DTO record를 사용하는 이유는 무엇인가요?",
                "API 요청/응답 데이터 구조를 간결하고 명확하게 표현하기 위해서입니다.",
                "Java record는 불변 데이터 객체를 만들기 쉬우며, getter 성격의 메서드와 생성자 등이 자동으로 제공됩니다.",
            ),
            (
                "ResponseStatusException은 왜 사용하나요?",
                "서비스 로직에서 HTTP 상태 코드에 맞는 오류를 던질 수 있게 하기 위해서입니다.",
                "예를 들어 quizId가 없으면 BAD_REQUEST, 문서가 없으면 NOT_FOUND, Firestore 오류는 INTERNAL_SERVER_ERROR로 표현할 수 있습니다.",
            ),
            (
                "ObjectMapper는 어디에 쓰이나요?",
                "JSON 파일과 외부 API 응답을 Java 객체나 JsonNode로 변환할 때 사용됩니다.",
                "DictionaryLoader는 sign_dictionary.json을 읽고, ExternalLexiconApiClient와 OpenAiMorphologyNormalizerService는 API 응답 JSON을 파싱합니다.",
            ),
            (
                "RestTemplate은 어떤 역할인가요?",
                "외부 HTTP API를 호출하는 데 사용됩니다.",
                "ETRI, 우리말샘, 한국어기초사전, OpenAI API 호출에 사용되며, 실패 시 빈 결과를 반환하는 방식으로 전체 기능의 안정성을 유지합니다.",
            ),
            (
                "현재 백엔드에서 가장 복잡한 서비스는 무엇인가요?",
                "TranslationService와 UserService입니다.",
                "TranslationService는 여러 분석 후보와 사전/영상 매칭을 연결하고, UserService는 사용자 통계, 오답노트, 북마크, 프로필 등 많은 데이터 갱신을 담당합니다.",
            ),
            (
                "백엔드 성능 병목이 생긴다면 어디가 유력한가요?",
                "Storage 조회, 외부 API 호출, Firestore 반복 조회가 병목이 될 가능성이 큽니다.",
                "StorageVideoCache는 Storage 병목을 줄이기 위한 장치이고, 외부 API는 timeout과 fallback, Firestore는 batch나 캐싱 개선을 고려할 수 있습니다.",
            ),
        ],
    },
    {
        "title": "12. 데이터 모델과 API 설계 관련 질문",
        "items": [
            (
                "주요 Firestore 컬렉션은 무엇인가요?",
                "`quiz_items`, `users`, 그리고 사용자 하위의 `translator_bookmarks`, `wrong_note_saved`가 핵심입니다.",
                "quiz_items는 문제 데이터, users는 사용자 프로필과 통계, 하위 컬렉션은 개인별 저장 데이터를 관리합니다.",
            ),
            (
                "왜 사용자 하위 컬렉션을 사용하나요?",
                "북마크와 오답노트는 특정 사용자에게 종속된 데이터이기 때문입니다.",
                "`users/{uid}/wrong_note_saved/{quizId}`처럼 저장하면 사용자별 데이터 경계가 명확해지고 조회 경로도 직관적입니다.",
            ),
            (
                "퀴즈 문제 문서에는 어떤 필드가 필요한가요?",
                "questionText, choices, correctChoiceId, videoUrl, isActive, category, level, attempt_count, correct_count, wrong_count, difficulty_level 등이 필요합니다.",
                "choices는 4개여야 하며 correctChoiceId는 A/B/C/D 중 하나여야 정상적으로 퀴즈 UI와 정답 판정이 동작합니다.",
            ),
            (
                "API 응답에서 불필요한 정보를 숨기는 예가 있나요?",
                "퀴즈 세션 응답에는 정답 id를 넣지 않습니다.",
                "정답은 사용자가 제출한 뒤 서버가 판단합니다. 클라이언트에 정답을 미리 주면 사용자가 개발자 도구로 정답을 볼 수 있습니다.",
            ),
            (
                "왜 `/translate`는 `/api/translate`가 아니라 루트 경로인가요?",
                "현재 코드상 TranslateController가 클래스 레벨 `/api` 매핑 없이 `@PostMapping('/translate')`만 사용하기 때문입니다.",
                "일관성을 위해 장기적으로는 `/api/translate`로 옮길 수도 있지만, 그 경우 프론트의 translate.ts와 SecurityConfig도 함께 수정해야 합니다.",
            ),
            (
                "API 네이밍에서 개선할 부분은 무엇인가요?",
                "퀴즈 북마크와 번역 북마크의 컬렉션/엔드포인트 이름을 더 명확히 분리하면 좋습니다.",
                "현재 일부 북마크 구조는 기존 데이터 호환성을 고려한 것으로 보이지만, 장기적으로는 `quiz_bookmarks`와 `translator_bookmarks`를 분리하는 편이 이해하기 쉽습니다.",
            ),
            (
                "Map을 많이 쓰는 필드는 어떤 장단점이 있나요?",
                "dailySolvedCounts나 incorrectQuestionCounts처럼 key-value 통계에는 Map이 편리합니다.",
                "다만 key가 많아지면 문서 크기가 커지고 부분 업데이트나 쿼리가 어려워질 수 있으므로, 규모가 커지면 별도 컬렉션으로 분리할 수 있습니다.",
            ),
            (
                "API 오류 처리는 충분한가요?",
                "기본적인 bad request, not found, internal server error 처리는 있지만, 운영 수준의 공통 에러 응답 포맷은 더 정리할 수 있습니다.",
                "예를 들어 모든 API가 `{ code, message, details }` 형태의 일관된 오류 응답을 사용하면 프론트 처리와 디버깅이 쉬워집니다.",
            ),
        ],
    },
    {
        "title": "13. 테스트와 검증 관련 질문",
        "items": [
            (
                "현재 어떤 테스트가 필요한가요?",
                "TextNormalizer, SignSentenceSimplifier, TranslationService, QuizService, UserService에 대한 단위 테스트가 우선 필요합니다.",
                "이 기능들은 결과가 사용자 경험에 직접 영향을 주므로, 입력별 예상 결과를 고정해 회귀를 막는 테스트가 중요합니다.",
            ),
            (
                "번역 기능은 어떻게 테스트할 수 있나요?",
                "대표 입력 문장과 기대 토큰 목록, unknown/noVideoWords 결과를 비교하는 방식으로 테스트할 수 있습니다.",
                "OpenAI나 외부 API는 mock 처리하고, 내부 규칙 기반 결과와 사전 매칭 로직을 안정적으로 검증하는 것이 좋습니다.",
            ),
            (
                "퀴즈 기능은 어떻게 테스트할 수 있나요?",
                "정답 선택지와 오답 선택지를 넣었을 때 `isCorrect`, `correctChoiceId`, `correctChoiceText`가 올바른지 확인하면 됩니다.",
                "Firestore 의존성은 mock 또는 테스트용 emulator를 사용해 분리하는 것이 좋습니다.",
            ),
            (
                "사용자 통계 갱신은 어떻게 검증하나요?",
                "tryQuestion 호출 전후의 totalQuestionNum, correctQuestionNum, dailySolvedCounts, incorrectQuestionCounts를 비교하면 됩니다.",
                "정답 케이스와 오답 케이스를 따로 테스트해야 합니다.",
            ),
            (
                "프론트엔드에서는 어떤 테스트가 필요한가요?",
                "API 클라이언트의 URL 변환, AuthContext의 저장/복원, 퀴즈 화면의 답 제출 흐름, 번역기 재생 목록 세팅을 테스트할 수 있습니다.",
                "특히 resolveBackendUrl은 배포 환경에서 영상 재생 문제와 직결되므로 테스트 가치가 높습니다.",
            ),
            (
                "수동 시연 전 확인할 체크리스트는 무엇인가요?",
                "백엔드 서버 실행, 프론트 Expo 실행, Firebase credential 설정, quiz_items 데이터 존재, Storage 영상 URL 접근, Google client id 설정, JWT secret 설정을 확인해야 합니다.",
                "또한 로그인 계정으로 퀴즈를 풀고 마이페이지 통계가 변하는지, 오답노트와 북마크가 저장되는지 확인해야 합니다.",
            ),
            (
                "외부 API가 실패하면 테스트가 불안정하지 않나요?",
                "그렇기 때문에 외부 API 의존 테스트는 mock으로 분리하고, 실제 API 호출은 통합 테스트나 수동 점검으로 제한하는 것이 좋습니다.",
                "현재 코드도 외부 API 실패 시 빈 리스트를 반환하도록 되어 있어 기능 전체가 멈추지 않게 설계되어 있습니다.",
            ),
            (
                "품질을 더 높이려면 어떤 검증이 필요하나요?",
                "영상 재생 실패율, 번역 토큰 매칭률, unknown 발생률, noVideoWords 발생률, 퀴즈 정답률, 사용자 재방문율을 측정하면 좋습니다.",
                "기술적 로그와 사용자 행동 데이터를 함께 보면 어떤 사전과 영상을 우선 보강해야 할지 판단할 수 있습니다.",
            ),
        ],
    },
    {
        "title": "14. 배포와 운영 관련 질문",
        "items": [
            (
                "로컬에서 실행하려면 무엇이 필요한가요?",
                "Java 17, Gradle, Node/Expo 환경, Firebase credential, Google OAuth client id, JWT secret이 필요합니다.",
                "백엔드는 Gradle bootRun으로 실행하고, 프론트는 `npx expo start` 또는 npm script로 실행합니다.",
            ),
            (
                "환경변수는 어떤 것들이 중요한가요?",
                "GOOGLE_CLIENT_ID, JWT_SECRET, Firebase project/storage 관련 credential, ETRI/우리말샘/KRDict/OpenAI API key 등이 중요합니다.",
                "필수 기능만 보면 Google 로그인, JWT, Firebase/Storage 설정이 가장 중요하고, 외부 형태소/사전/OpenAI는 선택적 보정 기능입니다.",
            ),
            (
                "배포 시 가장 주의할 점은 무엇인가요?",
                "credential과 API key를 코드에 직접 넣지 않고 환경변수나 Secret Manager로 관리해야 합니다.",
                "또한 CORS, API base URL, Firebase Storage 접근 권한, JWT secret 길이와 보안을 확인해야 합니다.",
            ),
            (
                "프론트와 백엔드 주소가 달라지면 어떻게 처리하나요?",
                "프론트의 `EXPO_PUBLIC_API_BASE_URL` 환경변수로 백엔드 주소를 설정합니다.",
                "base-url.ts의 getBaseUrl과 resolveBackendUrl이 이 값을 기준으로 API와 영상 URL을 보정합니다.",
            ),
            (
                "운영 중 영상 파일을 추가하면 코드 수정이 필요한가요?",
                "파일명 규칙과 Storage prefix가 맞다면 코드 수정 없이 Storage에 파일을 추가하고 사전 데이터를 보강하면 됩니다.",
                "다만 새 단어가 sign_dictionary.json 또는 Firestore 문제 데이터와 연결되어야 번역/퀴즈에서 사용할 수 있습니다.",
            ),
            (
                "운영 모니터링은 무엇을 보면 좋나요?",
                "API 에러율, 번역 요청 수, unknown/noVideoWords 발생률, Storage 조회 실패 로그, 로그인 실패율, Firestore 오류를 보면 좋습니다.",
                "특히 영상 기반 서비스이므로 Storage URL 생성 실패와 영상 재생 실패는 중요한 모니터링 대상입니다.",
            ),
            (
                "데이터 백업은 어떻게 고려해야 하나요?",
                "Firestore 사용자 데이터와 quiz_items 데이터, Storage 영상 파일은 정기 백업이 필요합니다.",
                "특히 사용자 학습 기록과 오답노트는 개인화 서비스의 핵심 데이터이므로 삭제나 손상에 대비해야 합니다.",
            ),
            (
                "확장성이 필요한 경우 어디를 먼저 개선해야 하나요?",
                "외부 API 호출 캐싱, Firestore 조회 최적화, Storage URL 캐싱 만료 정책, API 인증 필터, 로깅/모니터링을 우선 개선하는 것이 좋습니다.",
                "사용자가 많아질수록 매 요청마다 외부 API와 Storage를 조회하는 구조는 병목이 될 수 있습니다.",
            ),
        ],
    },
    {
        "title": "15. 한계와 개선 방향 관련 질문",
        "items": [
            (
                "현재 프로젝트의 가장 큰 한계는 무엇인가요?",
                "완전한 수어 문법 번역이 아니라 단어 단위 영상 연결 방식이라는 점입니다.",
                "수어는 독립적인 문법과 비수지 표현이 중요하므로, 단어를 순서대로 보여주는 것만으로는 자연스러운 통역을 완전히 구현하기 어렵습니다.",
            ),
            (
                "번역 품질을 높이려면 무엇을 해야 하나요?",
                "수어 사전 확장, 문장 패턴 규칙 강화, unknown/noVideoWords 분석, 사용자 피드백 반영, 문장 단위 수어 표현 데이터 확보가 필요합니다.",
                "현재 구조에서는 내부 사전과 영상 데이터가 많아질수록 실제 재생 가능한 결과가 좋아집니다.",
            ),
            (
                "영상 기반 방식의 한계는 무엇인가요?",
                "단어별 영상이 자연스럽게 이어지지 않을 수 있고, 문맥에 따라 같은 단어도 다른 표현이 필요할 수 있습니다.",
                "향후에는 3D 아바타, 동작 보간, 문장 단위 애니메이션 생성으로 개선할 수 있습니다.",
            ),
            (
                "개인화 학습을 더 강화하려면 어떻게 해야 하나요?",
                "사용자의 오답률, 최근 풀이 시점, 학습 빈도, 단어 난이도를 바탕으로 문제를 추천하는 알고리즘을 추가할 수 있습니다.",
                "현재는 랜덤과 오답 복습 중심이므로, 다음 단계는 개인별 취약 단어 우선 추천입니다.",
            ),
            (
                "접근성을 더 높이려면 어떤 기능이 필요할까요?",
                "자막, 영상 속도 조절, 반복 재생, 큰 글씨 모드, 색 대비 개선, 키보드 접근성, 스크린리더 라벨 보강이 필요합니다.",
                "수어 학습 서비스는 접근성 자체가 중요한 주제이므로 UI 접근성도 기능만큼 중요합니다.",
            ),
            (
                "보안을 강화하려면 무엇을 추가해야 하나요?",
                "JWT 검증 필터, 사용자별 uid 권한 확인, refresh token 관리, rate limiting, CORS 제한, 민감정보 최소화가 필요합니다.",
                "현재 구현은 기능 중심에 가깝고, 운영 서비스로 가려면 인증/인가 계층을 강화해야 합니다.",
            ),
            (
                "데이터 모델을 개선한다면 어떻게 바꾸겠나요?",
                "퀴즈 북마크와 번역 북마크를 분리하고, dailySolvedCounts와 오답 기록을 규모에 따라 하위 컬렉션으로 분리할 수 있습니다.",
                "또한 quiz_items의 category, difficulty, tags를 더 체계화하면 추천과 검색 기능을 확장하기 쉽습니다.",
            ),
            (
                "향후 추가하고 싶은 기능은 무엇인가요?",
                "단어 검색 사전, 영상 느리게 보기, 사용자 발음/동작 따라 하기, 카메라 기반 수어 인식, 개인화 추천 퀴즈, 관리자용 문제 업로드 도구를 추가할 수 있습니다.",
                "특히 관리자 도구가 생기면 개발자가 직접 Firestore를 수정하지 않고도 퀴즈와 영상을 확장할 수 있습니다.",
            ),
            (
                "AI 기능을 더 발전시킨다면 어떻게 하겠나요?",
                "OpenAI를 단순 형태소 보정이 아니라 문장 의도 분석, 수어식 문장 변환 후보 생성, unknown 단어 대체 제안, 학습 설명 생성에 활용할 수 있습니다.",
                "다만 AI 결과는 검증 없이 그대로 쓰기보다 내부 사전, 수어 전문가 검토, 사용자 피드백과 함께 사용하는 것이 안전합니다.",
            ),
        ],
    },
    {
        "title": "16. 발표와 시연에서 나올 질문",
        "items": [
            (
                "시연 순서는 어떻게 잡는 것이 좋나요?",
                "로그인 또는 게스트 시작, 홈 확인, 학습 화면에서 퀴즈 풀이, 오답 저장, 번역기 문장 입력, 북마크 저장, 마이페이지 통계 확인 순서가 좋습니다.",
                "이 순서는 프로젝트의 핵심 가치인 학습, 표현 확인, 개인화 기록을 모두 보여줄 수 있습니다.",
            ),
            (
                "시연 중 가장 강조해야 할 기능은 무엇인가요?",
                "수어 영상 기반 퀴즈와 텍스트 입력 기반 영상 재생 기능입니다.",
                "이 두 기능이 프로젝트의 핵심입니다. 오답노트, 북마크, 마이페이지는 핵심 기능을 반복 학습으로 연결하는 보조 기능으로 설명하면 좋습니다.",
            ),
            (
                "심사자가 '실제 번역이 맞나요?'라고 물으면 어떻게 답하나요?",
                "현재는 전문 통역 수준의 완전한 수어 번역이 아니라, 입력 문장을 수어 사전 단어로 정규화하고 해당 수어 영상을 연결해 보여주는 학습 보조 기능이라고 답하면 됩니다.",
                "정직하게 한계를 인정하되, 규칙 기반 정규화와 외부 형태소 후보, OpenAI 보정을 결합해 매칭률을 높이려 했다는 기술적 시도를 함께 설명하면 좋습니다.",
            ),
            (
                "심사자가 '왜 카메라 인식은 없나요?'라고 물으면 어떻게 답하나요?",
                "이번 버전의 목표는 수어를 인식하는 것보다 초보자가 수어를 보고 학습하고 표현을 확인하는 흐름을 완성하는 것이었다고 답하면 됩니다.",
                "카메라 기반 수어 인식은 향후 확장 방향으로 적합하지만, 데이터셋과 모델 학습, 정확도 검증이 필요한 별도 큰 과제입니다.",
            ),
            (
                "데모가 실패할 때 대비 답변은 무엇인가요?",
                "영상이 나오지 않으면 Storage credential, bucket, 파일명 규칙, URL 접근 권한 문제일 가능성이 있다고 설명할 수 있습니다.",
                "번역 결과가 일부 누락되면 unknown과 noVideoWords를 구분해 사전 확장 또는 영상 업로드가 필요한 상태라고 설명하면 됩니다.",
            ),
            (
                "프로젝트에서 본인이 맡은 부분을 어떻게 설명하면 좋나요?",
                "맡은 역할에 따라 프론트엔드 UI/상태관리/API 연동, 백엔드 API/Firestore/번역 로직, 3D/영상 데이터 제작, 기획/문서화로 나누어 설명하면 됩니다.",
                "중요한 것은 단순히 파일을 만들었다고 말하는 것이 아니라 어떤 사용자 문제를 해결하기 위해 어떤 기능을 구현했는지 말하는 것입니다.",
            ),
            (
                "기술적으로 가장 어려웠던 점은 무엇이라고 답하면 좋나요?",
                "입력 문장을 내부 수어 사전과 영상 파일에 맞게 정규화하고, 영상 URL을 안정적으로 찾아 프론트에서 재생 가능한 형태로 연결하는 점이 어려웠다고 답할 수 있습니다.",
                "한국어 입력은 조사와 어미가 다양하고, 영상 파일은 Storage 경로/파일명 규칙과 맞아야 하므로 여러 계층을 함께 맞춰야 했습니다.",
            ),
            (
                "기획적으로 가장 고민한 점은 무엇인가요?",
                "초보자가 부담 없이 수어를 접하고 반복 학습할 수 있는 흐름을 만드는 것이었습니다.",
                "그래서 로그인 장벽을 낮추기 위한 게스트 모드, 학습 지속성을 위한 일일 목표, 취약점 복습을 위한 오답노트, 다시 보기 위한 북마크를 넣었습니다.",
            ),
        ],
    },
    {
        "title": "17. 꼬리질문 대비용 짧은 답변",
        "items": [
            (
                "왜 monorepo처럼 frontendcodes와 backend를 한 저장소에 두었나요?",
                "프론트와 백엔드가 강하게 연동되는 프로젝트라 한 저장소에서 API 계약과 기능 흐름을 같이 관리하기 쉽기 때문입니다.",
                "다만 규모가 커지면 별도 repository 또는 workspace 기반 monorepo 구조로 정리할 수 있습니다.",
            ),
            (
                "왜 README와 실제 구조가 일부 다른가요?",
                "프로젝트가 진행되면서 초기 설명과 실제 구현 구조가 달라진 부분이 있을 수 있습니다.",
                "그래서 추가 문서에서는 현재 코드 기준으로 Spring Boot Gradle 백엔드와 Expo React Native 프론트 구조를 다시 정리했습니다.",
            ),
            (
                "왜 코드에 일부 깨진 한글 주석이 보이나요?",
                "인코딩 문제로 보입니다.",
                "기능 자체와 별개로 문서화와 유지보수를 위해 UTF-8 인코딩으로 주석을 정리하는 작업이 필요합니다.",
            ),
            (
                "왜 일부 API가 Map을 직접 반환하나요?",
                "빠르게 기능을 구현하는 과정에서 유연하게 응답을 만들기 위해 사용한 것으로 보입니다.",
                "장기적으로는 명확한 response DTO로 통일하는 것이 타입 안정성과 문서화에 더 좋습니다.",
            ),
            (
                "왜 Spring Security를 쓰면서 API가 permitAll인가요?",
                "현재는 기능 구현과 연동 편의를 우선한 상태로 보입니다.",
                "운영 서비스로 발전시키려면 JWT 인증 필터와 사용자별 접근 제어를 추가해야 합니다.",
            ),
            (
                "왜 signed URL 만료가 365일인가요?",
                "클라이언트가 장기간 영상을 안정적으로 재생할 수 있게 하려는 목적입니다.",
                "보안을 더 중시하면 더 짧은 만료와 URL 재발급 구조를 고려할 수 있습니다.",
            ),
            (
                "왜 문제별 통계와 사용자별 통계를 둘 다 저장하나요?",
                "문제별 통계는 문제 난이도와 전체 정답률 분석에 필요하고, 사용자별 통계는 개인 학습 기록에 필요합니다.",
                "두 통계는 목적이 다르므로 분리해서 저장하는 것이 좋습니다.",
            ),
            (
                "왜 수어 영상을 서버 정적 파일이 아니라 Storage에서 찾나요?",
                "영상 파일은 많아지고 용량이 커질 수 있으므로 Storage가 더 적합합니다.",
                "또한 서버 배포와 영상 자산 관리를 분리할 수 있습니다.",
            ),
            (
                "왜 번역 결과를 바로 저장하지 않고 북마크 버튼을 따로 두었나요?",
                "모든 번역 요청을 저장하면 불필요한 데이터가 많아질 수 있기 때문입니다.",
                "사용자가 의미 있다고 판단한 문장만 저장하게 하는 것이 데이터 관리와 사용자 경험 면에서 더 적절합니다.",
            ),
            (
                "왜 일일 목표를 10개로 잡았나요?",
                "초보자가 부담 없이 반복 학습할 수 있는 작은 단위 목표로 볼 수 있습니다.",
                "향후에는 사용자가 목표치를 직접 설정하거나 난이도에 따라 목표를 조정할 수 있습니다.",
            ),
        ],
    },
]


def add_qa(doc, number, question, answer, detail):
    q = para(doc, f"Q{number}. {question}", style="Heading 3")
    for run in q.runs:
        set_font(run, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    r = p.add_run("핵심 답변: ")
    set_font(r, bold=True, color=BLUE)
    r = p.add_run(answer)
    set_font(r)

    p = doc.add_paragraph()
    r = p.add_run("자세한 설명: ")
    set_font(r, bold=True, color=BLUE)
    r = p.add_run(detail)
    set_font(r)

    p = doc.add_paragraph()
    r = p.add_run("답변 팁: ")
    set_font(r, bold=True, color=BLUE)
    r = p.add_run("질문을 받았을 때는 먼저 한 문장으로 결론을 말하고, 그다음 코드나 기능 흐름을 예로 들어 설명하면 좋습니다.")
    set_font(r, color=MUTED)


def build():
    doc = configure()
    doc.add_paragraph("Sign-Language 프로젝트 예상 질문과 답변 완전 정리", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    r = subtitle.add_run("기획, 기술, 구현, 보안, 데이터, 시연, 한계와 개선 방향까지 발표와 심사에서 나올 수 있는 질문을 Q&A 형태로 정리한 문서")
    set_font(r)

    para(
        doc,
        "이 문서는 Sign-Language 프로젝트에 대해 발표, 심사, 면접, 보고서 질의응답에서 나올 수 있는 질문을 최대한 넓게 모아 정리한 자료이다. 단순 암기용 답변이 아니라, 질문의 의도에 맞게 프로젝트 구조와 코드 구현 근거를 함께 설명할 수 있도록 구성했다.",
    )
    callout(
        doc,
        "사용 방법",
        [
            "발표 전에는 각 섹션의 핵심 답변만 먼저 읽어 전체 흐름을 잡는다.",
            "기술 질문이 나오면 관련 섹션의 자세한 설명을 바탕으로 Controller, Service, Firestore, Storage, 프론트 API 흐름을 연결해 말한다.",
            "한계 질문이 나오면 숨기지 말고 현재 범위와 향후 개선 방향을 함께 말한다.",
            "심사자가 꼬리질문을 하면 '현재 구현은 이렇게 되어 있고, 확장한다면 이렇게 할 수 있다'는 구조로 답한다.",
        ],
    )
    add_table(
        doc,
        ["분류", "질문 성격", "답변에서 강조할 점"],
        [
            ["기획", "왜 만들었는가, 누구를 위한 서비스인가", "수어 접근성, 초보자 반복 학습, 실사용 보조"],
            ["기술", "어떤 구조와 기술을 썼는가", "Expo, Spring Boot, Firebase, Storage, JWT, 형태소 정규화"],
            ["구현", "기능이 실제로 어떻게 동작하는가", "프론트 API 호출 -> 백엔드 Service -> Firestore/Storage"],
            ["한계", "무엇이 부족하고 어떻게 개선할 것인가", "단어 영상 연결 방식의 한계와 향후 3D/문장 단위 개선"],
            ["시연", "데모에서 무엇을 보여줄 것인가", "퀴즈, 번역, 오답노트, 북마크, 마이페이지 흐름"],
        ],
        widths=[1.1, 2.4, 3.0],
    )

    count = 1
    for section in SECTIONS:
        doc.add_paragraph(section["title"], style="Heading 1")
        for question, answer, detail in section["items"]:
            add_qa(doc, count, question, answer, detail)
            count += 1

    doc.add_paragraph("마지막 정리", style="Heading 1")
    para(
        doc,
        "이 프로젝트를 설명할 때 가장 중요한 말은 '수어 학습을 영상 중심으로 쉽게 접근하게 하고, 사용자의 학습 기록을 바탕으로 반복 학습을 돕는 서비스'라는 점이다. 기술적으로는 프론트엔드가 사용자 경험을 만들고, 백엔드가 번역/퀴즈/사용자 데이터를 처리하며, Firestore와 Storage가 데이터와 영상을 담당한다. 질문을 받을 때는 항상 이 큰 구조로 돌아와 답하면 안정적으로 설명할 수 있다.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
