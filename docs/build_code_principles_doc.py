from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "sign-language-code-principles-complete-ko.md"
OUTPUT = ROOT / "docs" / "sign-language-code-principles-complete-ko.docx"

FONT = "Malgun Gothic"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
CODE_SHADE = "F4F6F9"
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "AAB7C4"


def set_east_asian_font(run, font_name=FONT):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, size_pt=None, color=None, bold=None):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=TABLE_BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def parse_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_east_asian_font(run, "Consolas")
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x30, 0x37, 0x3D)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_east_asian_font(run)
            run.bold = True
        else:
            run = paragraph.add_run(part)
            set_east_asian_font(run)


def clean_table_cell(value):
    value = value.strip()
    if value.startswith("`") and value.endswith("`"):
        value = value[1:-1]
    return value.replace("\\|", "|")


def clean_heading_text(value):
    return value.strip().replace("`", "").replace("**", "")


def is_table_separator(line):
    stripped = line.strip()
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return False
    return all(set(part.strip()) <= {"-", ":"} and "-" in part for part in stripped.strip("|").split("|"))


def is_table_line(line):
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]


def split_table_row(line):
    return [clean_table_cell(part) for part in line.strip().strip("|").split("|")]


def add_markdown_table(doc, rows):
    if len(rows) < 2:
        return
    headers = split_table_row(rows[0])
    body = [split_table_row(row) for row in rows[2:] if is_table_line(row)]
    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)

    col_width = 6.5 / len(headers)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(col_width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_east_asian_font(run)
        run.bold = True
        run.font.size = Pt(9.5)
    set_repeat_table_header(table.rows[0])

    for raw_row in body:
        normalized = raw_row[: len(headers)] + [""] * max(0, len(headers) - len(raw_row))
        cells = table.add_row().cells
        for idx, value in enumerate(normalized[: len(headers)]):
            cell = cells[idx]
            cell.width = Inches(col_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            parse_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(9)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_code_block(doc, code_lines):
    if not code_lines:
        return
    text = "\n".join(code_lines)
    p = doc.add_paragraph(style="CodeBlock")
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.06)

    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_SHADE)
    p_pr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DADCE0")
    borders.append(bottom)
    p_pr.append(borders)


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, RGBColor(0, 0, 0))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    set_style_font(title, 24, RGBColor(0x0B, 0x25, 0x45), True)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.1

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, 10.5, MUTED, False)
    subtitle.paragraph_format.space_after = Pt(16)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, 11, RGBColor(0, 0, 0))
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "CodeBlock" not in [s.name for s in styles]:
        code_style = styles.add_style("CodeBlock", 1)
    else:
        code_style = styles["CodeBlock"]
    set_style_font(code_style, 9, RGBColor(0x30, 0x37, 0x3D))
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.text = "Sign-Language 코드 구조와 동작 원리"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_east_asian_font(run)
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        set_east_asian_font(run)
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED

    return doc


def build_docx():
    doc = configure_document()
    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    doc.add_paragraph("Sign-Language 프로젝트 코드 구조와 동작 원리 완전 상세 설명서", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Spring Boot 백엔드, Expo/React Native 프론트엔드, Firebase 데이터 흐름, 번역 파이프라인, 퀴즈/오답/북마크 구조를 코드 기준으로 정리한 상세 기술 문서")
    set_east_asian_font(subtitle.runs[0])

    i = 1 if lines and lines[0].startswith("# ") else 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if is_table_line(line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_heading_text(heading_match.group(2))
            if level == 1:
                p = doc.add_paragraph(text, style="Heading 1")
            elif level == 2:
                p = doc.add_paragraph(text, style="Heading 1")
            elif level == 3:
                p = doc.add_paragraph(text, style="Heading 2")
            else:
                p = doc.add_paragraph(text, style="Heading 3")
            for run in p.runs:
                set_east_asian_font(run)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, numbered.group(1))
            i += 1
            continue
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, bullet.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
